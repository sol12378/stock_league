from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

from src.config import AppConfig, load_config
from src.utils.logging import setup_logger


def _markdown_table(df: pd.DataFrame, max_rows: int = 20) -> str:
    if df.empty:
        return "_No data available._"
    show = df.head(max_rows).copy()
    show = show.fillna("")
    columns = list(show.columns)
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join(["---"] * len(columns)) + " |",
    ]
    for _, row in show.iterrows():
        lines.append("| " + " | ".join(str(row[col]) for col in columns) + " |")
    return "\n".join(lines)


def _screening_summary_for_report(summary: pd.DataFrame) -> pd.DataFrame:
    labels = {
        "universe": "母集団",
        "price_available": "株価取得可能",
        "liquid_20m_60d": "流動性条件通過",
        "investment_eligible": "投資適格性通過",
        "scored": "スコア算出対象",
        "candidates_top80": "統合スコア上位80社",
        "portfolio_candidates": "最終20銘柄",
    }
    table = summary.copy()
    table["stage"] = table["stage"].map(labels).fillna(table["stage"])
    return table.rename(columns={"stage": "段階", "count": "社数"})


def _company_display_name(df: pd.DataFrame) -> pd.Series:
    fallback = df["company_name"] if "company_name" in df.columns else df["ticker"]
    if "company_name_ja" not in df.columns:
        return fallback
    japanese = df["company_name_ja"].fillna("").astype(str).str.strip()
    return japanese.where(japanese.str.len() > 0, fallback)


def _portfolio_report_table(portfolio: pd.DataFrame) -> pd.DataFrame:
    table = portfolio.copy()
    table["企業名"] = _company_display_name(table)
    return table[
        [
            "code",
            "ticker",
            "企業名",
            "sector_33",
            "category",
            "previous_close",
            "shares",
            "actual_investment",
            "actual_weight",
        ]
    ].rename(
        columns={
            "code": "コード",
            "ticker": "Ticker",
            "sector_33": "業種",
            "category": "カテゴリ",
            "previous_close": "前営業日終値",
            "shares": "株数",
            "actual_investment": "投資額",
            "actual_weight": "比率",
        }
    )


def _top_candidates_report_table(top80: pd.DataFrame) -> pd.DataFrame:
    table = top80.copy()
    table["企業名"] = _company_display_name(table)
    return table[
        [
            "code",
            "ticker",
            "企業名",
            "sector_33",
            "category",
            "adjusted_bb_score",
        ]
    ].rename(
        columns={
            "code": "コード",
            "ticker": "Ticker",
            "sector_33": "業種",
            "category": "カテゴリ",
            "adjusted_bb_score": "Adjusted BB Score",
        }
    )


def generate_markdown(config: AppConfig) -> Path:
    logger = setup_logger("generate_markdown", config.logs_dir)
    summary = pd.read_csv(config.data_processed_dir / "screening_summary.csv")
    portfolio = pd.read_csv(config.data_processed_dir / "portfolio.csv")
    performance = pd.read_csv(config.data_processed_dir / "performance_summary.csv")
    top80 = pd.read_csv(config.data_processed_dir / "candidates_top80.csv")
    edinet_path = config.data_processed_dir / "edinet_documents.csv"
    edinet_count = len(pd.read_csv(edinet_path)) if edinet_path.exists() else 0
    edinet_note = (
        "今回の自動実行ではEDINET APIの認証が通らず、有価証券報告書XBRLによる財務補完は0件であった。"
        "そのため、初版のスコアはJPX上場銘柄一覧、yfinance価格データ、yfinanceで取得できた一部指標、業種・企業名ベースのテーマ露出を中心に構成している。"
        if edinet_count == 0
        else f"EDINETから有価証券報告書候補を{edinet_count}件取得し、取得可能なXBRL項目を財務指標に反映した。"
    )

    lines = [
        "# BEYOND BUFFETT 分析レポート草案",
        "",
        "## 1. 分析の位置づけ",
        "",
        "本分析は、公開データを用いて東証上場企業をスクリーニングし、既存の競争優位、変革余地、Future Moat、バリュエーション規律を統合した BEYOND BUFFETT Score によって候補企業を選定したものである。",
        "",
        "本資料は日経STOCKリーグ提出レポートの材料であり、個別銘柄の投資助言を目的としない。",
        "",
        "## 2. スクリーニング結果",
        "",
        _markdown_table(_screening_summary_for_report(summary)),
        "",
        "### EDINET取得状況",
        "",
        edinet_note,
        "",
        "## 3. スコア上位候補",
        "",
        _markdown_table(_top_candidates_report_table(top80), max_rows=30),
        "",
        "![Score Distribution](../figures/score_distribution.png)",
        "",
        "## 4. 500万円ポートフォリオ",
        "",
        _markdown_table(_portfolio_report_table(portfolio), max_rows=30),
        "",
        "![Sector Allocation](../figures/sector_allocation.png)",
        "",
        "![Category Allocation](../figures/category_allocation.png)",
        "",
        "## 5. バックテスト・リスク分析",
        "",
        _markdown_table(performance),
        "",
        "![Cumulative Return](../figures/cumulative_return.png)",
        "",
        "![Drawdown](../figures/drawdown.png)",
        "",
        "![Contribution by Stock](../figures/contribution_by_stock.png)",
        "",
        "![Risk Return Scatter](../figures/risk_return_scatter.png)",
        "",
        "## 6. データ上の限界",
        "",
        "本分析では、無料・公開データを中心に用いたため、一部の財務指標やガバナンス指標については欠損や基準日の差異が存在する。そのため、一次スクリーニングでは取得可能な定量指標を用い、最終選定段階では各社の有価証券報告書、決算説明資料、中期経営計画を確認することで、データ上の限界を補完する必要がある。また、バックテストは現在上場している企業を対象としているため、上場廃止企業を含まない生存者バイアスが存在する可能性がある。",
        "",
    ]

    output = config.reports_draft_dir / "report_draft.md"
    output.write_text("\n".join(lines), encoding="utf-8")
    logger.info("Wrote %s", output)
    return output


def _add_dataframe_table(document: Document, df: pd.DataFrame, max_rows: int = 20) -> None:
    show = df.head(max_rows).fillna("")
    table = document.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(show.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(show.columns):
            cells[idx].text = str(row[column])


def generate_docx(config: AppConfig) -> Path:
    summary = pd.read_csv(config.data_processed_dir / "screening_summary.csv")
    portfolio = pd.read_csv(config.data_processed_dir / "portfolio.csv")
    performance = pd.read_csv(config.data_processed_dir / "performance_summary.csv")
    edinet_path = config.data_processed_dir / "edinet_documents.csv"
    edinet_count = len(pd.read_csv(edinet_path)) if edinet_path.exists() else 0

    doc = Document()
    doc.add_heading("BEYOND BUFFETT 分析レポート草案", 0)
    doc.add_heading("スクリーニング結果", level=1)
    _add_dataframe_table(doc, _screening_summary_for_report(summary))
    doc.add_heading("EDINET取得状況", level=2)
    if edinet_count == 0:
        doc.add_paragraph(
            "今回の自動実行ではEDINET APIの認証が通らず、有価証券報告書XBRLによる財務補完は0件であった。"
            "初版のスコアはJPX上場銘柄一覧、yfinance価格データ、yfinanceで取得できた一部指標を中心に構成している。"
        )
    else:
        doc.add_paragraph(f"EDINETから有価証券報告書候補を{edinet_count}件取得した。")
    doc.add_heading("500万円ポートフォリオ", level=1)
    _add_dataframe_table(
        doc,
        _portfolio_report_table(portfolio).drop(columns=["前営業日終値"]),
        max_rows=30,
    )
    doc.add_heading("パフォーマンス", level=1)
    _add_dataframe_table(doc, performance)
    doc.add_heading("図表", level=1)
    for name in [
        "cumulative_return.png",
        "drawdown.png",
        "sector_allocation.png",
        "category_allocation.png",
        "contribution_by_stock.png",
    ]:
        path = config.reports_figures_dir / name
        if path.exists():
            doc.add_paragraph(name)
            doc.add_picture(str(path), width=Inches(6.0))
    doc.add_heading("データ上の限界", level=1)
    doc.add_paragraph(
        "本分析では、無料・公開データを中心に用いたため、一部の財務指標やガバナンス指標については欠損や基準日の差異が存在する。"
        "バックテストは現在上場している企業を対象としているため、上場廃止企業を含まない生存者バイアスが存在する可能性がある。"
    )
    output = config.reports_draft_dir / "beyond_buffett_report.docx"
    doc.save(output)
    return output


def main() -> None:
    config = load_config()
    generate_markdown(config)
    generate_docx(config)


if __name__ == "__main__":
    main()
