from __future__ import annotations

from pathlib import Path
import sys

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.config import load_config  # noqa: E402
from src.portfolio.metrics import performance_row  # noqa: E402
from src.report.generate_pdf import PdfPageWriter, _main_risk, _selection_reason  # noqa: E402
from src.utils.prices import repair_split_jumps  # noqa: E402


OUTPUT_MD = ROOT / "reports" / "draft" / "submission_report_draft.md"
OUTPUT_DOCX = ROOT / "reports" / "draft" / "submission_report_draft.docx"
OUTPUT_PDF = ROOT / "reports" / "draft" / "submission_report_draft.pdf"
SENSITIVITY_CSV = ROOT / "reports" / "tables" / "contribution_sensitivity.csv"


def pct(value: object, digits: int = 1) -> str:
    try:
        if pd.isna(value):
            return "-"
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return "-"


def num(value: object, digits: int = 2) -> str:
    try:
        if pd.isna(value):
            return "-"
        return f"{float(value):.{digits}f}"
    except Exception:
        return "-"


def yen(value: object) -> str:
    try:
        if pd.isna(value):
            return "-"
        return f"{int(round(float(value))):,}円"
    except Exception:
        return "-"


def md_table(df: pd.DataFrame, max_rows: int | None = None) -> str:
    if max_rows is not None:
        df = df.head(max_rows)
    if df.empty:
        return "_該当データなし_"
    columns = [str(column) for column in df.columns]
    rows = []
    for _, row in df.iterrows():
        rows.append(["" if pd.isna(value) else str(value) for value in row.tolist()])
    out = ["| " + " | ".join(columns) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    out.extend("| " + " | ".join(values) + " |" for values in rows)
    return "\n".join(out)


def screening_summary_for_report(summary: pd.DataFrame) -> pd.DataFrame:
    labels = {
        "universe": "母集団",
        "price_available": "株価取得可能",
        "liquid_20m_60d": "流動性条件通過",
        "investment_eligible": "投資適格性通過",
        "scored": "スコア算出対象",
        "candidates_top80": "統合スコア上位80社",
        "portfolio_candidates": "最終20銘柄",
    }
    table = summary.copy()
    table["stage"] = table["stage"].map(labels).fillna(table["stage"])
    return table.rename(columns={"stage": "段階", "count": "社数"})


def display_name(row: pd.Series) -> str:
    name = str(row.get("company_name_ja", "") or "").strip()
    return name or str(row.get("company_name", "") or row.get("ticker", "")).strip()


def category_theme(category: str) -> str:
    if category == "Core Moat":
        return "既存の競争優位・収益基盤・ブランドを評価する安定枠"
    if category == "Transformation Moat":
        return "資本効率改善、低PBR是正、事業構造改革による再評価枠"
    if category == "Future Moat":
        return "AI、半導体、光通信、電力、データ基盤など将来の産業構造への接続枠"
    return "市場に見落とされやすい発見枠"


def future_moat_axis(row: pd.Series) -> str:
    code = str(row.get("code", ""))
    mapping = {
        "6777": "光通信・光計測・データセンター周辺",
        "6627": "AI半導体需要を支える半導体検査・後工程",
        "6524": "光通信・電子部品・AIインフラ周辺部材",
        "6387": "半導体製造装置・化合物半導体",
        "6800": "通信・車載・電子部品",
        "8050": "精密機器・電子デバイス・ブランド資産",
        "9984": "AI投資、半導体、デジタルプラットフォーム",
        "9501": "電力インフラ・データセンター需要",
        "9513": "電力インフラ・脱炭素投資",
    }
    if code in mapping:
        return mapping[code]
    sector = str(row.get("sector_33", ""))
    if "Electric Appliances" in sector:
        return "電子部品・半導体・通信インフラ"
    if "Machinery" in sector:
        return "省人化・半導体製造・設備投資"
    if "Electric Power" in sector:
        return "電力インフラ・データセンター需要"
    if "Information" in sector:
        return "AI・ソフトウェア・データ基盤"
    return "資本効率改善または既存事業の堀"


def load_tables() -> dict[str, pd.DataFrame]:
    base = ROOT / "data" / "processed"
    return {
        "portfolio": pd.read_csv(base / "portfolio.csv", dtype={"code": str}),
        "performance": pd.read_csv(base / "performance_summary.csv"),
        "ablation": pd.read_csv(base / "ablation_performance.csv"),
        "category": pd.read_csv(base / "category_returns.csv"),
        "market": pd.read_csv(base / "screening_by_market.csv"),
        "sector": pd.read_csv(base / "screening_by_sector.csv"),
        "correlation": pd.read_csv(base / "score_correlation.csv"),
        "summary": screening_summary_for_report(pd.read_csv(base / "screening_summary.csv")),
        "contribution": pd.read_csv(base / "contribution_by_stock.csv"),
    }


def load_returns(tickers: list[str]) -> pd.DataFrame:
    prices = pd.read_parquet(ROOT / "data" / "processed" / "prices_daily.parquet")
    prices = prices.copy()
    prices["price_for_return"] = prices["adj_close"].fillna(prices["close"])
    pivot = (
        prices[prices["ticker"].isin(tickers)]
        .pivot_table(index="date", columns="ticker", values="price_for_return", aggfunc="last")
        .sort_index()
    )
    pivot.index = pd.to_datetime(pivot.index)
    pivot = pivot.apply(repair_split_jumps, axis=0)
    return pivot.pct_change().replace([np.inf, -np.inf], np.nan)


def build_sensitivity(portfolio: pd.DataFrame, contribution: pd.DataFrame) -> pd.DataFrame:
    config = load_config()
    tickers = portfolio["ticker"].astype(str).tolist()
    returns = load_returns(tickers + [config.topix_proxy])
    weights = portfolio.set_index("ticker")["actual_weight"].astype(float)
    top_contributors = contribution.sort_values("contribution", ascending=False)["ticker"].astype(str).tolist()
    scenarios = {
        "final BEYOND BUFFETT portfolio": [],
        "santec Holdings excluded": ["6777.T"],
        "top 3 contributors excluded": top_contributors[:3],
        "top 5 contributors excluded": top_contributors[:5],
    }
    rows: list[dict[str, object]] = []
    benchmark = returns[config.topix_proxy].dropna() if config.topix_proxy in returns.columns else None
    for label, excluded in scenarios.items():
        keep = [ticker for ticker in tickers if ticker in returns.columns and ticker not in excluded]
        scenario_weights = weights.reindex(keep).fillna(0)
        if scenario_weights.sum() <= 0:
            continue
        scenario_weights = scenario_weights / scenario_weights.sum()
        series = returns[keep].mul(scenario_weights, axis=1).sum(axis=1, min_count=1).fillna(0)
        row = performance_row(label, series, benchmark)
        row["label"] = row.pop("name")
        row["excluded"] = ", ".join(excluded) if excluded else "-"
        rows.append(row)
    out = pd.DataFrame(rows)
    SENSITIVITY_CSV.parent.mkdir(parents=True, exist_ok=True)
    out.to_csv(SENSITIVITY_CSV, index=False)
    return out


def portfolio_reason_table(portfolio: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in portfolio.iterrows():
        rows.append(
            {
                "code": str(row["code"]),
                "company": display_name(row),
                "category": row["category"],
                "weight": pct(row["actual_weight"], 2),
                "future_axis": future_moat_axis(row),
                "reason": _selection_reason(row),
                "risk": _main_risk(row),
            }
        )
    return pd.DataFrame(rows)


def make_markdown(tables: dict[str, pd.DataFrame], sensitivity: pd.DataFrame) -> str:
    portfolio = tables["portfolio"]
    performance = tables["performance"]
    ablation = tables["ablation"]
    category = tables["category"]
    market = tables["market"]
    sector = tables["sector"]
    correlation = tables["correlation"]
    summary = tables["summary"]
    contribution = tables["contribution"]
    reasons = portfolio_reason_table(portfolio)

    final_perf = performance.loc[performance["name"] == "Portfolio"].iloc[0]
    topix = performance.loc[performance["name"] == "1306.T"].iloc[0]
    nikkei = performance.loc[performance["name"] == "^N225"].iloc[0]
    invested = portfolio["actual_investment"].sum()
    cash = portfolio["cash_remaining"].iloc[0] if "cash_remaining" in portfolio else 5_000_000 - invested
    future = category.loc[category["category"] == "Future Moat"].iloc[0]
    core = category.loc[category["category"] == "Core Moat"].iloc[0]
    transform = category.loc[category["category"] == "Transformation Moat"].iloc[0]
    santec = contribution.loc[contribution["ticker"] == "6777.T"].iloc[0]

    lines: list[str] = []
    lines.extend(
        [
            "# BEYOND BUFFETT 提出レポート草案",
            "",
            "## 0. 草案の位置づけ",
            "",
            "本草案は、東証上場企業3,649社を対象に実施した自動分析PDFと追加分析CSVをもとに、"
            "日経STOCKリーグ提出用レポートの「Ⅱ スクリーニング」「Ⅲ ポートフォリオ決定」に展開するための本文案である。"
            "ここでの目的は、単に高いバックテスト成績を示すことではない。どのような投資仮説で銘柄を選び、"
            "どのリスクを取った結果としてリターンが生じたのかを、審査員に説明できる論理へ変換することである。",
            "",
            "本分析は公開データに基づく調査であり、個別銘柄の投資助言ではない。また、現在入手可能な財務・株価データで"
            "選定した20社の過去リターンを検証しているため、厳密なリアルタイム運用成績ではなく、ポートフォリオの性質確認として扱う。",
            "",
            "## 1. 投資思想: BEYOND BUFFETT",
            "",
            "伝統的なBuffett型投資は、ブランド、規模、ネットワーク、価格決定力など、すでに確立された競争優位を重視する。"
            "しかし日本株では、低PBR是正、株主還元強化、政策保有株の縮減、資本効率改善、事業ポートフォリオ再編など、"
            "企業価値が後から顕在化する局面が増えている。さらにAI、半導体、光通信、電力、データセンター、省人化投資は、"
            "既存の堀とは異なる新しい産業基盤を形成しつつある。",
            "",
            "そこで本レポートでは、BEYOND BUFFETT企業を「過去に強かった企業」ではなく、"
            "既存の堀を持つか、資本効率改善によって再評価されるか、AI時代の産業構造に接続して将来の堀を深める企業と定義した。",
            "",
            "| 評価軸 | 意味 | レポートでの役割 |",
            "| --- | --- | --- |",
            "| Core Moat | 既存の競争優位、収益性、安定性 | 下落局面での土台 |",
            "| Transformation Moat | 低PBR是正、ROE改善、株主還元、事業改革 | 日本株再評価の受け皿 |",
            "| Future Moat | AI、半導体、光通信、電力、データ、セキュリティ、省人化 | 将来成長の源泉 |",
            "| Valuation Discipline | PER/PBR/配当利回りによる価格規律 | 良い企業を高値づかみしないための制約 |",
            "",
            "この設計により、本ポートフォリオは半導体・AI関連だけに偏るのではなく、電力、金融、鉄道、建設、製造業を含む。"
            "これはテーマが薄いからではなく、AI時代の企業価値を「半導体銘柄」単体ではなく、"
            "電力インフラ、金融の資本再配分、電子部品、交通・建設インフラまで含む産業構造として捉えたためである。",
            "",
            "## 2. Phase 1: 分析結果の読み解き",
            "",
            "### 2.1 スクリーニング全体像",
            "",
            "東証上場企業3,649社から、価格取得、流動性、投資適格性を確認し、最終的に上位80社、20銘柄へ絞り込んだ。",
            "",
            md_table(summary),
            "",
            "市場区分別では、最終20社のうち16社がプライム市場、4社がスタンダード市場、グロース市場は0社である。"
            "これはグロース市場を除外したという意味ではない。全市場を同じ母集団として見たうえで、"
            "流動性、財務安定性、価格データの継続性、500万円ポートフォリオとしての売買可能性を重視した結果である。",
            "",
            md_table(market),
            "",
            "業種別では、Electric Appliancesが候補80社中31社、最終20社中4社と最大である。"
            "一方で、Banks、Insurance、Electric Power and Gasも複数採用された。これは、Future Moatだけでなく、"
            "低PBR是正、金利環境、株主還元、電力需要というTransformation Moatを組み合わせた結果である。",
            "",
            md_table(sector, max_rows=15),
            "",
            "### 2.2 スコア相関の解釈",
            "",
            "スコア相関を見ると、Adjusted BB ScoreはBB Scoreと高く相関し、Moat Scoreとも強く連動している。"
            "一方で、Future Moat Score、Transformation Score、Valuation Scoreは完全には重ならない。"
            "これは、単に一つの指標で上位を拾うのではなく、異なる投資仮説を統合していることを示す。",
            "",
            md_table(correlation),
            "",
            "特にRisk ScoreはAdjusted BB Scoreと負の関係を持つ。これは高ボラティリティや大きなドローダウンを一定程度抑制し、"
            "Future Moat銘柄の成長性を取り込みながらも、無制限にリスクを許容しない設計である。",
            "",
            "### 2.3 アブレーション分析",
            "",
            "アブレーション分析では、最終BEYOND BUFFETTポートフォリオを、単独スコア上位20社、等ウェイト最終20社、"
            "TOPIX ETF、日経平均と比較した。",
            "",
            md_table(ablation),
            "",
            f"最終ポートフォリオの累積リターンは{pct(final_perf['cumulative_return'], 2)}、年率リターンは{pct(final_perf['annualized_return'], 2)}、"
            f"Sharpeは{num(final_perf['sharpe_ratio'], 2)}である。TOPIX ETFの累積リターン{pct(topix['cumulative_return'], 2)}、"
            f"日経平均の{pct(nikkei['cumulative_return'], 2)}を上回った。",
            "",
            "重要なのは、単独スコア上位20社も一定の成果を出しているが、最終ポートフォリオはそれらを上回るバランスを示した点である。"
            "Moatのみ、Transformationのみ、Future Moatのみ、Valuationのみでは、投資仮説が単線的になる。"
            "BEYOND BUFFETTでは、既存の堀、再評価余地、未来の堀、価格規律を組み合わせることで、"
            "リターン、リスク、分散のバランスを改善した。",
            "",
            "一方で、等ウェイト最終20社は累積リターン・Sharpeが最終配分をやや上回った。"
            "これは、スコア加重が必ずしも過去リターン最大化の配分ではないことを意味する。"
            "ただし提出レポートでは、等ウェイトを採用しなかった理由も説明できる。スコア加重は、過去リターン最大化ではなく、"
            "投資仮説の確信度、上限8%、1株単位、分散を同時に満たすためのルールである。",
            "",
            "### 2.4 カテゴリ別リターン",
            "",
            md_table(category),
            "",
            f"カテゴリ別では、Future Moatが累積{pct(future['cumulative_return'], 2)}と最も高い。"
            f"ただし最大ドローダウンは{pct(future['max_drawdown'], 2)}であり、最も大きな下落リスクを伴った。"
            f"Core Moatは累積{pct(core['cumulative_return'], 2)}、Sharpe {num(core['sharpe_ratio'], 2)}で、"
            "安定性とリターンの両面でポートフォリオの土台になった。"
            f"Transformation Moatは累積{pct(transform['cumulative_return'], 2)}と相対的には低いが、"
            "金融、電力、交通、鉄鋼など、日本企業の資本効率改善とマクロ環境変化を取り込む役割を担う。",
            "",
            "したがって本ポートフォリオの構造は、Future Moatが上昇余地を牽引し、Core Moatがリスク調整後の安定性を支え、"
            "Transformation Moatが日本株再評価の可能性を提供する三層構造である。",
            "",
            "### 2.5 santec依存度と集中リスク",
            "",
            f"銘柄別寄与度では、santec Holdingsの寄与が{num(santec['contribution'], 3)}と最大であり、"
            "本ポートフォリオの高い過去リターンが一部Future Moat銘柄に支えられていることは明確である。"
            "この点は弱点として隠すべきではない。",
            "",
            md_table(sensitivity),
            "",
            "感応度分析では、santec単独、上位3銘柄、上位5銘柄を除外した場合のパフォーマンスを確認した。"
            "この分析は、提出レポートでは「高リターンが一部銘柄に依存しているのか」を検証するための反証可能性として使う。"
            "もし除外後もTOPIXを上回るなら、銘柄選定の再現性を主張できる。大きく下がる場合でも、"
            "Future Moat銘柄の集中リスクを正直に認めることで、分析の信頼性はむしろ高まる。",
            "",
            "## 3. Phase 2: 20銘柄の投資仮説",
            "",
            f"500万円ポートフォリオは20銘柄で構成され、投資額は{yen(invested)}、残現金は{yen(cash)}である。"
            "1銘柄上限8%により、スコア上位銘柄に過度に集中しないよう制御した。",
            "",
            md_table(reasons[["code", "company", "category", "weight", "future_axis"]]),
            "",
            "各銘柄の採用理由は、単にスコアが高いからではなく、事業実態、財務指標、ポートフォリオ内の役割が一致しているかで説明する。",
            "",
        ]
    )
    for _, row in reasons.iterrows():
        lines.extend(
            [
                f"### {row['code']} {row['company']}",
                "",
                f"- 採用カテゴリ: {row['category']}。{category_theme(row['category'])}。",
                f"- Future Moat/事業軸: {row['future_axis']}。",
                f"- 採用理由: {row['reason']}",
                f"- 主なリスク: {row['risk']}",
                "",
            ]
        )

    lines.extend(
        [
            "## 4. Phase 3: 提出レポート本文の方針",
            "",
            "提出レポートでは、以下の章立てでまとめる。",
            "",
            "### 4.1 Ⅱ スクリーニング",
            "",
            "まず、BEYOND BUFFETTという投資テーマを定義する。"
            "ここでは、単なる割安株でも成長株でもなく、既存の堀、変革余地、未来の堀を統合する日本株戦略であることを示す。"
            "次に、東証3,649社から価格・流動性・投資適格性を確認し、候補80社へ絞り込む過程を説明する。"
            "この段階では市場区分別・業種別の通過社数を示し、グロース市場を見なかったのではなく、条件を満たす銘柄が少なかったことを明記する。",
            "",
            "スコア設計では、Moat、Transformation、Future Moat、Valuation、Momentum、Riskの役割を説明する。"
            "研究理論との接続として、Quality Minus JunkはMoat、Piotroski F-Score的発想はTransformation、"
            "Fama-FrenchはValuation、Carhart MomentumはMomentum、Markowitzは上限比率と分散、"
            "Sharpe RatioとJensen's Alphaは評価指標として位置づける。",
            "",
            "### 4.2 Ⅲ ポートフォリオ決定",
            "",
            "ポートフォリオ決定では、20銘柄をカテゴリ別に紹介する。"
            "Core Moatは安定性、Transformation Moatは再評価余地、Future Moatは将来成長を担うと整理する。"
            "金融銘柄については、一般事業会社とROICや営業CFで単純比較せず、PBR、ROE、金利環境、株主還元、"
            "リスク分散機能を中心に別枠で評価したと明記する。",
            "",
            "パフォーマンス分析では、累積リターン、年率リターン、Sharpe、最大ドローダウン、CAPM alpha/betaを示す。"
            f"本ポートフォリオは累積{pct(final_perf['cumulative_return'], 2)}、年率{pct(final_perf['annualized_return'], 2)}、"
            f"Sharpe {num(final_perf['sharpe_ratio'], 2)}、最大DD {pct(final_perf['max_drawdown'], 2)}である。"
            "ただし、最大DDはTOPIXより深く、Future Moat銘柄のボラティリティと集中リスクを伴う。"
            "この弱点を認めたうえで、リスクを理解して勝ちに行くポートフォリオとして表現する。",
            "",
            "### 4.3 限界の書き方",
            "",
            "最後に、以下の限界を本文中で明記する。",
            "",
            "- 現在データに基づく選定後に過去株価で検証しているため、厳密なリアルタイム運用成績ではない。",
            "- 現在上場企業中心の母集団であるため、生存者バイアスが残る。",
            "- EDINET XBRLは定量指標には有用だが、事業内容・リスク・ガバナンスなどの定性情報は最終的に有報本文で確認する必要がある。",
            "- 金融銘柄は一般事業会社と同じ財務指標で単純比較しない。",
            "- 高リターンは一部Future Moat銘柄の寄与を含むため、集中リスクを確認する。",
            "",
            "このように弱点を先回りして書くことで、レポートは「勝った結果の宣伝」ではなく、"
            "「リスクを理解したうえで投資仮説を検証した研究」になる。",
            "",
            "## 5. 結論",
            "",
            "BEYOND BUFFETTポートフォリオは、既存の優良企業だけを買う戦略ではない。"
            "日本企業の資本効率改善、低PBR是正、AI時代のインフラ需要、半導体・光通信・電力・金融の再評価を組み合わせ、"
            "これから堀を深める企業を選ぶ戦略である。",
            "",
            "提出レポートでは、高いバックテスト成績を前面に押し出しすぎず、"
            "なぜその成績が生まれたのか、どのリスクを取ったのか、なぜこの20社がテーマに合うのかを中心に論じる。"
            "最終的なメッセージは、無リスクで市場平均を上回るという主張ではなく、"
            "Future MoatとTransformation Moatに対して明示的にリスクを取り、そのリスクをCore Moatと分散ルールで制御する、"
            "日本株向けの発展型Buffett戦略であるという主張である。",
            "",
        ]
    )
    return "\n".join(lines)


def set_doc_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        heading = document.styles[name]
        heading.font.name = "Arial"
        heading.font.size = Pt(size)
        heading.font.color.rgb = color
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(5)


def add_table(document: Document, df: pd.DataFrame, max_rows: int | None = None) -> None:
    if max_rows is not None:
        df = df.head(max_rows)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)


def add_markdownish(document: Document, markdown: str) -> None:
    for line in markdown.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[2:])
            run.font.name = "Arial"
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue
        if line.startswith("|"):
            continue
        document.add_paragraph(line)


def build_docx(markdown: str, tables: dict[str, pd.DataFrame], sensitivity: pd.DataFrame) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_doc_style(document)

    add_markdownish(document, markdown)

    document.add_page_break()
    document.add_heading("付録A 主要表", level=1)
    document.add_heading("市場区分別スクリーニング", level=2)
    add_table(document, tables["market"])
    document.add_heading("カテゴリ別リターン", level=2)
    add_table(document, tables["category"])
    document.add_heading("アブレーション分析", level=2)
    add_table(document, tables["ablation"])
    document.add_heading("集中リスク感応度", level=2)
    add_table(document, sensitivity)

    document.add_heading("付録B 図表", level=1)
    for title, filename in [
        ("スコア相関", "score_correlation_heatmap.png"),
        ("アブレーション累積リターン", "ablation_cumulative_return.png"),
        ("カテゴリ別累積リターン", "category_cumulative_return.png"),
        ("累積リターン", "cumulative_return.png"),
        ("ドローダウン", "drawdown.png"),
        ("銘柄別寄与度", "contribution_by_stock.png"),
    ]:
        path = ROOT / "reports" / "figures" / filename
        if path.exists():
            document.add_heading(title, level=2)
            document.add_picture(str(path), width=Inches(6.2))

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


def build_pdf(markdown: str) -> None:
    pdf = PdfPageWriter.create()
    for line in markdown.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            pdf.heading(text[2:], level=0)
        elif text.startswith("## "):
            pdf.heading(text[3:], level=1)
        elif text.startswith("### "):
            pdf.heading(text[4:], level=2)
        elif text.startswith("|"):
            pdf.text(text, size=14, gap=2)
        elif text.startswith("- "):
            pdf.text("・" + text[2:], size=18, gap=4, indent=20)
        else:
            pdf.text(text, size=19, gap=8)
    pdf.save(OUTPUT_PDF)


def main() -> None:
    tables = load_tables()
    sensitivity = build_sensitivity(tables["portfolio"], tables["contribution"])
    markdown = make_markdown(tables, sensitivity)
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(markdown, encoding="utf-8")
    build_docx(markdown, tables, sensitivity)
    build_pdf(markdown)
    print(OUTPUT_MD)
    print(OUTPUT_DOCX)
    print(OUTPUT_PDF)
    print(SENSITIVITY_CSV)


if __name__ == "__main__":
    main()
