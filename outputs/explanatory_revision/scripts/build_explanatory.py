# -*- coding: utf-8 -*-
"""説明論文（初学者向け）の生成器。MODE で2系統を出す：
  detail  : 詳細解説版（頁無制限・全8項目・付録に設計式/原式/監査）
  contest : 30頁提出版（圧縮8項目・人間記入テンプレ・実使用式のみ本文）
使い方: python build_explanatory.py detail|contest
"""
import sys, json
from pathlib import Path

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
sys.path.insert(0, "/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/final_revision/scripts")
import report_lib as R
import docxlib as X
from eqns_explained import E as EQS
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FR = HERE.parent
FIG = FR / "final_figures"
DATA = json.load(open(FR / "data_real.json", encoding="utf-8"))
EQ = {e["num"]: e for e in EQS}
GO, MIN, BLACK, WHITE = X.GO, X.MIN, X.BLACK, X.WHITE
MODE = sys.argv[1] if len(sys.argv) > 1 else "detail"

doc = Document()
X.setup_styles(doc)


def H(t, lv):
    X.heading(doc, t, lv)


def UH(t):  # unnumbered heading
    p = doc.add_paragraph(style="Heading 1"); X._clear_indent(p)
    X.setfont(p.add_run(t), GO, 12, True, BLACK)


def B(t, indent=True):
    X.body(doc, t, first_indent=indent)


def PB():
    doc.add_page_break()


def FIGURE(fname, title, note):
    fp = FIG / fname
    pt = doc.add_paragraph(); X._clear_indent(pt); pt.paragraph_format.space_before = Pt(6)
    pt.paragraph_format.space_after = Pt(2); pt.paragraph_format.keep_with_next = True
    X.setfont(pt.add_run(title), GO, 9, True, BLACK)
    if fp.exists():
        from PIL import Image
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(pi)
        pi.paragraph_format.keep_with_next = True
        w, h = Image.open(fp).size; wmm = min(140.0, w / 220 * 25.4)
        if wmm * h / w > 150:
            pi.add_run().add_picture(str(fp), height=Mm(150))
        else:
            pi.add_run().add_picture(str(fp), width=Mm(wmm))
    pn = doc.add_paragraph(); X._clear_indent(pn); pn.paragraph_format.space_after = Pt(6)
    X.setfont(pn.add_run(note), MIN, 8, False, BLACK)


def FORM(num):
    R.formula_8item(doc, EQ[num], mode=MODE)


# ============================ 表紙 ============================
for _ in range(4):
    doc.add_paragraph()
for i, (t, sz, bold) in enumerate([("BEYOND BUFFETT", 16, True),
                                   ("― Moat の時間軸拡張による日本株ポートフォリオ ―", 12, True),
                                   ("だれでも検算できる、割安・優良・こわれにくい20社の選び方と配り方", 10.5, False)]):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p)
    p.paragraph_format.space_after = Pt(6 if i == 0 else 3)
    X.setfont(p.add_run(t), GO if i == 0 else MIN, sz, bold, BLACK)
sub = "詳細解説版（教師・専門家レビュー用）" if MODE == "detail" else "提出版（本文30ページ以内）"
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p)
p.paragraph_format.space_before = Pt(18)
X.setfont(p.add_run(sub), MIN, 10, False, BLACK)
for _ in range(4):
    doc.add_paragraph()
for lab in ["［学部・学年］", "［氏名］"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; X._clear_indent(p)
    X.setfont(p.add_run(lab), MIN, 11, False, WHITE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p)
p.paragraph_format.space_before = Pt(18)
X.setfont(p.add_run("（注：PDF 提出版は学部学年・氏名を白文字にする。本文・図表・注に個人名・所属ゼミ名・謝辞を記載しない。）"), MIN, 8, False, BLACK)
PB()

# ============================ 結論先出し ============================
UH("要点（結論先出し）")
B("本研究の目的を一文で言えば、「よい会社を、安く、こわれにくい形で20社選び、500万円を、だれでも検算できる形で配る方法を示すこと」である。")
B("Warren Buffett の投資の本質は、天才的なひらめきではなく、①割安に買い、②優良な会社を選び、③大きく損をしない、という再現可能な規律にある。本研究の題名「BEYOND BUFFETT（バフェットを超える）」は、この規律を否定するのではなく、そこで測る「堀（moat＝競争優位）」を、いまの姿だけでなく「これから変わる姿」「これから生まれる姿」まで時間軸を広げて測る、という意味である。")

H("最終結果 ― 選ばれた20社（役割別）", 2)
B("先に結論を示す。守・破・離の三段階で選んだ最終20社は、次の5つの役割で構成される。銘柄はいったん選んだ後、成績を見て入れ替えることはしていない。")
role_rows = []
order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
role_jp = {"Buffett Core": "完成した堀（守の中核5社・固定）", "Transformation Core": "変わる堀（5社）",
           "Emerging Core": "生まれる堀（5社）", "Dual Moat": "両立型（3社）", "Bridge / Diversifier": "分散役（2社）"}
for role in order:
    names = [f"{d['code']} {d['name'].split()[0][:12]}" for d in DATA.values() if d["role"] == role]
    wsum = sum(d["w"] for d in DATA.values() if d["role"] == role)
    role_rows.append([role_jp[role], "、".join(names), f"{wsum*100:.0f}%"])
R.kv_table(doc, ["役割（日本語）", "銘柄（コード・略称）", "配分合計"], role_rows, [40, 86, 20])
B("完成した堀（Buffett Core）は Phase1 で選んだ5社（3539 JMホールディングス、4350 メディカルシステムネットワーク、6430 大黒電機、7803 ブシロード、9470 学研ホールディングス）で、以降は固定する。合計500万円のうち約495万円を投資し、残りは現金として残る（消化率99.0%）。")

H("全体の流れ ― 3,099社から20社へ（守・破・離）", 2)
B("選定は三段階で進む。守（Phase1）は先行研究の式だけで「完成した堀」を厳しく5社に絞る。破（Phase2）は式を変えずに使い方だけを工夫し、次の段階で見る候補を1,200社に広げる。離（Phase3）は、その1,200社から「変わる堀」「生まれる堀」を測り、三世代あわせて20社を組む。社数の流れは次の通りである。")
FIGURE("phase1_funnel.png", "図表 要-1　守（Phase1）の絞り込み：3,099社→77社→Top5",
       "注：各段階の通過社数。守はここから Buffett Core 5社を固定する。破ではこれとは別に候補を1,200社へ広げ、離で最終20社を選ぶ。筆者作成。")

H("本研究が各社に問う5つの問い", 2)
B("各会社について、私たちは次の5つの問いを順に確かめる。前半3つ（守）はバフェット型の規律、後半2つ（離）が本研究の拡張である。それぞれ、答えを出すために使う式を対応させた。")
R.kv_table(doc, ["#", "問い", "答えを出す式"],
           [["①", "株価は割安か（純資産・利益に対して安いか）", "式 1.1 B/M、式 1.2 E/P"],
            ["②", "良い会社か（資産を効率よく利益に変えているか）", "式 1.3 Gross Profitability、式 1.4 Piotroski"],
            ["③", "財務的に安全か（こわれにくいか）", "式 1.5 Sloan、式 1.7 危機ガードレール"],
            ["④", "これから変わるか（再評価されうるか）", "式 3.2 Transformation Score"],
            ["⑤", "これから新しい堀が生まれるか", "式 3.3 Emerging Score"]],
           [8, 96, 42])
B("以下、用語をやさしく定義したうえで、各段階と各式を「①問い→②ひとことの意味→③式→④変数→⑤数値例→⑥使い方→⑦分かること→⑧限界」の順に説明する。")
PB()

# ============================ 用語集 ============================
UH("はじめに読む用語集")
B("本文で使う言葉を、専門知識のない読者向けに定義する。ここで定義しない略語・記号は本文でも使わない。")
gloss = [
 ("株価", "市場でその会社の1株につく値段（円）。"),
 ("発行済株式数", "その会社が発行した株式の総数。"),
 ("時価総額", "株価×発行済株式数。会社全体の市場価格。"),
 ("売上高", "1年間に商品・サービスを売って得た金額。"),
 ("売上総利益（粗利）", "売上高から売上原価を引いた、もうけの素。"),
 ("営業利益", "本業のもうけ。粗利から販売・管理費を引いたもの。"),
 ("純利益", "税金なども差し引いた最終的なもうけ。"),
 ("自己資本（純資産）", "総資産から負債を引いた、会社の正味の元手。マイナスだと債務超過。"),
 ("総資産", "会社が持つすべての資産（現金・設備・在庫など）。"),
 ("営業キャッシュフロー（営業CF）", "本業で実際に入ってきた現金の増減。"),
 ("PBR（株価純資産倍率）", "時価総額÷純資産。小さいほど純資産に比べ割安。B/M はこの逆数。"),
 ("PER（株価収益率）", "時価総額÷純利益。小さいほど利益に比べ割安。E/P はこの逆数。"),
 ("B/M", "純資産÷時価総額（式 1.1）。大きいほど割安。"),
 ("E/P", "純利益÷時価総額（式 1.2）。大きいほど割安。"),
 ("堀（Moat）", "他社がまねしにくい持続的な競争優位。城の周りの堀にたとえた言葉。"),
 ("ボラティリティ", "株価の変動の大きさ。大きいほどリスクが高い。"),
 ("ベンチマーク", "比較の基準となる市場指数。本研究では TOPIX と日経平均。"),
 ("超過リターン", "ベンチマークを上回った分のリターン。"),
 ("in-sample（標本内）", "分析に使ったのと同じ期間のデータで測った結果。将来の成績を保証しない。"),
 ("Evidence Level（証拠水準）", "スコアの根拠がどれだけ具体的か（式 3.4）。1=言及のみ〜3=数量的裏づけ。"),
]
if MODE == "contest":
    gloss = gloss  # 30頁版も全語掲載（1〜1.5頁に収まる）
R.kv_table(doc, ["用語", "やさしい定義"], [[a, b] for a, b in gloss], [34, 112])
PB()

# ============================ Ⅰ 投資テーマ ============================
H("Ⅰ．投資テーマ", 1)
H("1．なぜ「堀の時間軸」を広げるのか", 2)
B("低く評価されてきた会社が、資本効率の改善や株主還元の強化によって見直される動き（東京証券取引所, 2023；経済産業省, 2014 の資本効率改革）が、いまの日本株では強まっている。これが「変わる堀」である。同時に、AI・半導体・光通信などの構造変化で新しい強みが育ちつつある。これが「生まれる堀」である。完成した堀だけを見ると、この二つを取りこぼす。だから堀を時間軸で三世代に広げて測る。")
B("本研究の独自性は、新しい万能の式を発明することではなく、意味の確かな既存の財務・会計の式を「堀の時間軸」という枠組みに組み替えた点にある。")

# ============================ Ⅱ 守破離 ============================
H("Ⅱ．守・破・離による選定方法", 1)
B("本章が中核である。守で完成した堀を厳しく抽出し、破で候補を広げ、離で三世代の堀を組む。各段階のはじめに「入力・問い・使った式・処理・出力・結論」を表で示す。")

H("1．守 ― 完成した堀を抽出する（Phase1）", 2)
R.phase_io_table(doc, [
    ("入力", "東証上場の非金融普通株 3,099 社と、その財務・株価データ。"),
    ("問い", "いま完成した堀を持ち、割安で、こわれにくい会社はどれか。"),
    ("使った式", "式 1.0 時価総額、1.1 B/M、1.2 E/P、1.3 GP、1.4 Piotroski、1.5 Sloan、1.6 流動性、1.7 危機ガードレール。"),
    ("処理", "各式の関門を順に通し、通過77社から固定順で上位5社を選定。同一業種は原則2社まで。"),
    ("出力", "Buffett Core 5社（3539・4350・6430・7803・9470）。以降固定。"),
    ("結論", "先行研究の式だけで、割安×優良×安全を満たす完成した堀を再現できた。"),
])
for n in ["1.0", "1.1", "1.2", "1.3", "1.4", "1.5", "1.6", "1.7"]:
    FORM(n)

H("2．破 ― 候補を1,200社に広げる（Phase2）", 2)
R.phase_io_table(doc, [
    ("入力", "守と同じ式・データ（3,099社）。"),
    ("問い", "守の厳しさで見落とす「変わる／生まれる堀の候補」を、どうやって広く拾うか。"),
    ("使った式", "式 2.1 正規化、2.2 コンセンサス、2.3 Phase2 信頼度。守の式（1.1〜1.6）は変えない。"),
    ("処理", "式の定義は変えず、閾値・分位・候補数だけを最適化。人がレビューできる規模として Top1200 を採用。"),
    ("出力", "候補1,200社（正式母集団）。Top2000 は取りこぼし確認用の参照。"),
    ("結論", "式を変えずに候補を広げられた。この『広げる判断』が最終選定に最も効いている（Ⅳ章アブレーション）。"),
])
B("なぜ1,200社かは、数式が出した唯一の正解ではない。候補を広げるほど取りこぼしは減るが、人が一社ずつ確認できる規模には限りがある。品質・広さ・レビュー可能性の三つを両立させる運営上の判断として1,200社を選び、その判断過程を開示している。")
for n in ["2.1", "2.2", "2.3"]:
    FORM(n)

H("3．離 ― 三世代の堀を組む（Phase3）", 2)
R.phase_io_table(doc, [
    ("入力", "破の候補1,200社と、各社の確認フラグ・信頼度。"),
    ("問い", "変わる堀・生まれる堀を測り、完成・変化・新生の三世代で20社をどう組むか。"),
    ("使った式", "式 3.2 Transformation（実装形）、3.3 Emerging、3.4 Evidence Level。（設計形 3.1 は付録）"),
    ("処理", "Transformation/Emerging を測り、証拠水準を分離。役割 5/5/5/3/2 に割当。除外862社は理由別に記録。"),
    ("出力", "最終20社（Buffett5＋Transformation5＋Emerging5＋Dual3＋Bridge2）。"),
    ("結論", "低PBR単独でもAIキーワード単独でもない、証拠に裏づけられた三世代ポートフォリオを構成できた。"),
])
FIGURE("moat_timeaxis.png", "図表 Ⅱ-1　堀の時間軸拡張（完成・変化・新生）",
       "注：完成した堀（守5社）・変わる堀（Transformation）・生まれる堀（Emerging）の三世代。筆者作成。")
B("以下に、実際の選定に使った実装形（式 3.2）・生まれる堀（式 3.3）・証拠水準（式 3.4）を示す。理想として設計したが使えなかった完全形（式 3.1）は、混乱を避けるため付録に回す。")
for n in ["3.2", "3.3", "3.4"]:
    FORM(n)

H("4．選定の頑健性（アブレーション）", 2)
B("選定が特定の一要素や一テーマ、後付けの調整に依存していないかを、条件を一つずつ外した16通りの再選定で確かめた。各条件を外したときに最終20社と何社一致するか（重複）を見る。重複が最も小さいのは「候補を Top100 に狭めた」場合（7社）で、破で候補を1,200社に広げた判断が最も効いていることを示す。")
FIGURE("ablation_overlap.png", "図表 Ⅱ-2　条件を一つ外したときの最終20社との重複",
       "注：重複は20社中の一致数。A8（Top100に限定）が最小の7。in-sample の構造検査であり成績の主張ではない。出所：ablation_results.csv。")

# ============================ Ⅲ 配分 ============================
H("Ⅲ．最終20社と500万円の配分", 1)
R.phase_io_table(doc, [
    ("入力", "選定済みの20社（役割つき）と、各社の株価・変動・流動性・証拠・信頼度。"),
    ("問い", "500万円を、成績予想に頼らずどう配るか。"),
    ("使った式", "式 4.1 役割予算×リスク調整配分、4.2 単元株調整。"),
    ("処理", "役割ごとに予算を決め、役割内でリスク・流動性・証拠・信頼度で傾け、8%上限と単元株に丸める。"),
    ("出力", "20社の目標比率と購入株数。投資額 4,949,198 円・残現金 50,801 円（消化率99.0%）。"),
    ("結論", "リターン予想を使わず、説明できる形で配分を決められた。"),
])
B("配分の考え方は4段階である。第一に20社を役割ごとに分ける。第二に役割ごとの予算を決める（完成・変化・新生を各25%、両立型15%、分散役10%）。第三に同じ役割の中で、変動が小さく・売買しやすく・証拠が強く・信頼できる会社を厚くする。第四に、1銘柄8%の上限と、実際に買える株数（単元株）に丸める。")
for n in ["4.1", "4.2"]:
    FORM(n)
B("最終配分を図表 Ⅲ-1 に示す。役割の合計は25/25/25/15/10、最大保有はゼンリン7.46%で、いずれも制約内である。")
# portfolio table (real)
prows = []
for role in order:
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        prows.append([d["code"], d["name"][:16], role_jp[role].split("（")[0], d["theme"],
                      f"{d['w']*100:.2f}%", str(d["qtyL1"]), f"{d['amtL1']:,}"])
R.kv_table(doc, ["コード", "企業名", "役割", "テーマ", "比率", "株数(L=1)", "金額(L=1)"],
           prows, [13, 34, 22, 22, 13, 15, 20])
B("図表 Ⅲ-1　最終ポートフォリオ（配分案・単元未満株 L=1 基準）。出所：allocation_final.csv。")

if MODE == "contest":
    H("最終20社の紹介（記入テンプレート）", 2)
    B("各社の事業概要・堀の根拠・採用理由・主なリスク・一次資料の出典は、提出者が各社IR・有価証券報告書に基づいて記入する（捏造を避け空欄とした）。左の定量値は正典データから自動記入済み。")
    mrows = []
    for role in order:
        for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
            mrows.append([f"{d['code']} {d['name'][:14]}／{d['sector']}", f"役割={role_jp[role].split('（')[0]}／T{d['tsc']}/E{d['esc']}/L{d['evid']}", "（記入）"])
    R.kv_table(doc, ["コード・企業名／業種", "役割・定量（自動）", "事業概要・堀の根拠・採用理由・リスク・出典（記入）"],
               mrows, [40, 40, 66])

# ============================ Ⅳ 検証 ============================
H("Ⅳ．検証・アブレーション・限界", 1)
H("1．検証の位置づけ（in-sample・成績の主張ではない）", 2)
B("本ポートフォリオは 2026 年 6 月時点のデータで作ったため、過去をさかのぼった成績計算はすべて標本内（in-sample）である。したがって以下の数値は「過去データ上のリスクの姿」の確認であって、将来の成績や優位性の証明ではない。ベンチマークの 1306.T には未調整の株式分割があったため補正した（注1）。")
B("3年（in-sample）の主な数値は、年率リターン +30.8%、リスク（年率変動）21.9%、最大下落 −24.9%、市場連動度ベータ 0.925、Sharpe 1.41、Jensen α +7.3%/年である。いずれも標本内の値であり、優位性の主張ではない。とくに直近1年のインフォメーション・レシオは −0.405（負）で、市場に負けている（式 5.4）。")
for n in ["5.1", "5.2", "5.3", "5.4", "5.5"]:
    FORM(n)
FIGURE("drawdown_chart.png", "図表 Ⅳ-1　ドローダウン（3年・in-sample）",
       "注：リスク特性の確認であり成績の主張ではない。1306.T は分割補正済み。出所：phase5_validation_summary.json。")
H("2．主要なリスクと限界", 2)
B("本研究の主なリスクと限界は次の通りである。①将来の堀（Emerging Core）への寄与集中、②AIテーマの過熱（テーマ集中 HHI 0.402）、③変わる堀の割安のワナ、④開示不足による過小・過大評価（Emergingの証拠 Level 2 以上は14社の手作業精査に依存）、⑤流動性・小型株リスク、⑥業種集中、⑦単元株の丸めによる歪み（実単元では9社買えず、単元未満株が前提）、⑧標本内評価の限界（直近1年の超過は負）、⑨独自合成式の恣意性（重みは設計値で、±20%動かしても選定は不変）。これらはすべて本文に明記し、過大な主張を避けた。")

# ============================ Ⅴ 一次情報と学び ============================
H("Ⅴ．一次情報と学び", 1)
if MODE == "contest":
    H("1．企業・専門家インタビュー（記入欄）", 2)
    B("選定が実態を反映しているかを確かめるため、投資対象企業へのインタビュー・アンケートを行う。未実施の回答は捏造しない。以下は記入欄である。")
    for lab in ["対象者（企業名・部署・役職）", "実施日・実施方法", "質問①〜③", "回答①〜③", "分析への反映"]:
        pp = doc.add_paragraph(style="Normal"); X._clear_indent(pp); pp.paragraph_format.space_after = Pt(1)
        X.setfont(pp.add_run(f"・{lab}：［記入］"), MIN, 9)
    H("2．研究を通じた学びと振り返り（記入欄）", 2)
    for lab in ["仮説が変化した点", "分析で失敗した点・工夫した点", "チームで学んだ点", "今後の課題"]:
        pp = doc.add_paragraph(style="Normal"); X._clear_indent(pp); pp.paragraph_format.space_after = Pt(1)
        X.setfont(pp.add_run(f"・{lab}：［記入］"), MIN, 9)
else:
    B("インタビュー・アンケート・振り返りは提出者が実体験に基づき記入する（捏造しない）。詳細解説版では記入欄を省略し、代わりに次章の通し計算例で全式のつながりを示す。")

# ============================ 通し計算例 ============================
H("Ⅵ．通し計算例 ― 学研ホールディングス（9470）で全式をつなぐ" if MODE == "detail" else "参考：通し計算例（9470 学研ホールディングス）", 1)
d = DATA["9470"]
B("守で選ばれた完成した堀の一社、9470 学研ホールディングスの実際の値を使って、時価総額から最終投資額までを一つにつなぐ。全候補にすべての指標が計算されるため、1社で全ステップをたどれる。")
steps = [
 ("① 時価総額", f"株価 {d['price']:.0f}円 × 約4,141万株 ≒ {d['mktcap_oku']}億円。会社全体の値段。"),
 ("② B/M（割安さ・純資産）", f"{d['bm']}。純資産が時価総額の約1.5倍。市場上位30%の割安さで守の関門①を通過。"),
 ("③ E/P（割安さ・利益）", f"{d['ep']}（PER約8倍）。黒字で割安。関門②を通過。"),
 ("④ Gross Profitability（質）", f"{d['gp']}。資産効率は市場中央値以上。関門③を通過。"),
 ("⑤ Piotroski（財務健全性）", f"{d['piotroski']}。計算できた全シグナルに合格。関門④を通過。"),
 ("⑥ Sloan（利益の質）", f"{d['sloan']}。マイナス＝現金の裏づけが厚く良質。関門⑤を通過。"),
 ("⑦ 危機ガードレール", f"債務超過なし・3期連続赤字なし → D=0。関門を通過。"),
 ("⑧ Phase2 信頼度", f"{d['conf']}。データが安定し要確認が少ない。"),
 ("⑨ 変わる堀／生まれる堀", f"Transformation={d['tsc']}（高め）だが Emerging={d['esc']}（低い）。生まれる堀では選ばれない。"),
 ("⑩ 役割割当", f"守 Top5 として固定 → 完成した堀（Buffett Core）。Evidence Level は L{d['evid']}。"),
 ("⑪ 配分優先度", f"変動小・流動性高（ℓ=1.00）・証拠 L3（e=1.05）・信頼度{d['conf']} → 役割内で厚め。"),
 ("⑫ 最終投資額", f"目標比率 {d['w']*100:.2f}% → 941円で {d['qtyL1']}株 → {d['amtL1']:,}円（実測）。"),
]
R.kv_table(doc, ["ステップ", "9470 学研HD の実際の値と判定"], [[a, b] for a, b in steps], [40, 106])
B("このように、守の5つの関門を通り、変わる堀の評価は高いが生まれる堀は低いため、9470 は「完成した堀」の役割に落ち着き、約34万円が配分される。役割が分かれる判定点（ここでは Emerging が低いこと）を示すことで、恣意的でないことを確かめられる。")

# ============================ 注・参考文献 ============================
UH("注")
for t in ["1) ベンチマーク 1306.T の未調整分割（2026-03-30）を補正した（同日以降を10倍）。構成銘柄に分割様の不連続はない。",
          "2) 株数は基準ケースで単元未満株（L=1）を前提とする。実単元（L=100）では株価の高い9社が買えず消化率46.7%へ低下する。",
          "3) すべての数値は 2026 年 6 月時点のデータに基づく。履歴計算は標本内であり将来を保証しない。"]:
    B(t, indent=False)
UH("参考文献")
try:
    refs = (Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/beyond_buffett_fable_loop_final/phase7_final_report/final_references.md").read_text(encoding="utf-8").splitlines())
    import re
    emit = False
    for rl in refs:
        r = rl.rstrip()
        if r.startswith("## 英語文献"):
            pp = doc.add_paragraph(); X._clear_indent(pp); X.setfont(pp.add_run("英語文献"), GO, 9, True, BLACK); emit = True; continue
        if r.startswith("## 日本語文献"):
            pp = doc.add_paragraph(); X._clear_indent(pp); X.setfont(pp.add_run("日本語文献"), GO, 9, True, BLACK); continue
        if r.startswith("## 本文引用"):
            break
        if emit and r.strip() and not r.startswith("#") and r.strip() != "---":
            from docx.oxml import OxmlElement
            pp = doc.add_paragraph(); ppr = pp._p.get_or_add_pPr()
            ind = OxmlElement("w:ind"); ind.set(X.qn("w:left"), "360"); ind.set(X.qn("w:hanging"), "360"); ppr.append(ind)
            pp.paragraph_format.space_after = Pt(1)
            X.setfont(pp.add_run(re.sub(r"\*(.+?)\*", r"\1", r.strip())), MIN, 8, False, BLACK)
except Exception as ex:
    B(f"（参考文献の読み込みに失敗：{ex}）")

# ============================ 付録（詳細版のみ） ============================
if MODE == "detail":
    PB()
    UH("付録A　設計完全形（実際の選定には未使用）")
    B("Transformation Moat Score の設計完全形（式 3.1）は、株主還元・改革開示の正式データが入力に存在しないため実際の選定には使えず、概念上の参照式にとどまる。実装済みのように扱ってはならない。")
    FORM("3.1")
    UH("付録B　倒産予測の原式（Ohlson・Altman）と実装できなかった理由")
    B("倒産予測の代表式に Ohlson O-Score（1980）と Altman Z-Score（1968）がある。O-Score はロジスティック回帰で倒産確率を、Z-Score は複数の財務比率の加重和で安全度を出す。いずれも、時価・負債の比率や運転資本など、本研究の入力ではすべての企業でそろわない変数を必要とした。不完全な変数で原式を名乗ると誤った倒産確率を示す恐れがあるため、原式は用いず、そろうデータで作れる最低限の除外規則（式 1.7）にとどめた。式 1.7 は倒産確率を出せない点で原式に劣るが、債務超過と連続赤字という明確な危険だけは確実に外せる。")

out = FR / (f"beyond_buffett_explanatory_report.docx" if MODE == "detail" else "contest_report_30pages.docx")
doc.save(str(out))
print("saved", out.name, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables), "| mode:", MODE)
