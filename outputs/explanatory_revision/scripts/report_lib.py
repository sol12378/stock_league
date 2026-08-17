# -*- coding: utf-8 -*-
"""説明論文用レンダリング補助。既存 docxlib（v11式ボックス・OMML・スタイル）を再利用し、
8項目式ブロック・変数表・Phase入出力表・用語表を追加する。"""
import sys
from pathlib import Path

LOOP = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/final_revision/scripts")
sys.path.insert(0, str(LOOP))
import docxlib as X  # noqa

from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

GO, MIN, GRAY, WHITE, BLACK = X.GO, X.MIN, X.GRAY, X.WHITE, X.BLACK


def label_para(doc, label, text, size=9, after=3):
    """『ラベル：本文』の段落。ラベルは太字。"""
    p = doc.add_paragraph(style="Normal"); X._clear_indent(p)
    p.paragraph_format.space_after = Pt(after)
    X.setfont(p.add_run(label), MIN, size, True, BLACK)
    if text:
        X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def var_table(doc, rows, width_mm=146.0):
    """④詳細変数表：記号/日本語名/使うデータ/計算方法/範囲/高い・1の意味。"""
    hdr = ["記号", "日本語名", "使うデータ", "計算方法", "範囲", "高い/1 の意味"]
    t = doc.add_table(rows=1, cols=6); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [16, 26, 30, 28, 20, 26]
    for j, h in enumerate(hdr):
        c = t.rows[0].cells[j]; X.shade_cell(c, GRAY)
        pp = c.paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(h), GO, 7.5, True, WHITE)
    for r in rows:
        cells = t.add_row().cells
        for j in range(6):
            v = r[j] if j < len(r) else ""
            pp = cells[j].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
            X.setfont(pp.add_run(str(v)), MIN, 7.5, False, BLACK)
    for r in t.rows:
        X.cant_split(r)
        for j, w in enumerate(widths):
            r.cells[j].width = Mm(w)
    sp = doc.add_paragraph(); X._clear_indent(sp); sp.paragraph_format.space_after = Pt(3)


def short_symbol_line(symbols):
    """③式ボックス第3行用：記号→和名の短い対応表（説明はしない）。"""
    return "　".join(f"{s[0]}＝{s[1]}" for s in symbols[:6])


def formula_8item(doc, e, mode="detail"):
    """8項目ブロック。mode='detail' は全項目、'contest' は圧縮（④⑤を省略/短縮）。"""
    # 見出し（第4見出し・番号なし本文継続扱いだがここは小見出しとして太字）
    h = doc.add_paragraph(style="Heading 4"); X._clear_indent(h)
    X.setfont(h.add_run(f"（{e['num']}）{e['title'].split('（')[0]}　［{e['cls']}］"), GO, 9, True, BLACK)
    # ① 問い
    label_para(doc, "問い　", e["q"])
    # ② 一文の意味
    label_para(doc, "ひとことで言うと　", e["meaning"])
    # ③ v11 式ボックス（第3行＝短い記号対応表）
    X.formula_box(doc, e["num"], e["title"], e["latex"], short_symbol_line(e["symbols"]))
    if mode == "detail":
        # ④ 詳細変数表
        label_para(doc, "変数の意味　", "（下表）", after=1)
        var_table(doc, e["symbols"])
        # ⑤ 数値例
        label_para(doc, "数値例　", e["example"])
    # ⑥ 使用方法
    label_para(doc, "本研究での使い方　", e["usage"])
    # ⑦ 分かること
    label_para(doc, "この式で分かること　", e["finding"])
    # ⑧ 限界
    label_para(doc, "この式では分からないこと・限界　", e["limit"], after=6)


def phase_io_table(doc, rows, width_mm=146.0):
    """各Phase冒頭の 入力/問い/使用式/処理/出力/結論 表。rows=[(項目,内容)...]"""
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = t.add_row().cells
        X.shade_cell(cells[0], "D9D9D9")
        pp = cells[0].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(k), GO, 8.5, True, BLACK)
        pp2 = cells[1].paragraphs[0]; X._clear_indent(pp2); pp2.paragraph_format.space_after = Pt(0)
        X.setfont(pp2.add_run(v), MIN, 8.5, False, BLACK)
    for r in t.rows:
        X.cant_split(r); r.cells[0].width = Mm(24); r.cells[1].width = Mm(122)
    sp = doc.add_paragraph(); X._clear_indent(sp); sp.paragraph_format.space_after = Pt(4)


def kv_table(doc, hdr, rows, widths):
    """汎用表（見出し灰）。"""
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hh in enumerate(hdr):
        c = t.rows[0].cells[j]; X.shade_cell(c, GRAY)
        pp = c.paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(hh), GO, 8, True, WHITE)
    for r in rows:
        cells = t.add_row().cells
        for j in range(len(hdr)):
            pp = cells[j].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
            X.setfont(pp.add_run(str(r[j]) if j < len(r) else ""), MIN, 8, False, BLACK)
    for r in t.rows:
        X.cant_split(r)
        for j, w in enumerate(widths):
            r.cells[j].width = Mm(w)
    sp = doc.add_paragraph(); X._clear_indent(sp); sp.paragraph_format.space_after = Pt(4)
