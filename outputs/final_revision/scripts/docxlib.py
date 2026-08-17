"""Shared builders for the student-prize DOCX: academic styles, OMML equations
(via pandoc), and the v11-style gray 1x3 formula box."""
import copy
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

MIN = "ＭＳ 明朝"
GO = "ＭＳ ゴシック"
MATH = "Cambria Math"
GRAY = "808080"
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------- OMML
_omml_cache = {}


def latex_to_omml(latex):
    """Convert a LaTeX string to an OMML <m:oMathPara> element (centered display)."""
    if latex in _omml_cache:
        return copy.deepcopy(_omml_cache[latex])
    with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(f"$${latex}$$\n")
        mp = f.name
    dp = mp + ".docx"
    subprocess.run(["pandoc", mp, "-o", dp], check=True, capture_output=True)
    xml = zipfile.ZipFile(dp).read("word/document.xml")
    root = etree.fromstring(xml)
    node = root.find(f".//{{{M_NS}}}oMathPara")
    if node is None:                      # fall back to bare oMath, wrap it
        omath = root.find(f".//{{{M_NS}}}oMath")
        node = etree.SubElement(etree.Element(f"{{{M_NS}}}oMathPara"), f"{{{M_NS}}}oMath")
        node = node.getparent()
        node.append(copy.deepcopy(omath))
    _omml_cache[latex] = node
    return copy.deepcopy(node)


# ---------------------------------------------------------------- fonts / shading
def setfont(run, name=MIN, size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)
    latin = "Times New Roman" if name == MIN else name
    rf.set(qn("w:ascii"), latin); rf.set(qn("w:hAnsi"), latin)
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:cantSplit"); trpr.append(e)


def _set_border(tcpr, edge, sz, color, val="single"):
    tb = tcpr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcpr.append(tb)
    e = tb.find(qn(f"w:{edge}"))
    if e is None:
        e = OxmlElement(f"w:{edge}"); tb.append(e)
    e.set(qn("w:val"), val); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)


def cell_borders(cell, top=False, bottom=False, sz=6, color="000000"):
    """0.75pt = sz 6 (eighths of a point). No vertical lines (spec)."""
    tcpr = cell._tc.get_or_add_tcPr()
    if top:
        _set_border(tcpr, "top", sz, color)
    if bottom:
        _set_border(tcpr, "bottom", sz, color)
    _set_border(tcpr, "left", sz, color)
    _set_border(tcpr, "right", sz, color)


# ---------------------------------------------------------------- styles
def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(32)
    sec.top_margin = sec.bottom_margin = Mm(40.5)
    # Normal = MS Mincho 9pt, justified, first-line indent 1 char
    normal = doc.styles["Normal"]
    normal.font.name = MIN
    normal.font.size = Pt(9)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), MIN); rf.set(qn("w:ascii"), "Times New Roman"); rf.set(qn("w:hAnsi"), "Times New Roman")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    # first-line indent 1 Japanese char
    ppr = normal.element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind"); ppr.append(ind)
    ind.set(qn("w:firstLineChars"), "100")   # 1 char
    ind.set(qn("w:firstLine"), "180")
    # Headings 1-4 = MS Gothic, monochrome black, left, sizes 12/11/10/9
    for lvl, size in [(1, 12), (2, 11), (3, 10), (4, 9)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = GO; st.font.size = Pt(size); st.font.bold = True
        st.font.color.rgb = BLACK
        r = st.element.get_or_add_rPr()
        rf2 = r.find(qn("w:rFonts"))
        if rf2 is None:
            rf2 = OxmlElement("w:rFonts"); r.append(rf2)
        rf2.set(qn("w:eastAsia"), GO); rf2.set(qn("w:ascii"), GO); rf2.set(qn("w:hAnsi"), GO)
        p = st.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.space_before = Pt(10 if lvl <= 2 else 6)
        p.space_after = Pt(3)
        p.keep_with_next = True
        # no first-line indent on headings
        pp = st.element.get_or_add_pPr()
        i2 = pp.find(qn("w:ind"))
        if i2 is None:
            i2 = OxmlElement("w:ind"); pp.append(i2)
        i2.set(qn("w:firstLineChars"), "0"); i2.set(qn("w:firstLine"), "0")


# ---------------------------------------------------------------- paragraphs
def body(doc, text, first_indent=True):
    p = doc.add_paragraph(style="Normal")
    if not first_indent:
        pp = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind"); ind.set(qn("w:firstLine"), "0"); ind.set(qn("w:firstLineChars"), "0")
        pp.append(ind)
    if text:
        setfont(p.add_run(text), MIN, 9)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    setfont(p.add_run(text), GO, {1: 12, 2: 11, 3: 10, 4: 9}[level], True, BLACK)
    return p


# ---------------------------------------------------------------- v11 formula box
def formula_box(doc, num, title, latex, defn, width_mm=146.0):
    t = doc.add_table(rows=3, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for r in t.rows:
        cant_split(r)
        r.cells[0].width = Mm(width_mm)
    # row1: gray title band, white MSGothic bold
    c0 = t.rows[0].cells[0]
    shade_cell(c0, GRAY); cell_borders(c0, top=True, bottom=True)
    p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(0); p0.paragraph_format.keep_with_next = True
    _clear_indent(p0)
    setfont(p0.add_run(f"（{num}）{title}"), GO, 9, True, WHITE)
    # row2: white, centered OMML equation
    c1 = t.rows[1].cells[0]
    cell_borders(c1, bottom=True)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(5); p1.paragraph_format.space_after = Pt(5)
    p1.paragraph_format.keep_with_next = True
    _clear_indent(p1)
    omath = latex_to_omml(latex)
    # center the oMathPara
    jc = omath.find(f"{{{M_NS}}}oMathParaPr/{{{M_NS}}}jc")
    if jc is None:
        pr = omath.find(f"{{{M_NS}}}oMathParaPr")
        if pr is None:
            pr = etree.Element(f"{{{M_NS}}}oMathParaPr"); omath.insert(0, pr)
        jc = etree.SubElement(pr, f"{{{M_NS}}}jc")
    jc.set(f"{{{M_NS}}}val", "center")
    p1._p.append(omath)
    # row3: white, MS Mincho 8pt definition (left)
    c2 = t.rows[2].cells[0]
    cell_borders(c2, bottom=True)
    p2 = c2.paragraphs[0]; p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    _clear_indent(p2)
    setfont(p2.add_run(defn), MIN, 8)
    # spacing paragraph after box (<=1 line)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2); _clear_indent(sp)
    return t


def _clear_indent(p):
    pp = p._p.get_or_add_pPr()
    ind = pp.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind"); pp.append(ind)
    ind.set(qn("w:firstLine"), "0"); ind.set(qn("w:firstLineChars"), "0")
