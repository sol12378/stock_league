#!/usr/bin/env python
"""Assemble the student-prize DOCX from final_report_source.md.
Academic format: MS Mincho 9pt body, MS Gothic monochrome headings, 32/40.5mm
margins, v11 gray OMML formula boxes, monochrome figures/tables, fill-in
templates. Data tables/templates built from canonical CSVs."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
import docxlib as X
from eqns import EQUATIONS
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image
import pandas as pd

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
LF = ROOT / "outputs/beyond_buffett_fable_loop_final"
FR = ROOT / "outputs/final_revision"
FIG = FR / "final_figures"
EQ_BY = {e[0]: e for e in EQUATIONS}
GO, MIN, GRAY, WHITE, BLACK = X.GO, X.MIN, X.GRAY, X.WHITE, X.BLACK

doc = Document()
X.setup_styles(doc)


def unnum_heading(text):
    p = doc.add_paragraph(style="Heading 1")
    X._clear_indent(p)
    X.setfont(p.add_run(text), GO, 12, True, BLACK)
    return p


def figure(fname, title, note):
    fp = FIG / fname
    pt = doc.add_paragraph(); X._clear_indent(pt)
    pt.paragraph_format.space_before = Pt(6); pt.paragraph_format.space_after = Pt(2)
    pt.paragraph_format.keep_with_next = True
    X.setfont(pt.add_run(title), GO, 9, True, BLACK)   # title above, left
    if fp.exists():
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        X._clear_indent(pi); pi.paragraph_format.keep_with_next = True
        w_px, h_px = Image.open(fp).size
        w_mm = min(140.0, w_px / 220 * 25.4)
        if w_mm * h_px / w_px > 175:
            pi.add_run().add_picture(str(fp), height=Mm(175))
        else:
            pi.add_run().add_picture(str(fp), width=Mm(w_mm))
    pn = doc.add_paragraph(); X._clear_indent(pn); pn.paragraph_format.space_after = Pt(6)
    X.setfont(pn.add_run(note), MIN, 8, False, BLACK)


def _hdr_cell(cell, text):
    X.shade_cell(cell, GRAY)
    p = cell.paragraphs[0]; X._clear_indent(p); p.paragraph_format.space_after = Pt(0)
    X.setfont(p.add_run(text), GO, 8.5, True, WHITE)


def _cell(cell, text, size=8, bold=False, align=None, right=False):
    p = cell.paragraphs[0]; X._clear_indent(p); p.paragraph_format.space_after = Pt(0)
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        X.setfont(p.add_run(text), MIN, size, bold, BLACK)


def _load_final():
    a = pd.read_csv(LF / "phase4_portfolio_allocation/allocation_final.csv")
    order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
    a["__o"] = a.final_role.map({r: i for i, r in enumerate(order)})
    a = a.sort_values(["__o", "target_weight_final"], ascending=[True, False])
    try:
        ts = pd.read_csv(LF / "phase3_moat_construction/transformation_scores.csv")
        es = pd.read_csv(LF / "phase3_moat_construction/emerging_scores.csv")
        import re
        def nn(x): return re.search(r"\d+", str(x).replace(".0", "")).group(0).zfill(4)
        a["cn"] = a.code_n.map(nn)
        for df, col, new in [(ts, "transformation_score", "tsc"), (es, "emerging_score", "esc")]:
            if col in df.columns:
                key = "code" if "code" in df.columns else df.columns[0]
                df["cn"] = df[key].map(nn)
                a = a.merge(df[["cn", col]].drop_duplicates("cn"), on="cn", how="left").rename(columns={col: new})
    except Exception:
        a["tsc"] = None; a["esc"] = None
    return a


def portfolio_table():
    a = _load_final()
    hdr = ["コード", "企業名", "役割", "テーマ", "目標比率", "株数(L=1)", "購入金額(L=1)"]
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdr):
        _hdr_cell(t.rows[0].cells[j], h)
    for _, r in a.iterrows():
        c = t.add_row().cells
        vals = [int(r.code_n), str(r.company_name)[:18], r.final_role, r.theme,
                f"{r.target_weight_final*100:.2f}%", str(int(r.qty_L1)), f"{int(r.amount_L1):,}"]
        for j, v in enumerate(vals):
            _cell(c[j], str(v), right=(j >= 4))
    for r in t.rows:
        r.cells[0].width = Mm(13); r.cells[1].width = Mm(38); r.cells[2].width = Mm(30)
        r.cells[3].width = Mm(24); r.cells[4].width = Mm(15); r.cells[5].width = Mm(13); r.cells[6].width = Mm(20)
    cap = doc.add_paragraph(); X._clear_indent(cap); cap.paragraph_format.space_after = Pt(6)
    X.setfont(cap.add_run("出所：allocation_final.csv。総予算 500 万円、投資額 4,949,198 円、残現金 50,801 円（消化率 99.0%）。"), MIN, 8)


def meigara_template():
    a = _load_final()
    hdr = ["コード・企業名／業種／役割・テーマ", "定量（T/E スコア・種別・Ev.・比率）", "事業概要・Moat の根拠・採用理由・主なリスク・一次資料出典（記入）"]
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdr):
        _hdr_cell(t.rows[0].cells[j], h)
    for _, r in a.iterrows():
        c = t.add_row().cells
        left = f"{int(r.code_n)} {str(r.company_name)[:20]}\n{r.sector}\n{r.final_role}／{r.theme}"
        _cell(c[0], left, size=7.5)
        tsc = f"{r.tsc:.0f}" if pd.notna(r.get("tsc")) else "―"
        esc = f"{r.esc:.0f}" if pd.notna(r.get("esc")) else "―"
        quant = f"T={tsc} / E={esc}\nEv.=L{int(r.final_evidence_level)}\n比率={r.target_weight_final*100:.1f}%"
        _cell(c[1], quant, size=7.5)
        _cell(c[2], "", size=8)  # blank for submitter
    for r in t.rows:
        r.cells[0].width = Mm(40); r.cells[1].width = Mm(30); r.cells[2].width = Mm(76)
    cap = doc.add_paragraph(); X._clear_indent(cap); cap.paragraph_format.space_after = Pt(6)
    X.setfont(cap.add_run("注：左二列は正典データから自動記入。右列（定性記述・一次資料出典）は提出者が各社 IR・有価証券報告書に基づき記入する。"), MIN, 8)


def interview_template():
    for lab in ["対象者（企業名・部署・役職）", "実施日・実施方法", "質問①〜③", "回答①〜③", "分析への反映"]:
        p = doc.add_paragraph(); X._clear_indent(p); p.paragraph_format.space_after = Pt(1)
        X.setfont(p.add_run(f"・{lab}：［記入］"), MIN, 9)
    p = doc.add_paragraph(); X._clear_indent(p); p.paragraph_format.space_after = Pt(6)
    X.setfont(p.add_run("（対象 2〜4 社を目安に記入。未実施の場合は本節を「今後の課題」に移す。捏造しない。）"), MIN, 8)


def reflect_template():
    for lab in ["仮説が変化した点", "分析で失敗した点・工夫した点", "チームで学んだ点", "今後の課題"]:
        p = doc.add_paragraph(); X._clear_indent(p); p.paragraph_format.space_after = Pt(1)
        X.setfont(p.add_run(f"・{lab}：［記入］"), MIN, 9)
    p = doc.add_paragraph(); X._clear_indent(p); p.paragraph_format.space_after = Pt(6)
    X.setfont(p.add_run("（研究過程の記録から下書き可。最終的には提出者の実体験として記入・確認する。）"), MIN, 8)


def references():
    refs = (LF / "phase7_final_report/final_references.md").read_text(encoding="utf-8").splitlines()
    emit = False
    import re
    for rl in refs:
        r = rl.rstrip()
        if r.startswith("## 英語文献"):
            p = doc.add_paragraph(); X._clear_indent(p); X.setfont(p.add_run("英語文献"), GO, 9, True, BLACK); emit = True; continue
        if r.startswith("## 日本語文献"):
            p = doc.add_paragraph(); X._clear_indent(p); X.setfont(p.add_run("日本語文献"), GO, 9, True, BLACK); continue
        if r.startswith("## 本文引用"):
            break
        if emit and r.strip() and not r.startswith("#") and r.strip() != "---":
            txt = re.sub(r"\*(.+?)\*", r"\1", r.strip())
            p = doc.add_paragraph(); pp = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360"); pp.append(ind)
            p.paragraph_format.space_after = Pt(1)
            X.setfont(p.add_run(txt), MIN, 8)


def cover(lines):
    for _ in range(4):
        doc.add_paragraph()
    sizes = [16, 12, 10.5]
    for i, ln in enumerate(lines):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p)
        p.paragraph_format.space_after = Pt(6 if i == 0 else 3)
        X.setfont(p.add_run(ln), GO if i == 0 else MIN, sizes[min(i, 2)], i <= 1, BLACK)
    for _ in range(6):
        doc.add_paragraph()
    for lab in ["［学部・学年］", "［氏名］"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; X._clear_indent(p)
        p.paragraph_format.space_after = Pt(3)
        X.setfont(p.add_run(lab), MIN, 11, False, WHITE)   # white for PDF blind review
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p)
    p.paragraph_format.space_before = Pt(24)
    X.setfont(p.add_run("（注：PDF 提出版は学部学年・氏名を白文字にする。本文・図表・注に個人名・所属ゼミ名・謝辞を記載しない。）"), MIN, 8, False, BLACK)


# ---------------------------------------------------------------- parse
src = (FR / "final_report_source.md").read_text(encoding="utf-8").splitlines()
i = 0
while i < len(src):
    ln = src[i].rstrip()
    if ln == "@COVER":
        block = []
        i += 1
        while i < len(src) and src[i].strip() and not src[i].startswith("@"):
            block.append(src[i].strip()); i += 1
        cover(block); continue
    if ln == "@PB":
        doc.add_page_break(); i += 1; continue
    if ln.startswith("@ABST "):
        unnum_heading(ln[6:].strip()); i += 1; continue
    if ln.startswith("@1 "):
        X.heading(doc, ln[3:].strip(), 1); i += 1; continue
    if ln.startswith("@2 "):
        X.heading(doc, ln[3:].strip(), 2); i += 1; continue
    if ln.startswith("@3 "):
        X.heading(doc, ln[3:].strip(), 3); i += 1; continue
    if ln.startswith("@4 "):
        X.heading(doc, ln[3:].strip(), 4); i += 1; continue
    if ln.startswith("@EQ "):
        num = ln[4:].strip()
        n, title, latex, defn = EQ_BY[num]
        X.formula_box(doc, n, title, latex, defn); i += 1; continue
    if ln.startswith("@FIG "):
        parts = ln[5:].split("|")
        figure(parts[0], parts[1] if len(parts) > 1 else "", parts[2] if len(parts) > 2 else "")
        i += 1; continue
    if ln == "@PORTFOLIO":
        portfolio_table(); i += 1; continue
    if ln == "@MEIGARA":
        meigara_template(); i += 1; continue
    if ln == "@INTERVIEW":
        interview_template(); i += 1; continue
    if ln == "@REFLECT":
        reflect_template(); i += 1; continue
    if ln == "@REFS":
        references(); i += 1; continue
    if not ln.strip():
        i += 1; continue
    X.body(doc, ln.strip())
    i += 1

out = FR / "beyond_buffett_final_student_contest.docx"
doc.save(str(out))
print("saved", out, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
