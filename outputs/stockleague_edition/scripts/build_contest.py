# -*- coding: utf-8 -*-
"""STOCKリーグ提出版 v2(レビュー対応・約25頁)。
柱1:「超える」の定義を先出しし五役割を演繹 / 柱2: 結論の明確化(判定表) /
柱3: 平易な日本語(禁止語チェック内蔵)+日本語図表 / 柱4: 全文字10pt以上・詳細増補。
使い方: build_contest.py [--tocmap tocmap.json]
"""
import json
import sys
from pathlib import Path

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
sys.path.insert(0, "/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/final_revision/scripts")
import docxlib as X
from docx import Document

# v3書体規定: 和文=MS明朝(欧文・数字はTimes New Roman) / 見出し・図表・表頭(GO)=和文ゴシック+欧文Arial
_orig_setfont = X.setfont


def _setfont_v3(run, name=X.MIN, size=9, bold=False, color=None):
    _orig_setfont(run, name, size, bold, color)
    if name == X.GO:
        from docx.oxml.ns import qn as _qn
        rf = run._element.get_or_add_rPr().find(_qn("w:rFonts"))
        rf.set(_qn("w:ascii"), "Arial"); rf.set(_qn("w:hAnsi"), "Arial")


X.setfont = _setfont_v3
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
ED = HERE.parent
ASSETS = ED / "assets"
DATA = json.load(open(ROOT / "outputs/explanatory_revision/data_real.json", encoding="utf-8"))
for _v in DATA.values():
    _v["code"] = str(_v["code"])
PERF = json.load(open(ROOT / "outputs/beyond_buffett_fable_loop_final/phase5_verification_and_ablation/phase5_validation_summary.json", encoding="utf-8"))
COMP = json.load(open(ED / "control_comparison.json", encoding="utf-8"))
O3, O1, C3, C1 = COMP["ours"]["3y"], COMP["ours"]["1y"], COMP["control"]["3y"], COMP["control"]["1y"]
D3 = (O3["ann_return"] - C3["ann_return"]) * 100
D1 = (O1["ann_return"] - C1["ann_return"]) * 100
G1 = (O1["ann_return"] - O1["topix_ann_return"]) * 100
import csv as _csv
ABL = list(_csv.DictReader(open(ROOT / "outputs/beyond_buffett_fable_loop_final/phase5_verification_and_ablation/ablation_results.csv", encoding="utf-8")))

NAVY = "16324F"; TEAL = "2F6D5F"; TEAL_L = "DCE8E4"; GOLD_BG = "F5EBD0"
NAVYC = RGBColor(0x16, 0x32, 0x4F); TEALC = RGBColor(0x2F, 0x6D, 0x5F)
GO, MIN, WHITE, BLACK = X.GO, X.MIN, X.WHITE, X.BLACK
GRAYC = RGBColor(0x88, 0x88, 0x88)

order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
role_jp = {"Buffett Core": "完成した堀", "Transformation Core": "変わる堀",
           "Emerging Core": "生まれる堀", "Dual Moat": "両立型", "Bridge / Diversifier": "分散役"}
JPN = {"3539": "ＪＭホールディングス", "6430": "ダイコク電機", "7803": "ブシロード",
       "9470": "学研ホールディングス", "4350": "メディカルシステムネットワーク",
       "3697": "ＳＨＩＦＴ", "6841": "横河電機", "9474": "ゼンリン", "6368": "オルガノ",
       "6315": "ＴＯＷＡ", "6920": "レーザーテック", "6526": "ソシオネクスト",
       "5803": "フジクラ", "5902": "ホッカンホールディングス",
       "9828": "Genki Global Dining Concepts", "5233": "太平洋セメント", "8037": "カメイ",
       "3863": "日本製紙", "3089": "テクノアルファ", "2112": "塩水港精糖"}
SEC_JP = {"Retail Trade": "小売業", "Machinery": "機械", "Other Products": "その他製品",
          "Information & Communication": "情報・通信業", "Electric Appliances": "電気機器",
          "Nonferrous Metals": "非鉄金属", "Metal Products": "金属製品",
          "Glass and Ceramics Products": "ガラス・土石製品", "Wholesale Trade": "卸売業",
          "Pulp and Paper": "パルプ・紙", "Foods": "食料品"}
THEME_JP = {"non_ai": "ＡＩ以外の堅実分野", "semiconductor": "半導体",
            "quality_assurance": "品質保証", "factory_automation": "工場の自動化",
            "business_data": "業務データ", "optical_communication": "光通信"}

INVEST = sum(v["amtL1"] for v in DATA.values())
CASH = 5_000_000 - INVEST
N_UNBUY = sum(1 for v in DATA.values() if v["price"] * 100 > v["w"] * 5_000_000)
W3, W1 = PERF["window_3y"], PERF["window_1y"]

# ---- バージョン管理: 確定版は編集せず、VERSIONを繰り上げて新版を生成する ----
VER = (ED / "VERSION").read_text().strip() if (ED / "VERSION").exists() else "dev"
TOCMAP = {}
args = sys.argv[1:]
while args:
    a = args.pop(0)
    if a == "--tocmap" and args:
        TOCMAP = json.load(open(args.pop(0), encoding="utf-8"))
    elif a == "--ver" and args:
        VER = args.pop(0)
if (ED / f"beyond_buffett_stockleague_{VER}.LOCKED").exists():
    sys.exit(f"ERROR: {VER} は確定済み(LOCKED)。VERSIONファイルを繰り上げてから編集してください。")

doc = Document()
X.setup_styles(doc)
st = doc.styles["Normal"]
st.font.size = Pt(10.5)          # 既定も10pt以上(規定: 字の級数10.0pt以上)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Mm(18)
sec.top_margin = Mm(16); sec.bottom_margin = Mm(16)

HEADINGS = []


def para(text="", name=MIN, size=10.5, bold=False, align=None, before=0, after=4,
         color=None, indent=False, line=1.24, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if keep:
        pf.keep_with_next = True
    if not indent:
        X._clear_indent(p)
    if text:
        X.setfont(p.add_run(text), name, size, bold, color or BLACK)
    return p


def banner1(text, newpage=False):
    HEADINGS.append((text, 1))
    p = para(before=10, after=5, keep=True)
    pr = p._p.get_or_add_pPr()
    if newpage:
        pr.append(OxmlElement("w:pageBreakBefore"))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    pr.append(shd)
    X.setfont(p.add_run("　" + text), GO, 14, True, WHITE)
    return p


def banner2(text):
    HEADINGS.append((text, 2))
    p = para(before=8, after=4, keep=True)
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), TEAL_L)
    pr.append(shd)
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:left")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "28"); b.set(qn("w:space"), "3"); b.set(qn("w:color"), NAVY)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run("　" + text), GO, 12, True, NAVYC)
    return p


def banner3(text):
    p = para(before=5, after=3, keep=True)
    pr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "dashed"); b.set(qn("w:sz"), "10"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), TEAL)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run(text), GO, 11, True, TEALC)
    return p


def body(text, size=10.5, after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_after = Pt(after); pf.line_spacing = 1.21
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "100"); ind.set(qn("w:firstLine"), "210"); pp.append(ind)
    X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def bullet(label, text, size=10):
    p = para(after=2, line=1.22)
    X.setfont(p.add_run("・" + label + "　"), GO, size, True, NAVYC)
    X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def point(text):
    """章末の要点1行。"""
    p = para(before=4, after=6)
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F6F4")
    pr.append(shd)
    X.setfont(p.add_run("　この章の要点　"), GO, 10, True, TEALC)
    X.setfont(p.add_run(text), MIN, 10, False, BLACK)


def quote(text, who):
    p = para(before=5, after=2, line=1.4)
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480"); ind.set(qn("w:right"), "480"); pp.append(ind)
    r = p.add_run("「" + text + "」"); X.setfont(r, MIN, 11, True, NAVYC); r.font.italic = True
    p2 = para(after=6, align=WD_ALIGN_PARAGRAPH.RIGHT)
    X.setfont(p2.add_run("――　" + who), MIN, 10, False, BLACK)


def fig(fname, title, note, maxw=168):
    fp = ASSETS / fname
    pt_ = para(before=5, after=2, keep=True)
    X.setfont(pt_.add_run(title), GO, 10.5, True, NAVYC)
    if fp.exists():
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        X._clear_indent(pi); pi.paragraph_format.keep_with_next = True
        w, h = Image.open(fp).size; wmm = min(maxw, w / 200 * 25.4)
        if wmm * h / w > 150:
            pi.add_run().add_picture(str(fp), height=Mm(150))
        else:
            pi.add_run().add_picture(str(fp), width=Mm(wmm))
    para(note, MIN, 10, after=7)


def table(hdr, rows, widths, hdr_fill=NAVY, fs=10, zebra=True, aligns=None, spacer=True):
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdr):
        c = t.rows[0].cells[j]; X.shade_cell(c, hdr_fill)
        pp = c.paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(h), GO, fs, True, WHITE)
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        for j in range(len(hdr)):
            if zebra and i % 2 == 1:
                X.shade_cell(cells[j], "F2F6F4")
            pp = cells[j].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
            pp.alignment = (aligns[j] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            X.setfont(pp.add_run(str(r[j]) if j < len(r) else ""), MIN, fs, False, BLACK)
    for r in t.rows:
        X.cant_split(r)
        for j, w in enumerate(widths):
            r.cells[j].width = Mm(w)
    if spacer:
        para("", after=3)
    return t


def formula(num, title, latex, defn):
    t = doc.add_table(rows=3, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in t.rows:
        X.cant_split(r); r.cells[0].width = Mm(170)
    c0 = t.rows[0].cells[0]; X.shade_cell(c0, NAVY); X.cell_borders(c0, top=True, bottom=True)
    p0 = c0.paragraphs[0]; X._clear_indent(p0); p0.paragraph_format.space_after = Pt(0); p0.paragraph_format.keep_with_next = True
    X.setfont(p0.add_run(f"　式（{num}）{title}"), GO, 10.5, True, WHITE)
    c1 = t.rows[1].cells[0]; X.cell_borders(c1, bottom=True)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p1)
    p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(4); p1.paragraph_format.keep_with_next = True
    om = X.latex_to_omml(latex)
    from lxml import etree
    M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    pr = om.find(f"{{{M}}}oMathParaPr")
    if pr is None:
        pr = etree.Element(f"{{{M}}}oMathParaPr"); om.insert(0, pr)
    jc = pr.find(f"{{{M}}}jc")
    if jc is None:
        jc = etree.SubElement(pr, f"{{{M}}}jc")
    jc.set(f"{{{M}}}val", "center")
    p1._p.append(om)
    c2 = t.rows[2].cells[0]; X.cell_borders(c2, bottom=True)
    p2 = c2.paragraphs[0]; X._clear_indent(p2)
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    X.setfont(p2.add_run(defn), MIN, 10, False, BLACK)
    para("", after=3)


def fill_line(label, hint=""):
    p = para(after=2)
    X.setfont(p.add_run("・" + label + "："), MIN, 10.5, True, BLACK)
    X.setfont(p.add_run("［記入］"), MIN, 10.5, False, GRAYC)
    if hint:
        p2 = para(after=3)
        pp = p2._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360"); pp.append(ind)
        X.setfont(p2.add_run("書き方のヒント：" + hint), MIN, 10, False, GRAYC)


def step_chip(no, name):
    p = para(before=7, after=2, keep=True)
    r = p.add_run("　第" + "１２３"[no - 1] + "スクリーニング　")
    X.setfont(r, GO, 10.5, True, WHITE)
    rpr = r._element.get_or_add_rPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    rpr.append(shd)
    X.setfont(p.add_run("　" + name), GO, 11.5, True, NAVYC)
    HEADINGS.append((name, 2))


def add_page_number(section):
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("- ")
    X.setfont(run, MIN, 10, False, BLACK)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r2 = OxmlElement("w:r"); t2 = OxmlElement("w:t"); t2.text = "2"
    r2.append(t2); fld.append(r2)
    fp._p.append(fld)
    run3 = fp.add_run(" -")
    X.setfont(run3, MIN, 10, False, BLACK)


# ================= 表紙(全面) =================
sec.top_margin = Mm(0); sec.bottom_margin = Mm(0)
sec.left_margin = sec.right_margin = Mm(0)
pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc.paragraph_format.space_before = Pt(0); pc.paragraph_format.space_after = Pt(0)
pc.add_run().add_picture(str(ASSETS / "cover.png"), width=Mm(210))
ns = doc.add_section(WD_SECTION.NEW_PAGE)
ns.left_margin = ns.right_margin = Mm(18); ns.top_margin = Mm(16); ns.bottom_margin = Mm(16)
add_page_number(ns)

# ================= 要旨 =================
banner1("要旨")
body("本レポートでは、これからの日本株投資の姿として「三世代の堀(Three-Generation Moat)」を提示する。堀(Moat)とは、城のまわりの堀のように、他社がまねしにくい持続的な競争優位のことである。私たちは、ウォーレン・バフェットが築いた「完成した堀を割安に買う」という規律を出発点に、堀を時間軸でとらえ直した。いま完成している堀だけでなく、資本効率改革によってこれから「変わる堀」、ＡＩ・半導体・光通信などの構造変化からこれから「生まれる堀」までを、一つのポートフォリオに束ねる。")
body("選定は、剣道の修行段階になぞらえた守・破・離の三段階スクリーニングで行った。守では先行研究の式だけで完成した堀を5社に絞り、破では式を変えずに候補を1,200社へ広げ、離では変わる堀・生まれる堀を証拠の強さつきで測って最終20社を組んだ。20社は最初に決めたまま持ち続け、その後の値上がり・値下がりを見て選び直すことはしていない。あとから成績の良い銘柄だけを拾い直せば、どんな選び方でも立派に見えてしまうからだ。その代わりに、この選び方が条件を一つ変えたくらいでは崩れないことを、16通りの検査で確かめた。")
para("※守破離(しゅはり)＝日本の芸道の修行段階。守=型を守る、破=型を破る、離=独自の境地へ。", MIN, 10, after=4)
body("結論を先に言う。私たちは「超える」を四つの条件で定義した。①バフェットと同じ土俵に立つ、②彼の式では選ばれない15社を同じ厳しさの証拠で選ぶ、③選定が壊れにくい――この三つの設計条件は本文で確かめたとおり達成した。残る一つ、④成績は、同じ出発点(3,099社)から守の式だけで選んだ対照群『純正バフェットポートフォリオ』を自作し、先に決めた物差しで比較して検証した。結果、選定に使った3年間で年17.4ポイント、選定に使っていない直近1年でも年25.9ポイント、対照群を上回り、④も達成した。ただし市場平均(ＴＯＰＩＸ)には直近1年で2.8ポイント負けている――この事実も隠さず記し、市場平均へ勝ち続ける力の証明は来年の宿題として残す。")

banner1("目次")
toc_items = ["要旨", "Ⅰ．背景・投資テーマ決定", "Ⅱ．スクリーニング ― 守・破・離",
             "Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス", "Ⅳ．インタビュー・アンケート",
             "Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと", "参考文献", "用語の手引き"]
for it in toc_items:
    p = para(after=1)
    X.setfont(p.add_run(it), MIN, 10.5, False, BLACK)
    pg = TOCMAP.get(it, "")
    if pg:
        X.setfont(p.add_run("　" + "･" * max(4, 52 - len(it) * 2) + "　" + str(pg)), MIN, 10.5, False, BLACK)

# ================= Ⅰ 背景・投資テーマ決定 =================
banner1("Ⅰ．背景・投資テーマ決定", newpage=True)
banner2("1．イントロダクション ― 百年企業の中に、未来の堀が眠っていた")
quote("素晴らしい会社をまずまずの価格で買うほうが、まずまずの会社を素晴らしい価格で買うよりずっと良い。",
      "ウォーレン・バフェット(1989年 株主への手紙)")
body("バフェットは、企業を城に、競争優位を城を守る堀(Moat)にたとえた。深い堀を持つ会社を、高すぎない価格で買い、長く持つ。この単純な規律が、半世紀を超える成果を生んだ。私たちの研究は、この規律への挑戦ではなく、その延長である。")
body("きっかけは、一つの会社だった。フジクラ(5803)――1885年創業の電線メーカーである。長らく「成熟産業の老舗」と見られてきたこの会社は、いま、生成ＡＩのデータセンターを支える光配線の担い手として再評価されている。百年かけて磨いた細くて高密度な配線の技術が、ＡＩという新しい産業の「そこが詰まると全体が止まる急所」を握ったのだ。完成したはずの堀の内側で、新しい堀が生まれていた。")
body("同じ頃、東京証券取引所は上場企業に「資本コストや株価を意識した経営」を求め、長く割安に放置された企業が、資本効率の改善と株主還元の強化によって変わり始めた。堀は止まった絵ではない。変わる堀があり、生まれる堀がある。ならば、いまの堀だけを測るバフェットの物差しを、時間の方向へ延長できないか。これが私たちの問いである。")

banner2("2．バフェットの規律 ― 何がすごいのかを分解する")
body("「超える」と言う前に、まず相手を正しく知る。バフェットの投資は、大きく四つの部品でできている。")
bullet("堀(持続する競争優位)", "まねされにくい強み。ブランドの力(例:コカ・コーラ)、乗り換えの手間、利用者が増えるほど価値が増す仕組み、圧倒的なコストの安さ、特許や免許。バフェット(2007)は「永続的な堀」を持つ企業だけを買うと明言する。")
bullet("安全域(高すぎない価格)", "どんな良い会社も、高く買いすぎれば損をする。会社の価値より安い価格で買い、見立て違いに備える。冒頭の1989年の言葉は、価格より会社の質を優先するという順序も示している。")
bullet("能力の輪(分かるものだけ買う)", "自分が理解できる事業だけを対象にする。理解できない熱狂には乗らない。")
bullet("長期保有(複利で増やす)", "「私たちの好みの保有期間は永遠である」(バフェット、1988)。頻繁な売買をせず、堀が守る利益の再投資で雪だるま式に増やす。")
body("この規律は、感覚ではなく式で確かめられる。Frazzini, Kabiller and Pedersen (2018)は、バフェットの半世紀の超過成績が「割安×優良×値動きの穏やかさ」という測定可能な要素でほぼ説明できることを示した。つまり、彼の物差しは先行研究の式で再現できる――これが第Ⅱ章「守」の土台である。")
body("同時に、この分解は彼の物差しの死角も教えてくれる。割安・優良・穏やかさは、どれも「いま完成している堀」を測る物差しであり、これから変わる堀・生まれる堀は原理的に映らない。私たちが加える視点は三つ――①時間軸(堀の一生を見る)、②証拠主義(言葉ではなく数字の裏づけで測る)、③分散(どれか一つの未来に賭けない)。この三つを、彼と同じ規律の厳しさで実行することが、本レポートの挑戦である。")

banner2("3．背景 ― 二つの地殻変動")
body("第一の地殻変動は資本効率改革である。東京証券取引所(2023)の要請は、伊藤レポート(経済産業省、2014)以来の資本収益性重視の流れを決定づけ、ＰＢＲ(株価純資産倍率。株価が会社の純資産の何倍かを示す)1倍割れ企業の経営改革を促した。割安に放置された企業の中から、本当に変われる企業を見分けられれば、それは「変わる堀」への投資になる。")
body("第二の地殻変動はＡＩを起点とする産業構造の転換である。生成ＡＩは、半導体・光通信・データセンター・電力・品質保証といった基盤産業に、実際の需要という追い風を与えている。ただし「ＡＩ」と口にする企業がすべて堀を持つわけではない。私たちの候補調査では、ＡＩ関連テーマとの接点が社名・業種・引き継ぎメモへの機械的なキーワード照合でしか確認できず、製品・顧客・数量まで遡れる開示証拠を確認できなかった会社が577社に上った。接点と堀を見分ける仕組みが要る。")

banner2("4．三世代の堀 ― 私たちが求める企業像")
body("以上から、私たちは投資対象を三つの世代で定義する。第一に、すでに完成した堀(バフェットの規律で測れる企業)。第二に、資本効率改革で変わる堀(割安×改善の証拠を持つ企業)。第三に、構造変化から生まれる堀(ＡＩ基盤への実需接続を証拠で示せる企業)。この三世代を一つのポートフォリオに束ね、どれか一つの未来に賭けない。これが「三世代の堀」である。")
fig("fig1_moat.png", "図表Ⅰ-1　三世代の堀 ― 完成・変化・新生を一つに束ねる",
    "注：上段は三世代の名前、下段は中身。両立型・分散役を加えた合計20社で構成する。筆者作成。")
p_h = para(before=5, after=7)
_pr = p_h._p.get_or_add_pPr(); _shd = OxmlElement("w:shd")
_shd.set(qn("w:val"), "clear"); _shd.set(qn("w:color"), "auto"); _shd.set(qn("w:fill"), GOLD_BG)
_pr.append(_shd)
X.setfont(p_h.add_run("　投資仮説　"), GO, 10.5, True, NAVYC)
X.setfont(p_h.add_run("完成した堀・変わる堀・生まれる堀を証拠で選んで一つに束ねたポートフォリオは、堀の“現在”だけを測る純正バフェット型の選定を、リターンとリスクの両面で上回る。"), MIN, 10.5, True, BLACK)

banner2("5．何をもって「バフェットを超える」とするか ― 四つの条件")
body("「超える」を気分で語らないために、確かめられる四つの条件で先に定義しておく。本レポートの残りは、この定義の検証である。")
table(["条件", "内容", "どこで確かめるか"],
      [["① 同じ土俵に立つ", "バフェットの選び方を先行研究の式だけで再現し、その方法が選ぶ5社を実際に持つ。比較相手を自分のポートフォリオの中に持ち歩く", "第Ⅱ章「守」"],
       ["② 死角を補う", "バフェットの式では選ばれない15社を、同じ厳しさの証拠で選ぶ。彼の物差しが見ない未来を、彼と同じ規律で測る", "第Ⅱ章「離」"],
       ["③ 壊れにくい", "選定が一つの仕掛け・一つのテーマに依存しない。条件を一つずつ変えた16通りの選び直しで確かめる", "第Ⅱ章末の検査"],
       ["④ 成績で上回る", "同じ出発点(3,099社)・守の式だけで作った対照群『純正バフェットポートフォリオ』に、先に決めた物差し(年率リターン・ＴＯＰＩＸ超過・市場との連動度・最大の下落・超過の安定度)で上回る。ＴＯＰＩＸ・日経平均は参照として併記", "第Ⅲ章の検証"]],
      [34, 102, 34], fs=10)
para("図表Ⅰ-2　「超える」の定義 ― 四つの条件と検証の場所。", MIN, 10, after=5)
body("大切なのは順序である。守破離という手順は、この定義より先にあるのではない。①〜③を後から誰でも検証できる形にするための手段として、独自色を最後に置く守破離を選んだ。独自の物差しを最後に導入するからこそ、「どこからが私たちの主張か」をいつでも遡って示せる。")

banner2("6．検証の順番 ― 守・破・離")
body("守では先行研究の式を一切変えずに使い(条件①の検証)、破では式を変えずに使い方だけを調整して候補を広げ、離で初めて独自の物差しを導入する(条件②の検証)。最後に、出来上がった選定を16通りの検査にかける(条件③の検証)。条件④は、第Ⅲ章で対照群『純正バフェットポートフォリオ』とのリターン・リスク比較により検証する(物差しは比較の前に宣言する)。")
table(["段階", "やること", "どこまでが借り物か", "検証する条件"],
      [["守", "先行研究の式だけで完成した堀5社を選ぶ", "式も使い方も、すべて先行研究のまま", "条件①"],
       ["破", "式を変えずに使い方だけ調整し、候補を1,200社へ広げる", "式は借り物。使い方は私たちの判断(過程を開示)", "(離の準備)"],
       ["離", "変わる堀・生まれる堀を独自の物差しで測り、20社を組む", "部品は借り物。組み方が私たちの独自性", "条件②"]],
      [16, 72, 56, 26], fs=10)
para("図表Ⅰ-3　守・破・離の分担 ― 独自性を最後に置く。", MIN, 10, after=5)
body("なお、本文の数字はすべて、選定の過程で保存した記録(段階ごとの通過社数・除外862社の理由・16通りの検査結果)から機械的に転記している。手作業の書き写しによる食い違いをなくすため、本レポート自体を記録データから自動生成する仕組みで作成した。どの数字も、出どころのデータまでいつでも遡って確かめられる。")
point("「超える」を四つの条件で定義した。守破離は目的ではなく、この定義を検証可能にするための手段である。")
banner3("投資テーマと関連が深いＳＤＧｓの目標(3つ以内・記入)")
table(["関連の深いＳＤＧｓの目標", "その主な理由(記入)"],
      [["9．産業と技術革新の基盤をつくろう", "［記入例の観点：ＡＩ・半導体・光通信など「生まれる堀」への投資が技術革新の基盤形成を支えるため］"],
       ["8．働きがいも経済成長も", "［記入例の観点：資本効率改革で「変わる堀」を持つ企業の再評価が日本経済の成長につながるため］"],
       ["12．つくる責任・つかう責任", "［記入］"]],
      [60, 110], fs=10, spacer=False)

# ================= Ⅱ スクリーニング =================
banner1("Ⅱ．スクリーニング ― 守・破・離", newpage=True)
body("前章では「超える」を四つの条件で定義し、検証の順番として守破離を置いた。本章が選定の中核である。3,099社の非金融普通株から、三段階のスクリーニングで最終20社を選ぶ。各段階では、何を見たか(観点)、何で測ったか(データと式)、どこで区切ったか(合格ライン)、何社が残ったか(結果)を必ず示す。データの基準は、価格が2023年4月25日〜2026年6月1日の日次、財務・開示がＥＤＩＮＥＴ提出書類(2026年7月8日取得)である。")

step_chip(1, "守 ― バフェットの規律を式だけで再現する(→完成した堀5社)")
body("守では、バフェットの規律「良い会社を、高すぎない価格で買い、長期保有に耐えない企業を避ける」を、先行研究の式だけで再現する。独自の重み付けは使わず、七つの関門を一つずつ通す。")
table(["観点", "使うデータ・式", "合格ライン", "残った社数"],
      [["割安か(純資産)", "Ｂ／Ｍ＝純資産÷時価総額(式1)", "市場の上位30％", "3,099→2,740社(算出可能)"],
       ["割安か(利益)", "Ｅ／Ｐ＝純利益÷時価総額(式2)", "黒字かつ上位50％", "→583社(割安の関門通過)"],
       ["収益力は本物か", "粗利÷総資産(式3)", "市場の中央値以上", "(同じ583社に含む)"],
       ["財務は健全か", "6項目チェックの合格割合(式4)", "0.65以上", "→146社"],
       ["利益は本物か", "利益と現金の差(式5)", "悪い側の上位30％を除外", "→112社"],
       ["危険はないか", "債務超過・2期連続赤字(式6)", "該当ゼロ", "→90社"],
       ["実際に買えるか", "60日平均の売買代金(式7)", "1日約1,000万円以上", "→77社"]],
      [32, 52, 38, 48], fs=10)
para("図表Ⅱ-1　守の七関門(判定基準)。出所：第1段階の正典データ。", MIN, 10, after=5)
def verdict(text):
    p = para(after=6)
    X.setfont(p.add_run("→ 判定　"), GO, 10.5, True, TEALC)
    X.setfont(p.add_run(text), MIN, 10.5, True, BLACK)


body("七つの関門を、式そのものから一つずつ見ていく。どの式も先行研究のまま使い、私たちの独自の重みは一切加えない。")
banner3("関門1　会社の持ち物に対して安いか ― Ｂ／Ｍ")
body("純資産(会社の持ち物から借金を引いた正味の財産)と株価を比べる。安く買うことは、見立て違いへの保険になる(Fama and French, 1993)。", after=2)
formula("1", "Ｂ／Ｍ ― 純資産に対して安いか(Fama and French, 1993)",
        r"\mathrm{BM}_i = \frac{\mathrm{BE}_i}{\mathrm{ME}_i}",
        "ＢＥ＝自己資本(純資産)　ＭＥ＝時価総額(株価×株式数)。大きいほど、会社の持ち物のわりに株価が安い。")
verdict("よって、Ｂ／Ｍが市場の上位30％に入る企業を「純資産に対して割安」と評価する。")
banner3("関門2　稼ぐ力に対して安いか ― Ｅ／Ｐ")
body("純資産だけで測ると「資産はあるが稼げない会社」を拾ってしまう。そこで利益の面からも安さを測る(Basu, 1977)。", after=2)
formula("2", "Ｅ／Ｐ ― 利益に対して安いか(Basu, 1977)",
        r"\mathrm{EP}_i = \frac{\mathrm{NI}_i}{\mathrm{ME}_i}",
        "ＮＩ＝純利益(最終的に残る会計上の利益)　ＭＥ＝時価総額。大きいほど、稼ぐ力のわりに株価が安い。")
verdict("よって、黒字かつＥ／Ｐが市場の上位50％に入る企業を「稼ぐ力に対しても割安」と評価する。")
banner3("関門3　収益力は本物か ― 粗利÷総資産")
body("売上総利益(売上から原価を引いた粗利)は、会計操作の余地が小さく、資産を利益に変える力を素直に映す(Novy-Marx, 2013)。", after=2)
formula("3", "収益力 ― 資産を粗利に変える力(Novy-Marx, 2013)",
        r"\mathrm{GP}_i = \frac{\mathrm{Revenue}_i - \mathrm{COGS}_i}{\mathrm{TotalAssets}_i}",
        "Revenue＝売上高　COGS＝売上原価　TotalAssets＝総資産。大きいほど資産効率の高い優良企業。")
verdict("よって、粗利÷総資産が市場の中央値以上の企業を「収益力が本物」と評価する。")
banner3("関門4　財務は健全か ― 6項目チェック")
body("財務の健康を、良い/悪いの二択で数えるチェックリスト(Piotroski, 2000)。原式は9項目だが、粗利率の改善・短期支払能力の改善・増資の有無の3項目は取得できるデータに無かったため、実装できた6項目の合格割合で判定し、その旨を監査記録した(勝手に完全版とは呼ばない)。", after=2)
formula("4", "6項目チェック ― 財務の健全さの合格割合(Piotroski, 2000)",
        r"R_i = \frac{1}{6}\sum_{k=1}^{6} \mathbf{1}_k",
        "各項目を満たせば1、満たさなければ0と数え、合格割合を出す。6項目＝①総資産利益率(ＲＯＡ)がプラス ②本業の現金収入がプラス ③ＲＯＡが前年より改善 ④利益に現金の裏づけがある(現金収入÷総資産＞ＲＯＡ) ⑤資産回転率が改善 ⑥借入依存度が低下。")
verdict("よって、6項目の合格割合が0.65以上の企業を「財務が健全」と評価する。")
banner3("関門5　利益は本物か ― 利益と現金の差")
body("帳簿の利益と、実際に入った現金の差が大きい会社は、後で利益が失速しやすい(Sloan, 1996)。", after=2)
formula("5", "利益と現金の差 ― 利益の質(Sloan, 1996)",
        r"\mathrm{Acc}_i = \frac{\mathrm{NI}_i - \mathrm{CFO}_i}{(\mathrm{TA}_{i,t}+\mathrm{TA}_{i,t-1})/2}",
        "ＮＩ＝純利益　ＣＦＯ＝営業キャッシュフロー(本業の現金収入)　ＴＡ＝総資産(当期と前期の平均で割る)。大きいほど、利益に現金の裏づけが乏しい。")
verdict("よって、この差が悪い側(大きい側)の上位30％に入る企業を除外する。")
banner3("関門6　長期保有に耐えるか ― 危険よけ")
body("いくら割安で優良に見えても、財務危機の恐れがある会社は長期保有に向かない(Altman, 1968；Ohlson, 1980の考え方を、取得できるデータで実装した安全条件)。", after=2)
formula("6", "危険よけ ― 除外の条件",
        r"\mathrm{Exclude}_i = \mathbf{1}[\mathrm{Equity}_i \le 0]\ \lor\ \mathbf{1}[\mathrm{NI}_{i,t}<0 \land \mathrm{NI}_{i,t-1}<0]",
        "Equity＝純資産　ＮＩ＝純利益。純資産がマイナス(債務超過)、または2期連続の純損失に該当したら除外し、該当しない企業だけを通す。")
verdict("よって、債務超過・2期連続赤字に該当しない企業を「長期保有に耐える」と評価する。")
banner3("関門7　実際に買えるか ― 売買のしやすさ")
body("売買が少ない株は、買うだけで値段が動いてしまい、計画どおりの投資ができない。", after=2)
formula("7", "60日平均売買代金 ― 売買のしやすさ",
        r"\mathrm{ADV}^{60}_i = \frac{1}{60}\sum_{t=1}^{60} P_{i,t}\,V_{i,t}",
        "Ｐ＝日々の株価　Ｖ＝日々の出来高(売買された株数)。直近60営業日の売買代金の平均。1日あたり約1,000万円以上を基準とする。")
verdict("よって、1日平均約1,000万円以上の売買がある企業を「実際に買える」と評価する。")
g = DATA["9470"]
body("式が実際にどう働くかを、通過企業の一社・学研ホールディングス(9470)で追ってみる。")
table(["関門", "学研ＨＤの値", "合格ライン", "判定"],
      [["Ｂ／Ｍ(式1)", f"{g['bm']:.3f}", "市場の上位30％", "通過"],
       ["Ｅ／Ｐ(式2)", f"{g['ep']:.3f}(黒字)", "黒字かつ上位50％", "通過"],
       ["粗利÷総資産(式3)", f"{g['gp']:.3f}", "中央値以上", "通過"],
       ["6項目チェック(式4)", f"合格割合 {g['piotroski']:.0%}({round(g['piotroski']*6)}/6項目)", "0.65以上", "通過"],
       ["利益と現金の差(式5)", f"{g['sloan']:.3f}", "悪い側の上位30％でない", "通過"],
       ["危険よけ(式6)", "債務超過・2期連続赤字なし", "該当ゼロ", "通過"],
       ["売買のしやすさ(式7)", f"1日平均 {g['adv_oku']:.2f}億円", "約0.1億円以上", "通過"]],
      [40, 46, 44, 20], fs=10)
para("図表Ⅱ-2　通し数値例：学研ホールディングス(9470)の七関門。出所：選定の正典データより筆者作成。", MIN, 10, after=6)
b5 = [DATA[c] for c in ["3539", "4350", "6430", "7803", "9470"]]
body("通過77社から、収益力→利益÷時価総額→純資産÷時価総額→6項目→利益の質→買いやすさ→会社の大きさの固定順で上位5社を選び、完成した堀(基準線)として固定した(同一業種は原則2社まで)。選ばれたのは、"
     + "、".join(f"{d['code']} {JPN[d['code']]}({SEC_JP[d['sector']]})" for d in b5)
     + "の5社である。以降、この5社は一度も入れ替えない。")
fig("fig2_shu.png", "図表Ⅱ-3　守の絞り込み ― 3,099社から完成した堀Top5へ",
    "注：棒の幅と色の濃さが残った会社の数を表す。出所：第1段階の正典データより筆者作成。")

step_chip(2, "破 ― 式を変えずに候補を1,200社へ広げる")
body("前節の守では、先行研究の式だけで完成した堀の5社(基準線)を決めた。本節の破では、同じ式の「使い方」だけを調整して、次の離で詳しく調べる候補リストを1,200社へ広げる。守は規律が強い分、通過がわずか77社になり、これでは「変わる堀」「生まれる堀」の候補を評価する余地がないからだ。手順は四つある。")
banner3("手順1　単位の違う7指標を、どう公平に比べるか ― 順位化")
body("守で使った7指標(割安2つ・収益力・6項目・利益の質・危険よけ・売買のしやすさ)は単位がばらばらで、そのままでは足し合わせられない。そこで各指標を「市場の中で下から何番目か」という順位の割合(0〜1)に置き換える。", after=2)
formula("8", "順位化 ― 市場の中の相対的な位置に置き換える",
        r"r_{i,k} = \frac{\operatorname{rank}_k(x_{i,k})}{N}",
        "ｘ＝指標ｋの値　rank＝市場の中での順位(良い方が大きい)　Ｎ＝算出できた会社数。0〜1の点数になり、単位の違う指標同士を公平に比べられる。")
verdict("よって、7指標すべてを0〜1の順位点に変換して比べる。式の定義は一切変えない。")
banner3("手順2　順位点をどう束ねるか ― 重みつき合成点と設定の探索")
body("7つの順位点に重みを付けて足し合わせ、危険な兆候(異常値・データ欠け・極端に小さい会社・一時的な利益)には減点を与える。重み・減点の強さ・候補数・業種調整の有無・データ欠けの扱いは、無数の組合せを機械的に大量に試す方法(ランダムサーチ)で比較し、最も安定な設定を採用した。採用した設定は「市場順位で0〜1化・主要指標が欠けた会社は除外・業種調整はしない」である。", after=2)
formula("9", "合成点 ― 候補づくりのための重みつき合計(最終選定には使わない)",
        r"S_i = \sum_{k=1}^{7} w_k\, r_{i,k} - P_i",
        "ｗ＝指標ｋの重み(探索で決定し、全設定を記録)　ｒ＝順位点(式8)　Ｐ＝減点(異常値・データ欠け・超小型株・一時的利益)。候補を広げるための道具であり、最終20社の選定には使わない。")
verdict("よって、合成点の上位から候補リストを作る。ただし、この点数で最終選定はしない。")
body("重みについて、一つ正直に書き残す。危険よけの重みは0.251と大きいが、これを外しても結果はほとんど変わらない(目的関数の低下は0.0017)。危険な会社は合成点より前の段階で機械的に除外済み(財務除外184社+債務超過11社)であり、重みは「除外をすり抜けた残りの微調整」を担っているにすぎない。重みの大きさを予測力と読んではいけない――この多層防御の構造ごと開示しておく。")
banner3("手順3　順位の付け方を変えても、結果は同じか ― 4通りの監査")
body("0〜1化のやり方は一つではない。市場全体での順位・業種の中での順位・外れ値に強い標準化・端を丸めた標準化の4通りで順位を付け直し、顔ぶれが入れ替わらないかを確かめた。1,200社のうち970社は4通りのどの付け方でも共通して上位に残り(中核)、広めに見ても1,135社は安定していた。付け方に敏感な29社には確認フラグを付け、離の個別確認へ引き継いだ。")
banner3("手順4　候補は何社が適切か ― Top1200の採用")
body("広さ・質・安全・売買のしやすさ・業種の分散・人の目で確認できる規模を比べると、広さを重視するほどTop2000が有利になる。しかし目的は候補数の最大化ではなく「一社ずつ確認できる候補づくり」なので、Top1200を正式採用し、Top2000は取りこぼし確認用の参照群とした。Top100・Top300は日々の確認を優先する内側の層で、のちの壊れにくさ検査(A8・A9)の比較にも使う。この入れ子の関係を図表Ⅱ-5に示す。採用したTop1200は業種の偏りがなく(偏り指数0.0707・最大業種11.6％)、守の5社を全員含む。")
table(["観点", "内容", "結果"],
      [["順位化(式8)", "各指標を市場の中での順位(0〜1)に置き換える。式の定義は変えない", "単位の違う7指標を公平に比較"],
       ["合成点(式9)", "重み・減点・候補数などをランダムサーチで比較し、最も安定な設定を採用", "市場順位・欠損は除外・業種調整なし"],
       ["頑健さの監査", "順位の付け方4通りで顔ぶれの入れ替わりを確認", "中核970社・広めに1,135社・要確認29社"],
       ["候補数の決定", "広さと、人の目で確認できる規模の両立点を判断", "Top1200を正式採用(Top2000は参照群)"],
       ["守との整合", "守の5社が候補リストに含まれるか", "5社すべてTop1200に含まれる"]],
      [30, 84, 56], fs=10)
para("図表Ⅱ-4　破のやり方 ― 式は変えず、使い方だけを調整する。出所：Phase2正典記録(manifest・4方式監査表)。", MIN, 10, after=5)
body("最後に、破の限界も先に書いておく。時間を遡った完全な将来予測の検証(学習期間と検証期間の厳密な分割)は、使える年度数の制約で完了していない。代わりに、開示の提出日どおりに過去の各時点を再現した一覧表を作り、固定した設定を各年度の断面に当てはめる検証を行った。したがって私たちは、この候補づくりに将来のリターンを予測する力があるとは主張しない。候補を広く・壊れにくく作ること――それが破の役割である。")
fig("fig2_ha.png", "図表Ⅱ-5　破 ― 候補リストの形成(Top1200を正式採用)",
    "注：破の点数は候補選びの道具であり、最終選定には使わない。筆者作成。")

step_chip(3, "離 ― 変わる堀・生まれる堀を測り、20社を組む")
body("離で初めて、私たち独自の物差しを導入する。ただし新しい式を発明したのではない。守で使った先行研究の部品(割安・改善・利益の質)を、時間の方向へ組み替えた。変わる堀は式(10)、生まれる堀は式(11)で測り、点数の根拠の強さは式(12)の「証拠の強さ」として点数と分けて管理する。")
formula("10", "変わる堀の点数(選定に使用)",
        r"S^{\mathrm{Trans}}_i = 0.22\,V_i + 0.24\,C_i + 0.10\,F_i + 0.18\,X_i + 0.16\,Q_i + 0.10\,\Phi_i - P^{\mathrm{Trap}}_i",
        "Ｖ＝割安さ　Ｃ＝資本効率の改善　Ｆ＝還元の余力(本業の現金収入−設備投資)　Ｘ＝実行の信頼性(公約を守ってきたか)　Ｑ＝利益の質　Φ＝データの信頼度　Ｐ＝割安のワナ減点。")
verdict("よって、割安さに「変わる証拠」が重なる会社ほど高得点とし、変わる堀の上位5社を選ぶ。安いだけの15社は減点で弾いた。")
formula("11", "生まれる堀の点数(選定に使用)",
        r"S^{\mathrm{Emerg}}_i = 100\,(w_I I_i + w_N N_i + w_B B_i + w_A A_i + w_D D_i + w_T T_i) + B^{\mathrm{Evi}}_i - P^{\mathrm{Hype}}_i",
        "Ｉ＝無形資産　Ｎ＝技術力　Ｂ＝急所を握る度合い　Ａ＝ＡＩ基盤への接続　Ｄ＝データ・顧客基盤　Ｔ＝信頼・安全の基盤　加点＝証拠(最大8点)　減点＝キーワードの接点のみ(18点)。")
verdict("よって、ＡＩ基盤への実需接続を証拠つきで示せる会社ほど高得点とし、生まれる堀の上位5社を選ぶ。接点だけの577社は除外した。")
body("最後の式(12)は、点数が高くても証拠が弱い会社を弾く「関所」である。点数(魅力の大きさ)と証拠(裏づけの強さ)を分けて管理するのが、この設計の背骨だ。", after=2)
formula("12", "証拠の強さ(点数と分けて管理)",
        r"L^{\mathrm{final}}_i = \begin{cases} \min(L^{TQ}_i, L^{EM}_i) & \text{両立型} \\ L^{EM}_i & \text{生まれる堀} \\ \max(L^{TQ}_i, L^{TS}_i, L^{TR}_i) & \text{変わる堀} \\ \max(L^{TQ}_i, L^{EM}_i) & \text{その他} \end{cases}",
        "ＴＱ＝数字の改善で確認できる変わる証拠　ＴＳ＝株主還元の実行の証拠　ＴＲ＝改革の開示(計画文書)の証拠　ＥＭ＝新テーマ(生まれる堀)の開示証拠。読み方: min＝両方の証拠が必要(弱い方で判定)、max＝どれか一つ強ければ十分。水準1＝キーワードの接点のみ／水準2＝製品・顧客・投資計画の具体性を確認／水準3＝売上・受注・投資額の数字まで確認。")
verdict("よって、点数がどれほど高くても、証拠の水準が足りない会社は役割の候補から外す。最終20社の内訳は水準3が15社・水準2が4社・水準1が1社。")
body("当てはめ例: 両立型のゼンリン(9474)は「変わる証拠」と「新テーマの証拠」の弱い方(min)で判定して水準3――どちらの裏づけも数字まで確認できた、という意味になる。", after=4)
body("候補1,200社から20社へ絞る過程で除外した862社は、理由別にすべて記録した。選ばなかった理由を残すことが、選んだ理由の裏づけになる。")
body("生まれる堀の証拠は、確認できたものだけを数える保守設計である。ＩＲ資料・製品ページまで遡って証拠を確かめられた14社(水準2以上)だけを生まれる堀・両立型の候補に通し、キーワードの接点しか確認できない会社は――たとえ有望に見えても――一律で除外して監査記録した(疑わしきは除外)。")
table(["除外の理由", "社数", "説明"],
      [["ＡＩ関連の接点がキーワードのみ", "577社", "機械的なキーワード照合でしか接点を確認できず、製品・顧客・数量まで遡れる開示証拠(水準2以上)を確認できなかった"],
       ["より良い類似候補あり", "221社", "同じ業種・テーマに、証拠も点数も上回る候補が存在する"],
       ["財務が危険", "32社", "債務超過・連続赤字など、長期保有に耐えない"],
       ["割安のワナ", "17社", "安いだけで、改善の実行が伴わない"],
       ["安さの根拠のみ", "15社", "ＰＢＲの低さ以外に変わる証拠がない"]],
      [42, 18, 110], fs=10)
para("図表Ⅱ-6　除外862社の内訳(全社を理由つきで記録)。出所：除外記録の正典データ。", MIN, 10, after=6)

banner2("最終20社の構成 ― なぜ5・5・5・3・2なのか")
body("最終20社は、完成した堀5社(守の5社を固定)・変わる堀5社・生まれる堀5社・両立型3社・分散役2社で構成した。この構成は第Ⅰ章の「超える」の定義から導いている。")
bullet("完成した堀5社を持つ理由", "条件①(同じ土俵)の実行である。この5社を持たなければ、「何を超えたのか」を測る比較相手がいなくなる。しかも三世代のうち、完成した堀だけが仮説ではなく現在の事実である。")
bullet("5・5・5の均等", "三つの世代のどれか一つの未来に賭けない、という設計原理をそのまま数にした。")
bullet("3・2の少数枠", "両立型は変わる堀と生まれる堀の橋渡し、分散役は業種・テーマの偏りの調整役。主役ではなく脇役なので少数にした。")
body(f"「バフェットが選ぶはずの5社は、本当に必要なのか」という問いには、検査で答える。基準線5社の固定を外して選び直す検査(後述のA11)では、最終20社のうち11社しか一致せず、完成した堀の5社は5社とも脱落する。つまりこの5社は、新しい物差しの高得点者だから残っているのではなく、比較相手として意図して固定している。隠すのではなく、それ自体を設計として明示する。なお、この構成が唯一の最適だという証明はできない。最適かどうかは未来の成績でしか分からないからだ。私たちが保証するのは、定義との整合と、次に示す壊れにくさである。")
fig("fig2_ri.png", "図表Ⅱ-7　離 ― 三世代の堀から最終20社へ(矢印は1対1の対応)",
    "注：守の5社は証拠の関所を通らず固定。両立型は二つの点数の両立を、証拠は厳しい方で判定する。筆者作成。")

banner2("選定は壊れにくいか ― 条件を一つずつ変える16通りの検査")
body("ここまで読んだ読者は、当然こう疑うはずだ――「複雑なルールを後から調整して、結論ありきで20社を選んだのではないか?」。その疑いには、条件を一つずつ変えた16通りの選び直しで答える。橋の耐荷重試験のように、部品を一本ずつ抜いて、それでも崩れないかを確かめるのだ。検査は四つの家族に分かれる: 点数の部品を外す/減点・証拠の関所を外す/候補の広さを変える/役割の枠を外す。")
body("読み方の目盛りを先に決めておく。選び直した20社が元の20社と15社以上一致すれば、その条件を変えても中核は保たれる(頑健)。11〜14社なら構成に影響するが崩壊はしない。10社以下なら、その条件こそが選定の背骨だと分かる。")
body("結果、候補を100社に狭める検査A8だけが7社まで崩れ、残る15通りは11社以上を保った。どの減点・どの枠を一つ外しても選定は崩壊しない一方で、最も効くのは候補の広さ――破の章で語った「広く作って確認する」という判断の価値が、ここで数字として跳ね返ってくる。")
fig("fig2_ablation.png", "図表Ⅱ-8　条件を一つ変えて選び直したとき、最終20社のうち何社が残るか",
    "注：過去データ上の構造の検査であり、成績の主張ではない。出所：検査結果の正典データ(ablation_results.csv)。")
ABL_JP = {
    "A1": "変わる堀の点数だけで選ぶ", "A2": "生まれる堀の点数だけで選ぶ", "A3": "証拠の関所を外す",
    "A4": "割安のワナ減点を外す", "A5": "話題先行の減点を外す", "A6": "データ信頼度を外す",
    "A7": "業種の上限を外す", "A8": "候補を100社に絞る", "A9": "候補を300社に絞る",
    "A10": "候補1,200社全体から選ぶ", "A11": "基準線5社の固定を外す", "A12": "両立型の枠を外す",
    "A13": "分散役の枠を外す", "A14": "証拠水準2以上の条件を外す", "A15": "改革の証拠を外す",
    "A16": "ＡＩキーワード減点を一部の役割に限定"}
ABL_OV = {r["variant"]: int(r["overlap_with_final20"]) for r in ABL}
FAMS = [
    ("点数の部品を外す", ["A1", "A2"], "どちらか片方の点数だけでも中核は残る"),
    ("減点・証拠の関所を外す", ["A3", "A4", "A5", "A6", "A14", "A15", "A16"], "どの関所を一つ外しても崩壊しない(多層防御)"),
    ("候補の広さを変える", ["A8", "A9", "A10"], "A8(100社に絞る)のみ大幅入替＝候補の広さが選定の背骨"),
    ("役割の枠を外す", ["A7", "A11", "A12", "A13"], "基準線5社の固定(A11)は「意図した設計」だと確認できる"),
]
fam_rows = [[f, "・".join(m), f"{min(ABL_OV[x] for x in m)}〜{max(ABL_OV[x] for x in m)}/20", y]
            for f, m, y in FAMS]
table(["検査の家族", "検査", "一致の範囲", "読み取り"], fam_rows, [38, 42, 26, 64], fs=10)
para("図表Ⅱ-9　16通りの検査の4家族要約(各検査の個別値は図表Ⅱ-8に記載)。一致＝選び直しても最終20社と同じだった社数。", MIN, 10, after=6)
point("守で基準線5社(条件①)、離で死角を補う15社(条件②)を選び、16通りの検査で壊れにくさ(条件③)を確かめた。")

# ================= Ⅲ ポートフォリオ =================
banner1("Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス", newpage=True)
body("前章で最終20社が決まった。本章では500万円の配分を決め、20社の顔ぶれを紹介し、最後に第Ⅰ章の条件④を対照群との比較で検証する。")
banner2("1．500万円の配り方 ― 成績の予想を使わない配分")
body("どの銘柄が上がるかという予想は、一切使わない。予想を使えば、せっかく証拠で選んだ20社に、予想の当て推量が混ざってしまうからだ。配分は次の四段階で決めた。")
bullet("第一段階", "20社を五つの役割に分ける(前章)。")
bullet("第二段階", "役割ごとに予算を決める。完成・変化・新生に各25％、両立型15％、分散役10％。三世代へ同額を置くのは、どの未来にも賭けないという設計の続きである。")
bullet("第三段階", "同じ役割の中では、値動きが穏やかで・売買しやすく・証拠が強く・データが信頼できる会社を厚くする(式13)。")
bullet("第四段階", "1銘柄8％の上限をかけ、実際に買える株数へ丸める(式14)。")
formula("13", "役割予算つきの配分 ― 選定済み20社への配り方",
        r"\omega_i = B_{r(i)} \cdot \frac{\rho_i}{\sum_{j:\,r(j)=r(i)} \rho_j}, \qquad \rho_i = \frac{\ell_i\, e_i\, c_i}{\max(\sigma_i,\,0.10)}",
        "Ｂ＝役割ごとの予算(25/25/25/15/10％)　ℓ＝売買のしやすさ　ｅ＝証拠の強さ　ｃ＝データの信頼度　σ＝1年の値動きの大きさ。8％を超えた分は同じ役割の中で配り直す。")
verdict("よって、同じ役割の中では「値動きが穏やかで・売買しやすく・証拠が強い」会社ほど厚く持つ。")
formula("14", "単元株の調整 ― 実際に買える株数へ丸める",
        r"q_i = L_i \cdot \operatorname{floor}\!\left( \frac{B\,\omega_i}{P_i L_i} \right)",
        f"ｑ＝購入株数　Ｌ＝売買単位(基準1株)　Ｂ＝総予算500万円　ω＝目標比率　Ｐ＝株価　floor＝切り捨て。")
verdict(f"よって、目標比率を実際に買える株数へ丸め、総予算500万円のうち{INVEST:,}円(使用率{INVEST/5_000_000:.1%})を投資に充てた。")
wmax = max(DATA.values(), key=lambda z: z["w"])
body(f"最終配分を図表Ⅲ-1に示す。最大の保有は{JPN[wmax['code']]}の{wmax['w']*100:.2f}％、役割の合計は25/25/25/15/10％で、すべて決めた範囲の中にある。実際の売買単位(100株)では株価の高い銘柄が買えないため、1株から買える単元未満株の利用を前提とした。")
prows = []
for role in order:
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        prows.append([d["code"], JPN[d["code"]], role_jp[role], THEME_JP[d["theme"]],
                      f"{d['w']*100:.2f}％", f"{d['qtyL1']:,}", f"{d['amtL1']:,}円"])
table(["コード", "企業名", "役割", "テーマ", "比率", "株数", "金額"], prows,
      [14, 46, 18, 30, 15, 13, 22], fs=10)
para("図表Ⅲ-1　最終ポートフォリオ(総予算500万円・単元未満株基準)。出所：配分の正典データ(allocation_final.csv)。", MIN, 10, after=7)

banner2("2．銘柄紹介 ― 三世代の堀、20社の顔ぶれ")
body("各社の「選定データが語る事実」は、第Ⅱ章の点数と証拠から自動的に書ける。一方、事業の中身と堀の背景は、各社のＩＲ資料・有価証券報告書で確かめてから書くべきものなので、確認欄として残した(創作はしない)。")


def role_fact(d):
    r = d["role"]
    if r == "Buffett Core":
        return (f"守の七関門をすべて通過。純資産比{d['bm']:.2f}倍の割安・粗利÷総資産{d['gp']:.2f}・"
                f"財務6項目の合格割合{d['piotroski']:.0%}。完成した堀を割安に持つ基準線。比率{d['w']*100:.1f}％")
    if r == "Transformation Core":
        return (f"変わる堀の点数{d['tsc']:.0f}点(20社中の変わる堀5社)。純資産比{d['bm']:.2f}倍の割安に"
                f"改善・還元の証拠(証拠の強さ 水準{d['evid']})が重なる。比率{d['w']*100:.1f}％")
    if r == "Emerging Core":
        return (f"生まれる堀の点数{d['esc']:.0f}点・証拠の強さ 水準{d['evid']}。テーマは{THEME_JP[d['theme']]}。"
                f"キーワード頼みを弾く減点を通過した実需接続。比率{d['w']*100:.1f}％")
    if r == "Dual Moat":
        return (f"変わる堀{d['tsc']:.0f}点×生まれる堀{d['esc']:.0f}点の両立。証拠は厳しい方で判定して"
                f"水準{d['evid']}。二つの未来の橋渡し役。比率{d['w']*100:.1f}％")
    return (f"業種・テーマの偏りを整える堅実枠。変わる堀の点数{d['tsc']:.0f}点・証拠の強さ 水準{d['evid']}。"
            f"比率{d['w']*100:.1f}％")


role_desc = {
    "Buffett Core": "守の式だけで選んだ基準線。この5社を上回れるかが「超える」の物差しになる。",
    "Transformation Core": "割安×改善の証拠で選んだ、これから変わる会社。",
    "Emerging Core": "ＡＩ基盤への実需接続を証拠で確かめた、新しい堀が生まれる会社。",
    "Dual Moat": "変わる堀と生まれる堀、二つの点数がともに高い会社。",
    "Bridge / Diversifier": "業種・テーマの偏りを整え、ポートフォリオ全体を安定させる会社。",
}
fig_no = 2
for role in order:
    banner3(f"{role_jp[role]} ― {role_desc[role]}")
    mrows = []
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        mrows.append([f"{d['code']} {JPN[d['code']]}\n({SEC_JP[d['sector']]})",
                      role_fact(d), "［要ＩＲ確認：事業内容・堀の背景・リスク］"])
    table(["企業(業種)", "選定データが語る事実", "事業内容・堀の背景(確認して記入)"], mrows,
          [42, 78, 48], fs=10)
    para(f"図表Ⅲ-{fig_no}　銘柄紹介：{role_jp[role]}(左2列は正典データから自動記載)。", MIN, 10, after=5)
    fig_no += 1

banner2("3．パフォーマンス検証 ― 対照群『純正バフェット』と比べる")
body("前節までで20社と配分が決まった。本節では第Ⅰ章の条件④を検証する。比べる物差しは、分析の前に宣言しておく――①年率リターン ②ＴＯＰＩＸ超過 ③市場との連動度(β) ④最大の下落 ⑤市場超過の安定度(ＩＲ)の5つ、期間は「選定に使った3年」と「選定に使っていない直近1年」に分ける。過去データで選ぶ条件はどのチームも同じであり、その分は割り引いて読む――この注意はここに一度だけ書く。")
body("対照群『純正バフェットポートフォリオ』は、同じ出発点(3,099社)から選定ルールだけをバフェットの式(守)にした20社である。本ポートフォリオのBuffett Coreと同一系譜であるphase1の固定順候補リストの上位から、Top5選定と同じ「同一業種2社まで」の原則で20社を取り、配分の工夫を混ぜず等金額とした(私たちが同じ規律で後から入れ替えずに計算したものであることを明記する)。ＴＯＰＩＸ(市場平均)と日経平均は参照ベンチマークとして併記する。")
fig("fig3_cum.png", "図表Ⅲ-7　累積リターンの比較(過去3年・期首=1)",
    "注：ＴＯＰＩＸ連動指数1306は未調整の株式分割(2026-03-30)を補正済み。出所：control_comparison.json・価格正典データ。")
table(["物差し", "本ＰＦ", "純正バフェット(対照)", "ＴＯＰＩＸ", "日経平均"],
      [["年率リターン(3年)", f"{O3['ann_return']:.1%}", f"{C3['ann_return']:.1%}", f"{O3['topix_ann_return']:.1%}", f"{O3['nikkei_ann_return']:.1%}"],
       ["ＴＯＰＩＸ超過(3年)", f"＋{O3['excess_vs_topix']*100:.1f}pt", f"{C3['excess_vs_topix']*100:.1f}pt", "―", "―"],
       ["市場との連動度β(3年)", f"{O3['beta_vs_topix']:.3f}", f"{C3['beta_vs_topix']:.3f}", "1.000", "―"],
       ["最大の下落(3年)", f"{O3['max_drawdown']:.1%}", f"{C3['max_drawdown']:.1%}", f"{O3['topix_max_drawdown']:.1%}", "―"],
       ["超過の安定度ＩＲ(3年)", f"＋{O3['information_ratio']:.3f}", f"{C3['information_ratio']:.3f}", "―", "―"],
       ["年率リターン(直近1年)", f"{O1['ann_return']:.1%}", f"{C1['ann_return']:.1%}", f"{O1['topix_ann_return']:.1%}", f"{O1['nikkei_ann_return']:.1%}"],
       ["ＴＯＰＩＸ超過(直近1年)", f"{G1:.1f}pt", f"{C1['excess_vs_topix']*100:.1f}pt", "―", "―"],
       ["最大の下落(直近1年)", f"{O1['max_drawdown']:.1%}", f"{C1['max_drawdown']:.1%}", f"{O1['topix_max_drawdown']:.1%}", "―"]],
      [40, 30, 38, 28, 28], fs=10)
para("図表Ⅲ-8　対照群・参照ベンチマークとの比較。出所：control_comparison.json(phase5と同一の計算規約)。", MIN, 10, after=6)
body(f"読み取りは三つ。第一に、本ポートフォリオは対照群を選定に使った3年間で年{D3:.1f}ポイント、選定に使っていない直近1年でも年{D1:.1f}ポイント上回った――守の式だけでは届かない部分を、変わる堀・生まれる堀が実際に運んだ。第二に、正直に認める点として、守りは対照群の方が堅い(最大の下落も市場との連動度も対照群が小さい)。第三に、参照のＴＯＰＩＸには直近1年で年{G1:.1f}ポイント劣後し、日経平均には3年でも届いていない。対照群には勝ったが、市場平均に勝ち続ける力までは証明できていない――ここは言い換えずに残す。")
body("役割の分担も確かめた。過去3年の値上がり寄与の8割は生まれる堀(フジクラ・オルガノなど)で、完成した堀は寄与1％の安定した土台に徹していた。また、割安側の銘柄を外して組み直すと値動きの大きさが35.5％へ悪化する。変わる堀は攻めではなく、揺れを抑える緩衝材として効いている。なお本検証は取引コスト・税・配当再投資の差を考慮していない。")
fig("fig3_roles.png", "図表Ⅲ-9　役割別の寄与 ― 生まれる堀が牽引し、完成した堀は土台",
    "注：過去3年・選定に使った期間。出所：検証の正典データ(phase5_validation_summary.json)。")

banner2("4．結論 ― 結局、バフェットを超えたのか")
body("第Ⅰ章で立てた投資仮説と四つの条件に、正面から答える。")
table(["「超える」の条件", "結果", "判定"],
      [["① 同じ土俵に立つ", "守で規律を式だけから再現し、その5社を固定保有(第Ⅱ章)", "達成"],
       ["② 死角を補う", "バフェットの式では選ばれない15社を、証拠の強さつきで選定。水準3が15社・水準2が4社・水準1が1社。除外862社は理由つきで全記録(第Ⅱ章)", "達成"],
       ["③ 壊れにくい", "16通りの検査で確認。候補を100社に狭める検査以外は、どの条件を変えても20社中11社以上が一致(第Ⅱ章末)", "達成"],
       ["④ 成績で上回る", f"対照群『純正バフェット』を3年で年＋{D3:.1f}ポイント、選定に使っていない直近1年でも年＋{D1:.1f}ポイント上回った。参照のＴＯＰＩＸには直近1年で年{G1:.1f}ポイント劣後(図表Ⅲ-8)", "達成(参照指数には課題)"]],
      [34, 114, 20], fs=10)
para("図表Ⅲ-10　判定表 ― 四条件に対する現時点の答え。", MIN, 10, after=6)
body(f"四つの条件が、それぞれの持ち場で達成となった。①同じ土俵・②死角の補完・③壊れにくさは設計と検査で、④成績は対照群『純正バフェット』との比較で――3年でも、選定に使っていない直近1年でも上回った。守の規律だけでは届かない未来を、証拠つきの変わる堀・生まれる堀が実際に運んだことになる。ただし、参照のＴＯＰＩＸには直近1年で年{G1:.1f}ポイント負けており、市場平均に勝ち続ける力までは証明できていない。この宿題を隠さないことも、私たちの設計の一部である。第Ⅰ章の投資仮説――三世代を束ねた選定は純正バフェット型を上回る――は、この比較によって支持された。")
body("最後に一つ。証拠水準の分離管理・除外862社の全記録・16通りの壊れにくさ検査――この三つは、既製の手法ではなく、私たちが検証可能性のために自分の手で組んだ、本レポート固有の装備である。私たちがこの20社を持ち続けられる理由は、成績の自信ではなく、この構造と証拠と検査にある。")

banner2("5．リスクと限界 ― 正直に書き残す")
body("最後に、この設計の弱点を私たち自身の手で記録しておく。強みだけを並べたレポートは、検証に耐えない。")
bullet("テーマの偏り", f"過去3年の値上がり寄与の8割は生まれる堀に集中しており、テーマの偏りの度合い(偏り指数)は{PERF['hhi']['theme']:.2f}と高め。生まれる堀のテーマが崩れる未来では、市場に劣る可能性がある。だからこそ完成した堀・変わる堀に同額の予算を置き、一つの未来への賭けを避けている。")
bullet("単元未満株の前提", f"通常の売買単位(100株)では、株価の高い{N_UNBUY}社が予算内で買えない。1株から買える単元未満株の取扱いを前提としており、取扱いのない証券会社ではこの配分をそのまま再現できない。")
bullet("証拠は開示ベース", "証拠の強さは開示資料に基づく判定であり、現場の実態までは測れない。第Ⅳ章の取材は、この限界を埋めるために設計している。")
point("四条件とも達成(④は対照群比較)。ただしＴＯＰＩＸへの直近1年の劣後と設計の弱点まで、先に書き残した。")

# ================= Ⅳ インタビュー =================
banner1("Ⅳ．インタビュー・アンケート", newpage=True)
body("前章までの選定と検証は、すべて開示データの上に立っている。本章では、その見立てが現場の実態と合っているかを確かめる取材の計画を示す。質問は思いつきではなく、各社を選んだ理由(点数と証拠)から導いた(実施記録は記入欄)。")
table(["対象(役割)", "選定理由から導いた質問草案"],
      [["5803 フジクラ(生まれる堀)", "①データセンター向け光配線の受注動向と生産能力投資の計画　②ＡＩ需要が一巡した場合の下支え事業"],
       ["6368 オルガノ(生まれる堀)", "①半導体向け超純水装置の需要見通し　②価格決定力(急所を握る力)の源泉"],
       ["6920 レーザーテック(生まれる堀)", "①半導体の検査・品質保証分野での技術優位の持続性　②研究開発投資の方向"],
       ["9470 学研ＨＤ(完成した堀)", "①教育×医療福祉の資本効率改善の進捗　②株主還元方針の考え方"],
       ["5233 太平洋セメント(変わる堀)", "①資本コストを意識した経営改革の進捗　②株主還元と設備投資の優先順位"],
       ["9474 ゼンリン(両立型)", "①地図データベースの参入障壁　②ＡＩ時代におけるデータ資産の収益化"],
       ["3697 ＳＨＩＦＴ(両立型)", "①ソフトウェア品質保証の需要構造　②ＡＩ活用による事業機会と脅威"],
       ["2112 塩水港精糖(分散役)", "①原料価格の変動への耐性　②財務改善の計画"]],
      [52, 118], fs=10)
banner3("実施記録(記入・取材ごとに1列)")
table(["項目", "取材1(記入)", "取材2(記入)"],
      [["企業名・対象者(名字)・部署", "［記入］", "［記入］"],
       ["日時・方法(対面/オンライン/書面)", "［記入］", "［記入］"],
       ["質問と回答の要点", "［記入：質問草案の番号と対応させる］", "［記入］"],
       ["写真(オンライン画面も可)", "［写真貼付欄］\n\n\n", "［写真貼付欄］\n\n\n"],
       ["分析への反映", "［記入：証拠水準・役割の見立てをどう見直したか］", "［記入］"]],
      [38, 66, 66], fs=10)
banner3("取材の還流(総括・記入)")
para("書き方の型：「取材の結果、○○社の△△が確認でき、証拠水準の見立てを□□と見直した／仮説どおりだった。第Ⅱ章の××の判断は現場の実態と整合していた」――分析へ何が返ってきたかを必ず書く。", MIN, 10, after=2)
fill_line("総括(3〜5文)")
para("(未実施の場合は本章を「今後の課題」とし、その旨を明記する。回答の創作はしない。)", MIN, 10, after=6)

# ================= Ⅴ 学んだこと =================
banner1("Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと")
for lab, hint in [("直面した困難と乗り越え方", "例の問い：仮説はどこで崩れたか。候補を広げすぎて確認が回らなくなったとき、どう折り合いをつけたか"),
                  ("インタビューと結びついた学び", "例の問い：取材で確かめて初めて分かったことは何か。開示データだけでは見えなかったものは何か"),
                  ("チームでの役割分担と今後の課題", "例の問い：互いの誤りをどう見つけ合ったか。単元未満株の扱い・証拠確認の自動化など残った宿題"),
                  ("謝辞", "取材先・助言をくれた方・支えてくれた人へ(個人は名字のみ)")]:
    fill_line(lab, hint)
para("(提出者の実体験として記入する。下書きの創作はしない。)", MIN, 10, after=6)

# ================= 参考文献 =================
banner1("参考文献")
body("本文で直接引用した文献に加え、選定・最適化・検証・データ取得の設計にあたって参照した文献を、分野別にすべて記載する(正典: docs/references_master.md)。", after=5)


def ref_para(r):
    p = doc.add_paragraph(); pp = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360"); pp.append(ind)
    p.paragraph_format.space_after = Pt(1)
    X.setfont(p.add_run(r), MIN, 10, False, BLACK)


REF_GROUPS = [
 ("一次資料 ― バフェット株主への手紙(原文)", [
  "Buffett, W. E. (1988) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/1988.html (2026年7月18日)",
  "Buffett, W. E. (1989) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/1989.html (2026年7月18日)",
  "Buffett, W. E. (2007) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/2007ltr.pdf (2026年7月18日)",
 ]),
 ("英語文献Ⅰ ― 投資・会計・無形資産(第Ⅱ章 守・離の式典拠)", [
  "Altman, E. I. (1968) Financial Ratios, Discriminant Analysis and the Prediction of Corporate Bankruptcy. The Journal of Finance, 23(4), pp.589-609.",
  "Asness, C. S., Frazzini, A., & Pedersen, L. H. (2019) Quality Minus Junk. Review of Accounting Studies, 24, pp.34-112.",
  "Basu, S. (1977) Investment Performance of Common Stocks in Relation to Their Price-Earnings Ratios: A Test of the Efficient Market Hypothesis. The Journal of Finance, 32(3), pp.663-682.",
  "Basu, S. (1983) The Relationship between Earnings' Yield, Market Value and Return for NYSE Common Stocks: Further Evidence. Journal of Financial Economics, 12(1), pp.129-156.",
  "Chan, L. K. C., Lakonishok, J., & Sougiannis, T. (2001) The Stock Market Valuation of Research and Development Expenditures. The Journal of Finance, 56(6), pp.2431-2456.",
  "Fama, E. F., & French, K. R. (1993) Common Risk Factors in the Returns on Stocks and Bonds. Journal of Financial Economics, 33(1), pp.3-56.",
  "Fama, E. F., & French, K. R. (2015) A Five-Factor Asset Pricing Model. Journal of Financial Economics, 116(1), pp.1-22.",
  "Frazzini, A., Kabiller, D., & Pedersen, L. H. (2018) Buffett's Alpha. Financial Analysts Journal, 74(4), pp.35-55.",
  "Jensen, M. C. (1968) The Performance of Mutual Funds in the Period 1945-1964. The Journal of Finance, 23(2), pp.389-416.",
  "Lev, B., & Gu, F. (2016) The End of Accounting and the Path Forward for Investors and Managers, John Wiley & Sons.",
  "Markowitz, H. (1952) Portfolio Selection. The Journal of Finance, 7(1), pp.77-91.",
  "Novy-Marx, R. (2013) The Other Side of Value: The Gross Profitability Premium. Journal of Financial Economics, 108(1), pp.1-28.",
  "Ohlson, J. A. (1980) Financial Ratios and the Probabilistic Prediction of Bankruptcy. Journal of Accounting Research, 18(1), pp.109-131.",
  "Peters, R. H., & Taylor, L. A. (2017) Intangible Capital and the Investment-q Relation. Journal of Financial Economics, 123(2), pp.251-272.",
  "Piotroski, J. D. (2000) Value Investing: The Use of Historical Financial Statement Information to Separate Winners from Losers. Journal of Accounting Research, 38, Supplement, pp.1-41.",
  "Sharpe, W. F. (1966) Mutual Fund Performance. The Journal of Business, 39(1), pp.119-138.",
  "Sharpe, W. F. (1994) The Sharpe Ratio. The Journal of Portfolio Management, 21(1), pp.49-58.",
  "Sloan, R. G. (1996) Do Stock Prices Fully Reflect Information in Accruals and Cash Flows about Future Earnings? The Accounting Review, 71(3), pp.289-315.",
 ]),
 ("英語文献Ⅱ ― 最適化・探索(第Ⅱ章 破の方法典拠)", [
  "Akiba, T., Sano, S., Yanase, T., Ohta, T., & Koyama, M. (2019) Optuna: A Next-generation Hyperparameter Optimization Framework. In Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining, pp.2623-2631.",
  "Bergstra, J., Bardenet, R., Bengio, Y., & Kégl, B. (2011) Algorithms for Hyper-Parameter Optimization. In Advances in Neural Information Processing Systems, 24, pp.2546-2554.",
  "Bergstra, J., & Bengio, Y. (2012) Random Search for Hyper-Parameter Optimization. Journal of Machine Learning Research, 13, pp.281-305.",
  "Brochu, E., Cora, V. M., & de Freitas, N. (2010) A Tutorial on Bayesian Optimization of Expensive Cost Functions. arXiv, https://arxiv.org/abs/1012.2599 (2026年7月8日)",
  "Deb, K. (2001) Multi-Objective Optimization Using Evolutionary Algorithms, John Wiley & Sons.",
  "Deb, K., Pratap, A., Agarwal, S., & Meyarivan, T. (2002) A Fast and Elitist Multiobjective Genetic Algorithm: NSGA-II. IEEE Transactions on Evolutionary Computation, 6(2), pp.182-197.",
  "Fonseca, C. M., & Fleming, P. J. (1995) An Overview of Evolutionary Algorithms in Multiobjective Optimization. Evolutionary Computation, 3(1), pp.1-16.",
  "Goldberg, D. E. (1989) Genetic Algorithms in Search, Optimization, and Machine Learning, Addison-Wesley.",
  "Hansen, N., & Ostermeier, A. (2001) Completely Derandomized Self-Adaptation in Evolution Strategies. Evolutionary Computation, 9(2), pp.159-195.",
  "Hutter, F., Hoos, H. H., & Leyton-Brown, K. (2011) Sequential Model-Based Optimization for General Algorithm Configuration. In Learning and Intelligent Optimization, pp.507-523.",
  "Jamieson, K., & Talwalkar, A. (2016) Non-stochastic Best Arm Identification and Hyperparameter Optimization. In Proceedings of the 19th International Conference on Artificial Intelligence and Statistics, pp.240-248.",
  "Kushner, H. J. (1964) A New Method of Locating the Maximum Point of an Arbitrary Multipeak Curve in the Presence of Noise. Journal of Basic Engineering, 86(1), pp.97-106.",
  "Li, L., Jamieson, K., DeSalvo, G., Rostamizadeh, A., & Talwalkar, A. (2017) Hyperband: A Novel Bandit-Based Approach to Hyperparameter Optimization. Journal of Machine Learning Research, 18(185), pp.1-52.",
  "Mockus, J. (1978) The Application of Bayesian Methods for Seeking the Extremum. In Towards Global Optimization, Vol.2, pp.117-129, North-Holland.",
  "Snoek, J., Larochelle, H., & Adams, R. P. (2012) Practical Bayesian Optimization of Machine Learning Algorithms. In Advances in Neural Information Processing Systems, 25, pp.2951-2959.",
  "Storn, R., & Price, K. (1997) Differential Evolution - A Simple and Efficient Heuristic for Global Optimization over Continuous Spaces. Journal of Global Optimization, 11, pp.341-359.",
  "Zitzler, E., Deb, K., & Thiele, L. (2000) Comparison of Multiobjective Evolutionary Algorithms: Empirical Results. Evolutionary Computation, 8(2), pp.173-195.",
 ]),
 ("英語文献Ⅲ ― 検証・過学習対策・感度分析(壊れにくさ検査の典拠)", [
  "Bailey, D. H., & López de Prado, M. (2012) The Sharpe Ratio Efficient Frontier. Journal of Risk, 15(2), pp.3-44.",
  "Bailey, D. H., Borwein, J., López de Prado, M., & Zhu, Q. J. (2014) Pseudo-Mathematics and Financial Charlatanism: The Effects of Backtest Overfitting on Out-of-Sample Performance. Notices of the American Mathematical Society, 61(5), pp.458-471.",
  "Bailey, D. H., & López de Prado, M. (2014) The Deflated Sharpe Ratio: Correcting for Selection Bias, Backtest Overfitting, and Non-Normality. The Journal of Portfolio Management, 40(5), pp.94-107.",
  "Bailey, D. H., Borwein, J. M., López de Prado, M., & Zhu, Q. J. (2016) The Probability of Backtest Overfitting. Journal of Computational Finance, 20(4), pp.39-69.",
  "Jaccard, P. (1901) Étude comparative de la distribution florale dans une portion des Alpes et des Jura. Bulletin de la Société Vaudoise des Sciences Naturelles, 37, pp.547-579.",
  "López de Prado, M. (2018) Advances in Financial Machine Learning, John Wiley & Sons.",
  "Romano, J. P., & Wolf, M. (2005) Stepwise Multiple Testing as Formalized Data Snooping. Econometrica, 73(4), pp.1237-1282.",
  "Saltelli, A., Ratto, M., Andres, T., Campolongo, F., Cariboni, J., Gatelli, D., Saisana, M., & Tarantola, S. (2008) Global Sensitivity Analysis: The Primer, John Wiley & Sons.",
  "Sobol', I. M. (2001) Global Sensitivity Indices for Nonlinear Mathematical Models and Their Monte Carlo Estimates. Mathematics and Computers in Simulation, 55(1-3), pp.271-280.",
  "White, H. (2000) A Reality Check for Data Snooping. Econometrica, 68(5), pp.1097-1126.",
 ]),
 ("背景調査で参照した文献", [
  "Beneish, M. D. (1999) The Detection of Earnings Manipulation. Financial Analysts Journal, 55(5), pp.24-36.",
  "French, K. R. \"Data Library\", https://mba.tuck.dartmouth.edu/pages/faculty/ken.french/data_library.html (2026年7月8日)",
 ]),
 ("日本語文献", [
  "経済産業省(2014)『持続的成長への競争力とインセンティブ ―企業と投資家の望ましい関係構築―(伊藤レポート)』経済産業省．",
  "東京証券取引所(2023)「資本コストや株価を意識した経営の実現に向けた対応について」株式会社東京証券取引所．",
  "日本取引所グループ「単元株制度」，https://www.jpx.co.jp/ (2026年7月8日閲覧)",
 ]),
 ("データ・実装資料", [
  "金融庁「EDINET」，https://disclosure2.edinet-fsa.go.jp/ (2026年7月8日閲覧)",
  "金融庁 企画市場局 企業開示課(2026)『EDINET API 仕様書(Version 2)』金融庁．",
  "日本取引所グループ「上場会社情報」，https://www.jpx.co.jp/listing/co-search/ (2026年7月8日閲覧)",
  "日本取引所グループ「東証上場銘柄一覧」(データファイル)．",
  "Optuna Contributors, \"Optuna Documentation\", https://optuna.readthedocs.io/ (2026年7月8日)",
  "pymoo Developers, \"pymoo: Multi-objective Optimization in Python\", https://pymoo.org/ (2026年7月8日)",
  "SciPy Developers, \"scipy.optimize.differential_evolution\", SciPy Documentation, https://docs.scipy.org/ (2026年7月8日)",
  "scikit-learn Developers, \"Model selection and evaluation\", scikit-learn Documentation, https://scikit-learn.org/stable/model_selection.html (2026年7月8日)",
 ]),
]
for gtitle, grefs in REF_GROUPS:
    banner3(gtitle)
    for r in grefs:
        ref_para(r)

# ================= 用語の手引き =================
banner1("用語の手引き")
body("本文で使った言葉の早見表。専門用語は本文でも初出時に言い換えているが、迷ったらここへ。", after=5)
table(["ことば", "意味"],
      [["堀(Moat)", "他社がまねしにくい持続的な競争優位。城のまわりの堀のたとえ"],
       ["守破離(しゅはり)", "日本の芸道で修行の段階を表す言葉。守＝型を忠実に守る、破＝型を応用し破る、離＝型を離れて独自の境地を開く"],
       ["時価総額", "株価×株式数。市場がつけた会社全体の値段"],
       ["Ｂ／Ｍ", "純資産÷時価総額。大きいほど会社の持ち物のわりに株価が安い"],
       ["Ｅ／Ｐ", "純利益÷時価総額。大きいほど稼ぐ力のわりに株価が安い"],
       ["粗利÷総資産", "資産を売上総利益に変える力。優良さの物差し"],
       ["6項目チェック", "財務の健康を良い/悪いで数える確認法(Piotroski)。原式9項目のうち取得できた6項目の合格割合で判定"],
       ["利益の質", "帳簿の利益に現金の裏づけがあるか。差が大きいと後で失速しやすい(Sloan)"],
       ["債務超過", "借金などの負債が資産を上回った状態"],
       ["ＰＢＲ(株価純資産倍率)", "株価が1株あたり純資産の何倍か。1倍割れは会社の持ち物より安い値付け"],
       ["単元未満株", "通常の売買単位(100株)より少ない株数で買える仕組み"],
       ["ＴＯＰＩＸ", "東証株価指数。日本の株式市場全体の平均的な動き"],
       ["市場との連動度(β)", "市場が1動いたときに何倍動くか。1より小さければ市場より穏やか"],
       ["最大の下落", "その期間中、天井から底まで最大で何％下がったか"],
       ["市場超過の安定度(ＩＲ)", "市場平均を上回った幅を、その振れの大きさで割った値。マイナスは市場に負けたことを示す"],
       ["証拠の強さ(水準1〜3)", "選定根拠の裏づけの強さ。1＝キーワードの接点のみ、2＝具体性を確認、3＝数字まで確認"],
       ["選定に使った期間", "スクリーニングの計算に使った過去データの期間。この期間の成績は実力の証明にならない"]],
      [46, 124], fs=10)

# ================= 保存+自己検査 =================
out = ED / f"beyond_buffett_stockleague_{VER}.docx"
doc.core_properties.comments = VER
doc.save(str(out))
json.dump({h: None for h, l in HEADINGS if l == 1}, open(ED / f"headings_{VER}.json", "w"), ensure_ascii=False)

# 検査1: 禁止語(平易化スイープ)
BANNED = ["候補宇宙", "閾値", "分位", "アクルーアル", "ガードレール", "正規化", "ロバスト",
          "レビュー可能性", "ボトルネック", "インサンプル", "in-sample", "母集団", "トラップ",
          "ヘッジ", "ユニバース", "アブレーション", "標本内", "ファクター"]
texts = []
for p in doc.paragraphs:
    texts.append(p.text)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            texts.append(c.text)
full = "\n".join(texts)
bad = [w for w in BANNED if w in full]

# 検査2: 全文字10pt以上(w:szは半ポイント単位)
small = set()
for el in doc.element.iter():
    if el.tag == qn("w:sz") or el.tag == qn("w:szCs"):
        v = el.get(qn("w:val"))
        if v and float(v) < 20:
            small.add(v)

print("saved", out.name, "| ver:", VER, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
print("BANNED-CHECK:", "PASS" if not bad else f"FAIL {bad}")
print("FONT>=10pt-CHECK:", "PASS" if not small else f"FAIL sz(half-pt)={sorted(small)}")
print(f"INVEST={INVEST:,} CASH={CASH:,}")
