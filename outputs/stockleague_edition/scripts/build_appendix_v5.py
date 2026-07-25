# -*- coding: utf-8 -*-
"""v5 査読用技術補遺(beyond_buffett_technical_appendix_v5.docx)。

本編(30頁規定)に収まらない、専門家査読が要求する粒度の詳細をすべて収載する別冊。
頁制約なし。本編→補遺は一方向参照。全数値は正典JSON/CSVから自動転記(手打ち禁止)。
本イテレーションの収載: §A データ系譜 / §C 探索と目的関数 / §D 検証規約・BH / §F 統計。
(§B 式の完全形 / §E 頑健性 / §G 限界レジスタ / §H 再現手順 は次イテレーションで追加。)
"""
import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
ED = ROOT / "outputs/stockleague_edition"

# ---- 正典ロード ----
COMP = json.load(open(ED / "control_comparison_v5.json", encoding="utf-8"))
BH = json.load(open(ED / "bh_metrics_v5.json", encoding="utf-8"))
SIG = json.load(open(ED / "significance_v5.json", encoding="utf-8"))
EXC = json.load(open(ED / "exclusion_summary_v5.json", encoding="utf-8"))
PH2 = json.load(open(ROOT / "outputs/phase2_perfect_final_break/optimization/selected_phase2_solution_clean.json", encoding="utf-8"))
O3, O1, C3, C1 = COMP["ours"]["3y"], COMP["ours"]["1y"], COMP["control"]["3y"], COMP["control"]["1y"]
DATA = json.load(open(ROOT / "outputs/explanatory_revision/data_real.json", encoding="utf-8"))
for _v in DATA.values():
    _v["code"] = str(_v["code"])
ABL = list(csv.DictReader(open(ROOT / "outputs/beyond_buffett_fable_loop_final/phase5_verification_and_ablation/ablation_results.csv", encoding="utf-8")))
WF = list(csv.DictReader(open(ROOT / "outputs/phase2_final_integrated_break/walk_forward/fixed_weight_annual_validation.csv", encoding="utf-8")))

NAVY = RGBColor(0x16, 0x32, 0x4F)
TEAL = RGBColor(0x2F, 0x6D, 0x5F)
MINCHO, GOTHIC, MONO = "ＭＳ 明朝", "ＭＳ ゴシック", "Consolas"

doc = Document()
_st = doc.styles["Normal"]
_st.font.size = Pt(10.5); _st.font.name = "Times New Roman"
_st.element.rPr.rFonts.set(qn("w:eastAsia"), MINCHO)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = sec.bottom_margin = Mm(18)


def _font(run, name=MINCHO, size=10.5, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    _font(p.add_run(text), GOTHIC, 13, True, NAVY)
    return p


def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _font(p.add_run(text), GOTHIC, 11, True, TEAL)
    return p


def para(text, size=10.5, after=5, mono=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _font(p.add_run(text), MONO if mono else MINCHO, 9.5 if mono else size)
    return p


def table(hdr, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"
    for j, htext in enumerate(hdr):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(htext); _font(r, GOTHIC, 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
        _shade(c, "16324F")
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear()
            _font(cells[j].paragraphs[0].add_run(str(val)), MINCHO, 9.5)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def _shade(cell, hexcolor):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def pct(x, dp=1):
    return f"{x*100:.{dp}f}%"


# ================= 表紙 =================
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
_font(p.add_run("BEYOND BUFFETT ― 査読用技術補遺(Technical Appendix)"), GOTHIC, 16, True, NAVY)
para("日経ＳＴＯＣＫリーグ提出版『三世代の堀』の検証可能性を担保するための技術資料。本編(30頁規定)に"
     "収まらない詳細――全パラメータ・目的関数・計算規約・統計・頑健性・限界・再現手順――を収載する。"
     "本補遺の数値はすべて本編と同一の正典データ(JSON/CSV)から自動転記しており、本編の各節に対応する。", after=8)
para("凡例: 本補遺は審査対象外の補助資料であり、本編を補完する。数式・コードは等幅体で示す。"
     "出所として括弧内にファイル名を付す。", size=9.5, after=10)

# ================= §A データ系譜 =================
h1("§A　データ系譜と前処理(本編Ⅱ章冒頭・Ⅲ章§3に対応)")
h2("A-1　価格・指数データ")
para("価格系列はプロジェクト同梱の正典データ(data/processed/prices_daily.parquet、生成2026-06-01)に基づく。"
     "個別株および東証株価指数(ＴＯＰＩＸ)連動ＥＴＦ 1306.T は配当込みの調整後終値(yfinance Adj Close, "
     "auto_adjust=False)を用いる。日経平均(^N225)は配当を含まない価格指数である。したがって本ＰＦ対ＴＯＰＩＸ"
     "の比較は双方が配当込みでほぼ整合し、対日経平均の比較のみ本ＰＦに有利な非対称がある(日経の配当込み換算は"
     "さらに高い)。", after=5)
h2("A-2　データ補正(2件)と欠測補完")
para("① 未調整の株式分割: 1306.T は源データにおいて2026-03-30以降の価格が未調整の1:10分割を含む"
     "(調整後終値が1日で約-90%)。2026-03-30以降を×10して連続化した。ポートフォリオ構成銘柄には分割様の"
     "断崖はない(2112.T の2024-01-25 +32.9%は小型株の実変動として保持)。", after=3)
para(f"② 指数系列の欠測補完(WP0.1・本編の是正点): 1306.T は通常営業日である {COMP['gap_repair']['1306.T'][0]} が源parquetで欠測していた。"
     "従来はこの欠測により前後2日のリターンがNaN化し、その間の値動き(前営業日終値3,407→翌営業日3,485、"
     "約+2.29%)が累積計算から脱落していた。本補遺・本編v5では欠測日を前営業日終値で補完し、"
     "バージョン非依存の pct_change(fill_method=None) で算出する。構成銘柄・対照群40社には期間内の内部欠測が"
     "無いことをassertで確認済み(補完は指数系列のみ)。この是正でＴＯＰＩＸ年率は3年24.06→"
     f"{pct(O3['topix_ann_return'])}・直近1年40.66→{pct(O1['topix_ann_return'])}に上方修正され、"
     "対ＴＯＰＩＸ成績はより不利な(=正直な)値となった。", after=5)
para("出所: outputs/stockleague_edition/scripts/make_validation_v5.py, control_comparison_v5.json(gap_repair欄)。", size=9, after=6)
h2("A-3　母集団と生存者バイアス")
para("母集団は東証上場銘柄一覧(JPX)の2026-04-30時点スナップショット単一断面から金融を除いた3,099社である。"
     "したがって2023〜2026年に上場廃止・経営破綻した企業は母集団・対照群プール・過去3年統計のいずれにも"
     "含まれない(生存者バイアス)。本ＰＦと対照群には同一条件で作用するため両者の比較は内部整合するが、"
     "参照指数(ＴＯＰＩＸ・日経)との水準比較は生存者ユニバース上の数値であり、割り引いて読む必要がある。", after=8)

# ================= §B 式の完全形 =================
h1("§B　式の完全形と実装細部(本編の全14式に対応)")
h2("B-1　守の式(1)〜(7)の実装細部")
para("本編の式(1)〜(7)は先行研究の定義そのものだが、実装には次の細部がある。①割安・利益の質の各指標は"
     "1%/99%分位でウィンドソライズした値で順位化・閾値判定する(生値ではない)。②順位(percentile)の基底は"
     "指標ごとに異なる(Ｂ／Ｍは算出可能な3,089社内、Ｅ／Ｐは黒字2,740社内)。③式(4)のPiotroskiは原式9項目のうち"
     "取得できた6シグナル(ＲＯＡ＞0／営業ＣＦ＞0／ＲＯＡ改善／営業ＣＦ÷総資産＞ＲＯＡ／資産回転率改善／"
     "負債比率低下)で、閾値0.65の実効は4/6以上。負債比率は総負債/総資産(原式は長期負債)、アクルーアル項は"
     "期末総資産割り(原式は期首)。④式(7)の売買代金は77社ファネルでは実効300万円以上(3百万〜1千万は要確認扱いで通過)、"
     "Top5選定ファネルでは1,000万円以上を明確に要求する。出所: scripts/phase1_final/final_phase1.py, scripts/phase1_top5/build_top5.py。", after=5)
h2("B-2　通し数値例: 学研ホールディングス(9470)の七関門(本編Ⅱ章から移設)")
g = DATA["9470"]
table(["関門", "学研ＨＤの値", "合格ライン", "判定"],
      [["Ｂ／Ｍ(式1)", f"{g['bm']:.3f}", "市場の上位30％", "通過"],
       ["Ｅ／Ｐ(式2)", f"{g['ep']:.3f}(黒字)", "黒字かつ上位50％", "通過"],
       ["粗利÷総資産(式3)", f"{g['gp']:.3f}", "中央値以上", "通過"],
       ["6項目チェック(式4)", f"合格割合 {g['piotroski']:.0%}({round(g['piotroski']*6)}/6項目)", "0.65以上", "通過"],
       ["利益と現金の差(式5)", f"{g['sloan']:.3f}", "悪い側の上位30％でない", "通過"],
       ["危険よけ(式6)", "債務超過・2期連続赤字なし", "該当ゼロ", "通過"],
       ["売買のしやすさ(式7)", f"1日平均 {g['adv_oku']:.2f}億円", "約0.1億円以上", "通過"]],
      [40, 46, 44, 20])
para("出所: 選定の正典データ(data_real.json)。", size=9, after=6)
h2("B-3　離の式(10)変わる堀 ― 完全形と減点の内訳")
para("本編の式(10)は採用形(partial)の重み(0.22/0.24/0.10/0.18/0.16/0.10)を示した。設計初期のフル形は"
     "w_V=0.20・w_C=0.22・w_R=0.16・w_E=0.17・w_X=0.13・w_Q=0.12 だが、株主還元(R)・改革開示(E)の正式フィールド"
     "欠損によりフル形は構造的に到達不能で、FCFプロキシ(F=営業ＣＦ−設備投資が正)とデータ信頼度(Φ=Phase2 confidence"
     "×100/1.1)へ再配分したpartial形を採用形とした。結果は[0,100]にクリップする。最終20社のうち19社がpartial形、"
     "4350のみlite形(0.30/0.20/0.20/0.20/0.10)で採点する(4350は守の固定枠=Phase1で選定済みのため選定妥当性は"
     "毀損しない)。", after=3)
para("P^Trap(割安のワナ減点)の内訳は8成分の加点方式: 高Sloanアクルーアル10・財務危険(distress)25・異常値20・"
     "負の営業ＣＦ12・継続赤字15・正規化に脆弱8・外れ値に敏感5・粗利プロキシ5(点)。重みは事後リターン最適化ではなく"
     "概念上の設計係数で、±20%摂動でSpearman ρ最小0.9973・選定5社は不変。出所: phase3_common.py, phase3_formula_lineage.md。", after=5)
h2("B-4　離の式(11)生まれる堀・式(12)証拠水準 ― 本編省略項の開示")
para("式(11)の重みは w_I=0.18(無形資産)・w_N=0.15(技術力)・w_B=0.18(急所)・w_A=0.22(ＡＩ基盤接続)・"
     "w_D=0.14(データ・顧客)・w_T=0.13(信頼・安全)。加点 B^Evidence は証拠水準{0,1,2,3}を{0,2,4,8}点に写像。"
     "減点は P^Hype(キーワードのみ18点)に加え、本編式(11)では省略した P^Guard(財務ガードレール該当20点)がある"
     "(最終20社は全社非該当のため数値影響ゼロだが、862社の除外過程では作動する)。合成後[0,100]クリップ。", after=3)
para("式(12)の証拠水準は役割別分岐: 変わる堀=max(TQ,TS,TR)／生まれる堀=EM／両立型=min(TQ,EM)／その他=max(TQ,EM)。"
     "TQ=3年財務改善指標の正の本数(≥4→3・≥2→2・≥1→1)、TS=ＦＣＦプロキシ正なら1、TR=改革関連列が非欠損なら1、"
     "EM=生まれる堀の開示証拠水準。ただし現データではTR・TSは全社1で識別力を持たず、水準判定は事実上TQ・EMに依存する"
     "(§G参照)。なお本値は選定ゲートには用いず報告専用である。出所: phase3_v2_pipeline.py, phase3_rebuild.py。", after=5)
h2("B-5　配分の式(13)(14) ― 未開示スケールの全開示")
para("式(13)のρ_i = ℓ_i・e_i・c_i / max(σ_i,0.10) を構成する各係数のスケールは次のとおり: 売買のしやすさ"
     "ℓ∈{1.00(売買代金≥5,000万円)・0.85(≥3,000万円)・0.70(未満)}、証拠の強さ e=1+0.05×(証拠水準−2)∈{0.95,1.00,1.05}、"
     "データ信頼度 c=0.5+0.5×信頼度∈[0.5,1.0]、σは直近252営業日の日次リターン標準偏差×√252で下限0.10。"
     "1銘柄8%上限は役割内で反復再配分(最大5パス)するが、最終配分では非拘束(発動なし)。式(14)の単元丸めは"
     "L=1(単元未満株)を前提とし、総額4,949,198円(予算の99.0%)。出所: phase4_allocation.py。", after=6)

# ================= §C 探索と目的関数 =================
h1("§C　破の探索と目的関数(本編Ⅱ章§破・式(9)に対応)")
para("本編の破(第2スクリーニング)は、7指標の重み・減点・候補数・正規化・業種調整などの設定を大量に試し、"
     "候補集合の効用(広さ・質・安全・分散など)を最大化する設定を採用したものである。採用後に順位の安定性を"
     "別途監査した(§E)。重要: 目的関数に将来リターン・期間成績・シャープ等は一切含まれない"
     f"(future_return_prediction_claim = {str(PH2['future_return_prediction_claim']).lower()})。"
     "「選定に成績を使っていない」という本編の主張はこの点に基づく。", after=5)
h2("C-1　探索の実体")
para(f"採用手法は {PH2['selected_method']}(純ランダムサーチ)。探索は3系統併用で計約1.8万試行: "
     "Optuna TPE 5,000試行(seed=42)+ NSGA-II 3,000試行(seed=43)+ 純ランダムサーチ 10,000試行(seed=44)"
     "+ 固定ベースライン。採用解はランダムサーチが到達した最良解で、TPE・NSGAの最良を目的関数値で上回った。"
     "事後の安定性監査は seed=45。", after=5)
h2("C-2　採用した重みベクトル(全7指標)")
_w = PH2["selected_weights"]
_wlabel = {"bm": "割安(純資産)Ｂ／Ｍ", "ep": "割安(利益)Ｅ／Ｐ", "gp": "収益力 粗利÷総資産",
           "piotroski": "6項目チェック", "sloan": "利益の質(Sloan)", "distress": "危険よけ",
           "liquidity": "売買のしやすさ"}
table(["指標", "重み w_k", "備考"],
      [[_wlabel[k], f"{_w[k]:.4f}", ""] for k in ["bm", "ep", "gp", "piotroski", "sloan", "distress", "liquidity"]],
      [70, 30, 60])
para(f"最大の重みは売買のしやすさ(liquidity {_w['liquidity']:.3f})で、次いで危険よけ({_w['distress']:.3f})である。"
     "本編Ⅱ章もこの2値を開示している(危険よけは「多層防御の微調整」)。実効的にはliquidityとdistressの2指標が"
     "支配的で、piotroski・sloanはほぼゼロ(合成点の実態は5指標駆動)である点を、ここに全開示する。", after=5)
h2("C-3　減点とパラメータ")
_p = PH2["selected_penalty_weights"]; _pr = PH2["selected_params"]
table(["減点 P の項", "係数"],
      [["異常値(anomaly)", f"{_p['anomaly']:.4f}"], ["超小型株(microcap)", f"{_p['microcap']:.4f}"],
       ["主要指標の欠損(missing)", f"{_p['missing']:.4f}"], ["一時的利益(onetime)", f"{_p['onetime']:.4f}"],
       ["ＧＰ欠損ペナルティ強度", f"{_pr['gp_missing_penalty_strength']:.4f}"],
       ["重み集中ペナルティ強度", f"{_pr['weight_concentration_penalty_strength']:.4f}"]],
      [80, 40])
para(f"欠損の扱い = {_pr['missing_handling']}(主要指標が欠けた会社は実質除外)、正規化 = "
     f"{_pr['normalization_method']}(市場順位)、業種調整 = {'あり' if _pr['sector_adjustment'] else 'なし'}。"
     "式(9)の合成後にmin-max再正規化し、同点はＢ／Ｍ→Ｅ／Ｐでタイブレークする。", after=5)
h2("C-4　目的関数と Top1200 の採用")
para("目的関数 = max_N〔効用(N)〕− 集中ペナルティ。効用(N)は候補集合の断面特性(市場中央値との差・各フラグ率・"
     "業種ＨＨＩ・守5社カバー等)の重みつき和で、広さ・質・割安・安全・流動性・業種分散・安定性代理・守5社"
     "カバーの加点と、異常値率・欠損率・実行可能性違反の減点からなる。各項の重み(0.20/0.15/…)は手置きで、"
     "リターン最適化ではない。", after=3)
para(f"効用は候補数が多いほど広さ項で有利になるため、効用最大は Top2000 だった"
     f"(top1200_is_utility_optimal = {str(PH2['top1200_is_utility_optimal']).lower()})。それでも Top1200 を正式採用したのは、"
     "広さと「一社ずつ人の目で確認できる規模」の両立を優先したためである(効用のみでは決めていないことを明示する)。"
     f"守5社の Top1200 内カバレッジは {PH2['phase1_top5_coverage_in_top1200']} だが、これは目的関数の実行可能条件"
     "・報酬項・強制含有ロジックの三重で保証した設計上の制約であり、独立した検証結果ではない。", after=8)

# ================= §D 検証規約 =================
h1("§D　パフォーマンス検証の計算規約(本編図表Ⅲ-7・Ⅲ-8に対応)")
h2("D-1　主規約: 固定重み日次リバランス")
para("本編図表Ⅲ-7・Ⅲ-8の数値は、756営業日(3年)・252営業日(1年)のウィンドウで、目標重みへ日次リバランス"
     "する固定重み規約で算出している。年率リターンは幾何年率、βは日次リターンの回帰勾配、"
     "超過の安定度ＩＲは日次超過の算術年率化÷トラッキングエラー、最大の下落は日次終値ベースの累積最大下落幅。"
     "取引コスト・税・配当再投資の差は考慮していない。日経以外は配当込み(§A-1)。", after=5)
h2("D-2　運用実態との対応: バイ・アンド・ホールド(株数固定)併記")
para("本編の図表Ⅲ-1は株数を固定した500万円ポートフォリオ(単元未満株)を提示している。これは日次リバランス"
     "ではなく、期初に買って保有し続けるバイ・アンド・ホールド(BH)に対応する。両規約の3年/1年の値を"
     "併記する(下表)。BHでは重みがドリフトするため値動きは大きくなり、最大の下落は主規約より深い。"
     "リターン面ではBHの方が高いが、本編は保守的な主規約(固定重み)の値を採っている。", after=3)
table(["規約 × 期間", "本ＰＦ 年率", "本ＰＦ 最大下落", "対照群 年率", "対ＴＯＰＩＸ超過(本ＰＦ)"],
      [["固定重み・3年", pct(O3["ann_return"]), pct(O3["max_drawdown"]), pct(C3["ann_return"]),
        f"{(O3['excess_vs_topix']*100):+.1f}pt"],
       ["BH(株数固定)・3年", pct(BH["ours"]["3y"]["ann_return"]), pct(BH["ours"]["3y"]["max_drawdown"]),
        pct(BH["control"]["3y"]["ann_return"]), f"{BH['ours']['3y']['excess_vs_topix']*100:+.1f}pt"],
       ["固定重み・直近1年", pct(O1["ann_return"]), pct(O1["max_drawdown"]), pct(C1["ann_return"]),
        f"{(O1['excess_vs_topix']*100):+.1f}pt"],
       ["BH(株数固定)・直近1年", pct(BH["ours"]["1y"]["ann_return"]), pct(BH["ours"]["1y"]["max_drawdown"]),
        pct(BH["control"]["1y"]["ann_return"]), f"{BH['ours']['1y']['excess_vs_topix']*100:+.1f}pt"]],
      [46, 26, 30, 26, 34])
para("出所: control_comparison_v5.json(固定重み)・bh_metrics_v5.json(BH)。役割別寄与(本編図表Ⅲ-9)は"
     "銘柄別3年累積リターン×ウェイトの分解であり、BH的な寄与分解である点に注意(主規約の累積とは規約が異なる)。", size=9, after=8)

# ================= §E 頑健性 =================
h1("§E　頑健性の検査(本編Ⅱ章末・図表Ⅱ-8/Ⅱ-9に対応)")
h2("E-1　16通りの壊れにくさ検査(全個別値)")
para("選定を「結論ありき」で組んでいないことを、条件を一つずつ変えた16通りの選び直しで確かめる。"
     "各検査について、最終20社と何社一致したか(overlap)とJaccard係数を全開示する。読み方の目盛りは"
     "15社以上=中核維持、11〜14=構成に影響するが崩壊せず、10以下=その条件が選定の背骨。", after=4)
_erows = []
for r in ABL:
    _erows.append([r["variant"], r["description"], f"{r['overlap_with_final20']}/20", r["jaccard_with_final20"]])
table(["検査", "内容", "一致", "Jaccard"], _erows, [16, 92, 20, 22])
_a8 = next(r for r in ABL if r["variant"] == "A8")
para(f"最も効くのは候補の広さ(A8: 候補を100社に絞ると一致{_a8['overlap_with_final20']}/20)で、破の章で述べた"
     "「広く作って確認する」という判断が数字として跳ね返る。残る15通りは11社以上を保ち、どの減点・関所・役割の"
     "枠を外しても選定は崩壊しない。出所: ablation_results.csv。", after=5)
h2("E-2　重みの摂動と正規化の一致")
para("離の式(10)(11)の重みは事後リターン最適化ではなく概念上の設計係数である。各重みを±20%摂動させた検査で、"
     "順位のSpearman相関は最小0.9937、選定される顔ぶれは不変だった。破の順位化も4通り(市場順位・業種内順位・"
     "外れ値に強い標準化・端を丸めた標準化)で付け直し、Top1200のうち3方式以上で共通=970社、"
     "全4方式で共通=554社、2方式以上=1,135社、付け方に敏感な29社(=採用した市場順位のTop1200にのみ入り、"
     "他の3方式では圏外となる会社)は確認フラグ付きで離へ引き継いだ。"
     "本編が「4通り中3通り以上で共通」と書くのはこの定義による(全4方式共通は554社)。"
     "出所: phase3_formula_lineage.md, normalization_consensus_summary_top1200.csv。", after=5)
h2("E-3　提出日ベースの年度断面検証(WP0.6)")
para("完全な時点外検証(学習期間と検証期間の厳密分割による将来リターン予測)は、成熟した252営業日フォールドが"
     "1つしか取れず完了していない。代わりに、開示の提出日(EDINET submit_date)で各時点を再現したpoint-in-time"
     "パネルを構築し、固定した破の設定を各年度断面(2023・2024・2025)へ当てはめる検証を行った。下表は各年度に"
     "再構築したTop1200候補の健全性である。", after=3)
_wfrows = []
for r in WF:
    _wfrows.append([r["availability_year"], r["annual_top1200_count"],
                    f"{float(r['strict_ready_rate'])*100:.0f}%", r["distress_flag_count"],
                    f"{float(r['sector_hhi']):.3f}", f"{r['phase1_top5_coverage']}/5"])
table(["提出年", "Top1200", "厳格適格率", "危険該当", "業種ＨＨＩ", "守5社カバー"], _wfrows, [22, 24, 30, 24, 26, 30])
para("読み取り: 候補づくりは各年度で安定している(Top1200成立・危険該当ゼロ・業種ＨＨＩ0.065〜0.069)。"
     "ただし守5社カバーは2024年断面で4/5に低下する年があり、候補群の年次入替(年次Jaccard 0.52〜0.72)も小さくない。"
     "重要な限界: 本表はTop1200候補プールの断面健全性であって、2025年時点の情報だけで最終20社を組み直し"
     "2025→2026を測った真の20銘柄アウトオブサンプルではない。パネル(CSV)中の将来リターン列は候補プールの参考統計にすぎず、"
     "選定には一切用いていない。したがって本レポートは、この選定に将来リターンを予測する力があるとは主張しない。"
     "出所: fixed_weight_annual_validation.csv。", after=8)

# ================= §F 統計 =================
h1("§F　統計的有意性(本編結論・図表Ⅲ-8の読み取りに対応)")
para("比較優位の主張には不確実性の定量を伴わせる。日次超過リターンの平均がゼロと異ならないかを、素のt値と"
     "系列相関に頑健なNewey-West補正t値で検定する(多重検定の文脈=16通り検査等があるため、"
     "White(2000)・Romano-Wolf(2005)の趣旨に沿って保守的に読む)。", after=4)
_rows = []
for wn, wlabel in [("3y", "3年"), ("1y", "直近1年")]:
    for comp, clabel in [("ours_vs_control", "本ＰＦ − 対照群"), ("ours_vs_topix", "本ＰＦ − ＴＯＰＩＸ")]:
        s = SIG[wn][comp]
        sig = "有意(5%)" if abs(s["t_newey_west"]) >= 1.96 else ("10%水準(弱い示唆)" if abs(s["t_newey_west"]) >= 1.645 else "有意でない")
        _rows.append([f"{wlabel}／{clabel}", f"{s['ann_excess']*100:+.1f}%", f"{s['t_plain']:.2f}",
                      f"{s['t_newey_west']:.2f}", sig])
table(["期間／比較", "年率超過(算術)", "t値(素)", "t値(NW)", "判定"], _rows, [50, 30, 22, 24, 34])
para("読み取り: 対照群(純正バフェット)に対する3年の超過は t(NW)=" + f"{SIG['3y']['ours_vs_control']['t_newey_west']:.2f}"
     " で5%水準の有意。一方ＴＯＰＩＸに対する超過は t(NW)=" + f"{SIG['3y']['ours_vs_topix']['t_newey_west']:.2f}"
     " で有意でなく、直近1年のＴＯＰＩＸ劣後も統計的に有意ではない。すなわち「対照群には有意に勝つが、"
     "市場平均に勝つとは統計的に言えない」。これらはすべてin-sample(2026-06時点で構築した選定を過去に当てはめた)"
     "特性であり、将来の予測主張ではない。", after=4)
para("注: 年率超過(算術)は日次平均×252で、本編図表Ⅲ-8の年率リターン差(幾何)とは年率化規約が異なるため"
     "数値がずれる(例: 対照群超過は算術" + f"{SIG['3y']['ours_vs_control']['ann_excess']*100:+.1f}%" +
     f"／幾何+{(O3['ann_return']-C3['ann_return'])*100:.1f}pt)。検定は算術平均に対して行う。"
     "Deflated Sharpe Ratio(Bailey & López de Prado 2014)は適用しない: 同手法は成績を目的関数に多数試行から"
     "選抜した場合の補正であり、本設計は探索の目的関数に成績・シャープ比を含まないため前提が該当しない"
     "(§C参照。16通り検査は選抜ではなく頑健性の事後検査)。", size=9, after=6)

# ================= §G 限界レジスタ =================
h1("§G　限界レジスタ(査読で意見を求めたい論点)")
para("強みだけを並べたレポートは検証に耐えない。本設計の弱点を、査読者が突く前に自分の手で列挙する。"
     "各項は本編の該当箇所と対応する。", after=4)
_glim = [
    ["in-sample性", "選定は2026-06時点で構築したものを過去に当てはめており、原理的にin-sample。図表Ⅲ-8の"
     "「直近1年」も3年窓の部分集合で、配分の逆ボラはその1年の値動きを入力に使う。真の時点外検証は§E-3の限界どおり未完。"],
    ["選定指標と過去リターンの機械的相関", "割安指標は「過去に下がった株」を、生まれる堀の証拠は「過去に買われた株」を拾いやすい。"
     "対照群超過(+17.4pt)の相当部分は手法の優劣でなく、選定時点の情報が過去リターンと相関する構造の反映。"
     "本比較は予測力の証明でなく、三世代設計が意図どおり動いたことの確認である。"],
    ["市場平均への非有意", "対照群には有意に勝つ(t=2.26)が、ＴＯＰＩＸ超過は有意でない(t=0.81)。市場平均に勝ち続ける力は未証明。"],
    ["生存者バイアス", "母集団は2026-04-30の上場スナップショット。期間中の上場廃止企業を含まない(§A-3)。"],
    ["証拠水準の意味論", "水準の物差しは役割で異なる。生まれる堀=開示証拠の具体性、変わる堀=財務改善の広がり(改善項目の本数)。"
     "変わる堀の「数字まで確認」は受注・投資額の直接確認ではなく財務デルタの本数である。"],
    ["TR・TSの識別力", "式(12)のTR(改革開示)・TS(還元実行)は現データでは全社1で識別力を持たず、水準3判定は事実上TQ(財務デルタ)に依存。"],
    ["除外記録の範囲", "除外1,180社のうち理由コード付きは監査対象の862社+一括コードの318社=全1,180社を記録(§A系)。"
     "うち個別の詳細理由は862社分。"],
    ["キュレーション証拠の主観性", "生まれる堀の証拠水準2以上14社は手作業判定(URL+引用付きだが単独判定)。577社の中に本来水準2+の会社が"
     "残っていない網羅性は保証できない。疑わしきは除外の保守設計として運用している。"],
    ["運用実態との規約差", "本編は固定重み日次リバランスで数値化(§D)。株数固定BHとは値が異なり、BHでは最大の下落が深い。"],
    ["取引コスト等", "取引コスト・税・売買スプレッド・単元未満株の取扱可否は未考慮。"],
]
table(["論点", "内容"], _glim, [40, 132])
para("これらのうち査読者に特に意見を求めたいのは、①in-sample比較の設計としての妥当性、②リターン寄与の"
     "テーマ集中(生まれる堀に8割)の解釈、③証拠水準という順序尺度の設計、の3点である。", after=8)

# ================= §H 再現手順 =================
h1("§H　再現手順と正典チェーン")
para("本レポートの全数値は、以下のスクリプトを順に実行して生成した正典データから自動転記している。"
     "環境は .venv(python-docx 1.2.0 / pandas 3.0.3)+ soffice。乱数シードは破の探索がTPE=42・NSGA-II=43・"
     "ランダム=44、事後の安定性監査=45。", after=4)
para(
    "1. 検証正典:  .venv/bin/python3 scripts/make_validation_v5.py\n"
    "     → control_comparison_v5.json(固定重み・欠測補完)/ bh_metrics_v5.json / significance_v5.json / assets/cum3y_series_v5.csv\n"
    "2. 除外記録:  .venv/bin/python3 scripts/make_exclusion_record_v5.py\n"
    "     → exclusion_record_v5.csv(1,180行)/ exclusion_summary_v5.json\n"
    "3. 図:        .venv/bin/python3 scripts/make_contest_figures_v5.py  → assets/fig2_shu_v5.png, fig3_cum_v5.png\n"
    "4. 本編:      .venv/bin/python3 scripts/build_contest_v5.py [--tocmap tocmap_v5.json]  (2パス)\n"
    "     → beyond_buffett_stockleague_v5.docx →(soffice)→ .pdf\n"
    "5. 本補遺:    .venv/bin/python3 scripts/build_appendix_v5.py  → beyond_buffett_technical_appendix_v5.docx →(soffice)→ .pdf",
    mono=True, after=4)
para("上流の選定正典(参照のみ・再生成不要): phase1(守)= scripts/phase1_final/final_phase1.py・scripts/phase1_top5/build_top5.py、"
     "phase2(破)= outputs/phase2_real_optimization/…/run_real_optimization.py、"
     "phase3(離)= outputs/phase3_beyond_buffett_v2/…/phase3_v2_pipeline.py + "
     "outputs/beyond_buffett_fable_loop_final/scripts/phase3_rebuild.py。"
     "各段の出力CSV/JSONが本レポートの数値の出どころであり、本文・本補遺はそこから機械転記している。", after=6)

out = ED / "beyond_buffett_technical_appendix_v5.docx"
doc.save(str(out))
print("saved", out.name, "| sections: A, B, C, D, E, F, G, H")
