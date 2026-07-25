# -*- coding: utf-8 -*-
"""STOCKリーグ提出版 v2(レビュー対応・約25頁)。
柱1:「超える」の定義を先出しし五役割を演繹 / 柱2: 結論の明確化(判定表) /
柱3: 平易な日本語(禁止語チェック内蔵)+日本語図表 / 柱4: 全文字10pt以上・詳細増補。
使い方: build_contest.py [--tocmap tocmap.json]
"""
import json
import sys
from pathlib import Path

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
sys.path.insert(0, "/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/final_revision/scripts")
import docxlib as X
from docx import Document

# v3書体規定: 和文=MS明朝(欧文・数字はTimes New Roman) / 見出し・図表・表頭(GO)=和文ゴシック+欧文Arial
_orig_setfont = X.setfont


def _setfont_v3(run, name=X.MIN, size=9, bold=False, color=None):
    _orig_setfont(run, name, size, bold, color)
    if name == X.GO:
        from docx.oxml.ns import qn as _qn
        rf = run._element.get_or_add_rPr().find(_qn("w:rFonts"))
        rf.set(_qn("w:ascii"), "Arial"); rf.set(_qn("w:hAnsi"), "Arial")


X.setfont = _setfont_v3
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
ED = HERE.parent
ASSETS = ED / "assets"
# === V7: データ読込を v7 正典へ切替(新バフェット主対照/グレアム参考/単元未満株の目標ウェイト) ===
DATA = json.load(open(ED / "data_real_v7.json", encoding="utf-8"))
for _v in DATA.values():
    _v["code"] = str(_v["code"]).zfill(4)
ROLE_CONTRIB = json.load(open(ED / "role_contribution_v7.json", encoding="utf-8"))
# V9: 全枝ファネル正典(単一base・build_portfolio_v7.pyと同一ガード・出口20社の一致assert済み)
F9 = json.load(open(ED / "funnel_branches_v9.json", encoding="utf-8"))
COMP = json.load(open(ED / "control_comparison_v7.json", encoding="utf-8"))
# 主対照=新バフェットTop12(C)、参考=純正グレアム20(CG)、BH=参考規約
O3, O1 = COMP["ours"]["3y"], COMP["ours"]["1y"]
C3, C1 = COMP["control_buffett"]["3y"], COMP["control_buffett"]["1y"]       # 新バフェット(主対照)
CG3, CG1 = COMP["control_graham"]["3y"], COMP["control_graham"]["1y"]       # 純正グレアム(参考)
BH = COMP["bh_reference"]
D3 = (O3["ann_return"] - C3["ann_return"]) * 100   # 対新バフェット超過(3年)
D1 = (O1["ann_return"] - C1["ann_return"]) * 100   # 対新バフェット超過(1年)
DG3 = (O3["ann_return"] - CG3["ann_return"]) * 100  # 対グレアム超過(3年)
DG1 = (O1["ann_return"] - CG1["ann_return"]) * 100  # 対グレアム超過(1年)
G1 = (O1["ann_return"] - O1["topix_ann_return"]) * 100
G3 = (O3["ann_return"] - O3["topix_ann_return"]) * 100
SIG = json.load(open(ED / "significance_v7.json", encoding="utf-8"))
# SIG構造: ours_vs_buffett / ours_vs_graham / ours_vs_topix (v5の ours_vs_control 相当=ours_vs_buffett)
ROB = json.load(open(ED / "robustness_v7.json", encoding="utf-8"))
# (v9: 旧phase1の7シグナル・ファネルCSV読込を廃止 ― 式番号の衝突と二重ファネルの混乱源だったため。
#  §Ⅱの全社数は F9=funnel_branches_v9.json に一本化)

NAVY = "16324F"; TEAL = "2F6D5F"; TEAL_L = "DCE8E4"; GOLD_BG = "F5EBD0"
NAVYC = RGBColor(0x16, 0x32, 0x4F); TEALC = RGBColor(0x2F, 0x6D, 0x5F)
GO, MIN, WHITE, BLACK = X.GO, X.MIN, X.WHITE, X.BLACK
GRAYC = RGBColor(0x88, 0x88, 0x88)

order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
role_jp = {"Buffett Core": "完成した堀", "Transformation Core": "変わる堀",
           "Emerging Core": "生まれる堀", "Dual Moat": "両立型", "Bridge / Diversifier": "分散役"}
JPN = {"3539": "ＪＭホールディングス", "6430": "ダイコク電機", "7803": "ブシロード",
       "9470": "学研ホールディングス", "4350": "メディカルシステムネットワーク",
       "3697": "ＳＨＩＦＴ", "6841": "横河電機", "9474": "ゼンリン", "6368": "オルガノ",
       "6315": "ＴＯＷＡ", "6920": "レーザーテック", "6526": "ソシオネクスト",
       "5803": "フジクラ", "5902": "ホッカンホールディングス",
       "9828": "Genki Global Dining Concepts", "5233": "太平洋セメント", "8037": "カメイ",
       "3863": "日本製紙", "3089": "テクノアルファ", "2112": "塩水港精糖"}
SEC_JP = {"Retail Trade": "小売業", "Machinery": "機械", "Other Products": "その他製品",
          "Transportation Equipment": "輸送用機器",
          "Information & Communication": "情報・通信業", "Electric Appliances": "電気機器",
          "Nonferrous Metals": "非鉄金属", "Metal Products": "金属製品",
          "Glass and Ceramics Products": "ガラス・土石製品", "Wholesale Trade": "卸売業",
          "Pulp and Paper": "パルプ・紙", "Foods": "食料品"}
THEME_JP = {"non_ai": "ＡＩ以外の堅実分野", "semiconductor": "半導体",
            "quality_assurance": "品質保証", "factory_automation": "工場の自動化",
            "business_data": "業務データ", "optical_communication": "光通信"}

INVEST = sum(v["amtL1"] for v in DATA.values())
CASH = 5_000_000 - INVEST
N_UNBUY = sum(1 for v in DATA.values() if v["price"] * 100 > v["w"] * 5_000_000)
W3, W1 = O3, O1

# ---- バージョン管理: 確定版は編集せず、VERSIONを繰り上げて新版を生成する ----
VER = (ED / "VERSION").read_text().strip() if (ED / "VERSION").exists() else "dev"
TOCMAP = {}
args = sys.argv[1:]
while args:
    a = args.pop(0)
    if a == "--tocmap" and args:
        TOCMAP = json.load(open(args.pop(0), encoding="utf-8"))
    elif a == "--ver" and args:
        VER = args.pop(0)
if (ED / f"beyond_buffett_stockleague_{VER}.LOCKED").exists():
    sys.exit(f"ERROR: {VER} は確定済み(LOCKED)。VERSIONファイルを繰り上げてから編集してください。")

doc = Document()
X.setup_styles(doc)
st = doc.styles["Normal"]
st.font.size = Pt(10.5)          # 既定も10pt以上(規定: 字の級数10.0pt以上)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Mm(18)
sec.top_margin = Mm(16); sec.bottom_margin = Mm(16)

HEADINGS = []


def para(text="", name=MIN, size=10.5, bold=False, align=None, before=0, after=4,
         color=None, indent=False, line=1.24, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if keep:
        pf.keep_with_next = True
    if not indent:
        X._clear_indent(p)
    if text:
        X.setfont(p.add_run(text), name, size, bold, color or BLACK)
    return p


def banner1(text, newpage=False):
    HEADINGS.append((text, 1))
    p = para(before=10, after=5, keep=True)
    pr = p._p.get_or_add_pPr()
    if newpage:
        pr.append(OxmlElement("w:pageBreakBefore"))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    pr.append(shd)
    X.setfont(p.add_run("　" + text), GO, 14, True, WHITE)
    return p


def banner2(text):
    HEADINGS.append((text, 2))
    p = para(before=8, after=4, keep=True)
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), TEAL_L)
    pr.append(shd)
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:left")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "28"); b.set(qn("w:space"), "3"); b.set(qn("w:color"), NAVY)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run("　" + text), GO, 12, True, NAVYC)
    return p


def banner3(text):
    p = para(before=5, after=3, keep=True)
    pr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "dashed"); b.set(qn("w:sz"), "10"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), TEAL)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run(text), GO, 11, True, TEALC)
    return p


def body(text, size=10.5, after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_after = Pt(after); pf.line_spacing = 1.17
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "100"); ind.set(qn("w:firstLine"), "210"); pp.append(ind)
    X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def bullet(label, text, size=10):
    p = para(after=2, line=1.22)
    X.setfont(p.add_run("・" + label + "　"), GO, size, True, NAVYC)
    X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def point(text):
    """章末の要点1行。"""
    p = para(before=4, after=6)
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F6F4")
    pr.append(shd)
    X.setfont(p.add_run("　この章の要点　"), GO, 10, True, TEALC)
    X.setfont(p.add_run(text), MIN, 10, False, BLACK)


def quote(text, who):
    p = para(before=5, after=2, line=1.4)
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480"); ind.set(qn("w:right"), "480"); pp.append(ind)
    r = p.add_run("「" + text + "」"); X.setfont(r, MIN, 11, True, NAVYC); r.font.italic = True
    p2 = para(after=6, align=WD_ALIGN_PARAGRAPH.RIGHT)
    X.setfont(p2.add_run("――　" + who), MIN, 10, False, BLACK)


def fig(fname, title, note, maxw=148):
    fp = ASSETS / fname
    pt_ = para(before=5, after=2, keep=True)
    X.setfont(pt_.add_run(title), GO, 10.5, True, NAVYC)
    if fp.exists():
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        X._clear_indent(pi); pi.paragraph_format.keep_with_next = True
        w, h = Image.open(fp).size; wmm = min(maxw, w / 200 * 25.4)
        if wmm * h / w > 150:
            pi.add_run().add_picture(str(fp), height=Mm(150))
        else:
            pi.add_run().add_picture(str(fp), width=Mm(wmm))
    para(note, MIN, 10, after=7)


def table(hdr, rows, widths, hdr_fill=NAVY, fs=10, zebra=True, aligns=None, spacer=True, keep_caption=False, fixed=False):
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if fixed:
        # 列幅を厳密に守る: tblLayout=fixed + tblGrid を明示(mm→twips)
        _tblPr = t._tbl.tblPr
        _lay = OxmlElement("w:tblLayout"); _lay.set(qn("w:type"), "fixed"); _tblPr.append(_lay)
        _grid = t._tbl.tblGrid
        for _gc, _w in zip(_grid.findall(qn("w:gridCol")), widths):
            _gc.set(qn("w:w"), str(int(_w / 25.4 * 1440)))
    for j, h in enumerate(hdr):
        c = t.rows[0].cells[j]; X.shade_cell(c, hdr_fill)
        pp = c.paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(h), GO, fs, True, WHITE)
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        for j in range(len(hdr)):
            if zebra and i % 2 == 1:
                X.shade_cell(cells[j], "F2F6F4")
            pp = cells[j].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
            pp.alignment = (aligns[j] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            X.setfont(pp.add_run(str(r[j]) if j < len(r) else ""), MIN, fs, False, BLACK)
    for r in t.rows:
        X.cant_split(r)
        for j, w in enumerate(widths):
            r.cells[j].width = Mm(w)
    if keep_caption:
        # C-5対策: 最終行の段落を後続(キャプション)と同頁に保つ
        for c in t.rows[-1].cells:
            for pp in c.paragraphs:
                pp.paragraph_format.keep_with_next = True
    if spacer:
        para("", after=3, keep=keep_caption)
    return t


def formula(num, title, latex, defn):
    t = doc.add_table(rows=3, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in t.rows:
        X.cant_split(r); r.cells[0].width = Mm(170)
    c0 = t.rows[0].cells[0]; X.shade_cell(c0, NAVY); X.cell_borders(c0, top=True, bottom=True)
    p0 = c0.paragraphs[0]; X._clear_indent(p0); p0.paragraph_format.space_after = Pt(0); p0.paragraph_format.keep_with_next = True
    X.setfont(p0.add_run(f"　式（{num}）{title}"), GO, 10.5, True, WHITE)
    c1 = t.rows[1].cells[0]; X.cell_borders(c1, bottom=True)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p1)
    p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(4); p1.paragraph_format.keep_with_next = True
    om = X.latex_to_omml(latex)
    from lxml import etree
    M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    pr = om.find(f"{{{M}}}oMathParaPr")
    if pr is None:
        pr = etree.Element(f"{{{M}}}oMathParaPr"); om.insert(0, pr)
    jc = pr.find(f"{{{M}}}jc")
    if jc is None:
        jc = etree.SubElement(pr, f"{{{M}}}jc")
    jc.set(f"{{{M}}}val", "center")
    p1._p.append(om)
    c2 = t.rows[2].cells[0]; X.cell_borders(c2, bottom=True)
    p2 = c2.paragraphs[0]; X._clear_indent(p2)
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    X.setfont(p2.add_run(defn), MIN, 10, False, BLACK)
    para("", after=3)


def fill_line(label, hint=""):
    p = para(after=2)
    X.setfont(p.add_run("・" + label + "："), MIN, 10.5, True, BLACK)
    X.setfont(p.add_run("［記入］"), MIN, 10.5, False, GRAYC)
    if hint:
        p2 = para(after=3)
        pp = p2._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360"); pp.append(ind)
        X.setfont(p2.add_run("書き方のヒント：" + hint), MIN, 10, False, GRAYC)


def step_chip(no, name):
    p = para(before=7, after=2, keep=True)
    r = p.add_run("　スクリーニング" + "０１２３４"[no] + "　")
    X.setfont(r, GO, 10.5, True, WHITE)
    rpr = r._element.get_or_add_rPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    rpr.append(shd)
    X.setfont(p.add_run("　" + name), GO, 11.5, True, NAVYC)
    HEADINGS.append((name, 2))


def add_page_number(section):
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("- ")
    X.setfont(run, MIN, 10, False, BLACK)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r2 = OxmlElement("w:r"); t2 = OxmlElement("w:t"); t2.text = "2"
    r2.append(t2); fld.append(r2)
    fp._p.append(fld)
    run3 = fp.add_run(" -")
    X.setfont(run3, MIN, 10, False, BLACK)


# ================= 表紙(全面) =================
sec.top_margin = Mm(0); sec.bottom_margin = Mm(0)
sec.left_margin = sec.right_margin = Mm(0)
pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc.paragraph_format.space_before = Pt(0); pc.paragraph_format.space_after = Pt(0)
pc.add_run().add_picture(str(ASSETS / "cover.png"), width=Mm(210))
ns = doc.add_section(WD_SECTION.NEW_PAGE)
ns.left_margin = ns.right_margin = Mm(18); ns.top_margin = Mm(16); ns.bottom_margin = Mm(16)
add_page_number(ns)

# ================= 要旨 =================
banner1("要旨")
body("本レポートでは、これからの日本株投資の姿として「三世代の堀(Three-Generation Moat)」を提示する。堀(Moat)とは、城のまわりの堀のように、他社がまねしにくい持続的な競争優位のことである。私たちは、ウォーレン・バフェットが築いた「完成した堀を割安に買う」という規律を出発点に、堀を時間軸でとらえ直した。いま完成している堀だけでなく、資本効率改革によってこれから「変わる堀」、ＡＩ・半導体・光通信などの構造変化からこれから「生まれる堀」までを、一つのポートフォリオに束ねる。")
body("本レポートの軸は一つだけである。目指す先が「三世代の堀」(何を買うか)、そこへ至る道順が「守破離」(どう選ぶか)であり、堀が主役で守破離は選ぶ手順の名前にすぎない。守ではバフェットの品質基準(先行研究の式)だけで完成した堀を5社選び、破ではその型を破って割安×変革の会社(変わる堀)を、離では独自の物差しで新しい堀(生まれる堀)を選び、両立型3社・分散役2社を加えて最終20社を組んだ。選定後に成績を見て銘柄を入れ替える「後からの選び直し」はしていない(あとから成績の良い銘柄だけを拾えば、どんな選び方でも立派に見えるからだ)。その代わり、この選び方が条件を一つ変えたくらいでは崩れないことを、三通りの頑健性検査(品質ゲートの判定の目安を変えること・配分の重み方式・相場局面)で確かめた(詳細は第Ⅱ章末)。")
para("※守破離(しゅはり)＝日本の芸道の修行段階。守=型を守る、破=型を破る、離=独自の境地へ。", MIN, 10, after=4)
body(f"結論を先に言う。私たちは「超える」を四つの条件で定義した。①バフェットと同じ土俵に立つ、②彼の式では見えない15社を同じ厳しさの証拠で選ぶ、③選定が壊れにくい――この三つの設計条件は本文で確かめた。残る④成績は、バフェットの品質基準で選んだ『新バフェット型』を自作し、先に決めた物差しで比較した。過去データの自己検証では、割安一辺倒のグレアム型と市場平均(ＴＯＰＩＸ)には、リスクを抑えつつ統計的にも有意に上回った(3年で対グレアム＋{DG3:.0f}ポイント)。師である新バフェット型に対しては、リターン・効率(シャープ)・下落の小ささのすべてで、しかも四つの相場局面すべてで一貫して上回った(3年で年＋{D3:.0f}ポイント)が、その差は統計的には確実でなく、正直に「互角〜やや上」と評価する。これはすべて過去の自己検証であり、将来の市場超過の証明ではない――この限界は隠さず記す。")

banner1("目次")
toc_items = ["要旨", "Ⅰ．背景・投資テーマ決定", "Ⅱ．スクリーニング ― 守・破・離",
             "Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス", "Ⅳ．インタビュー・アンケート",
             "Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと", "参考文献"]
for it in toc_items:
    p = para(after=1)
    X.setfont(p.add_run(it), MIN, 10.5, False, BLACK)
    pg = TOCMAP.get(it, "")
    if pg:
        X.setfont(p.add_run("　" + "･" * max(4, 52 - len(it) * 2) + "　" + str(pg)), MIN, 10.5, False, BLACK)

# ================= Ⅰ 背景・投資テーマ決定 =================
banner1("Ⅰ．背景・投資テーマ決定", newpage=True)
banner2("1．イントロダクション ― 百年企業の中に、未来の堀が眠っていた")
quote("素晴らしい会社をまずまずの価格で買うほうが、まずまずの会社を素晴らしい価格で買うよりずっと良い。",
      "ウォーレン・バフェット(1989年 株主への手紙)")
body("バフェットは、企業を城に、競争優位を城を守る堀(Moat)にたとえた。深い堀を持つ会社を、高すぎない価格で買い、長く持つ。この単純な規律が、半世紀を超える成果を生んだ。私たちの研究は、この規律への挑戦ではなく、その延長である。")
body("きっかけは、一つの会社だった。santec Holdings(6777)――光通信の部品をつくる会社である。地味な部品メーカーと見られてきたこの会社は、いま、生成ＡＩのデータセンター内をつなぐ高速の光配線(光インターコネクト)の担い手として需要が急拡大している。データが詰まると計算全体が止まる――その急所を、精密な光の部品技術が握った。完成したはずの部品事業の内側で、ＡＩという新しい堀が生まれていた。")
body("同じ頃、東京証券取引所は上場企業に「資本コストや株価を意識した経営」を求め、長く割安に放置された企業が、資本効率の改善と株主還元の強化によって変わり始めた。堀は止まった絵ではない。変わる堀があり、生まれる堀がある。ならば、いまの堀だけを測るバフェットの物差しを、時間の方向へ延長できないか。これが私たちの問いである。")

banner2("2．バフェットの規律 ― 何がすごいのかを分解する")
body("「超える」と言う前に、まず相手を正しく知る。バフェットの投資は、大きく四つの部品でできている。")
bullet("堀(持続する競争優位)", "まねされにくい強み。ブランドの力(例:コカ・コーラ)、乗り換えの手間、利用者が増えるほど価値が増す仕組み、圧倒的なコストの安さ、特許や免許。バフェット(2007)は「永続的な堀」を持つ企業だけを買うと明言する。")
bullet("安全域(高すぎない価格)", "どんな良い会社も、高く買いすぎれば損をする。会社の価値より安い価格で買い、見立て違いに備える。冒頭の1989年の言葉は、価格より会社の質を優先するという順序も示している。")
bullet("能力の輪(分かるものだけ買う)", "自分が理解できる事業だけを対象にする。理解できない熱狂には乗らない。")
bullet("長期保有(複利で増やす)", "「私たちの好みの保有期間は永遠である」(バフェット、1988)。頻繁な売買をせず、堀が守る利益の再投資で雪だるま式に増やす。")
body("この規律は、感覚ではなく式で確かめられる。Frazzini, Kabiller and Pedersen (2018)は、バフェットの半世紀の超過成績が「割安×優良×値動きの穏やかさ」という測定可能な要素でほぼ説明できることを示した。つまり、彼の物差しは先行研究の式で再現できる――これが第Ⅱ章「守」の土台である。")
body("同時に、この分解は彼の物差しの死角も教えてくれる。割安・優良・穏やかさは、どれも「いま完成している堀」を測る物差しであり、これから変わる堀・生まれる堀は原理的に映らない。私たちが加える視点は三つ――①時間軸(堀の一生を見る)、②証拠主義(言葉ではなく数字の裏づけで測る)、③分散(どれか一つの未来に賭けない)。この三つを、彼と同じ規律の厳しさで実行することが、本レポートの挑戦である。")

banner2("3．背景 ― 二つの地殻変動")
body("第一の地殻変動は資本効率改革である。東京証券取引所(2023)の要請は、伊藤レポート(経済産業省、2014)以来の資本収益性重視の流れを決定づけ、ＰＢＲ(株価純資産倍率。株価が会社の純資産の何倍かを示す)1倍割れ企業の経営改革を促した。割安に放置された企業の中から、本当に変われる企業を見分けられれば、それは「変わる堀」への投資になる。")
body("第二の地殻変動はＡＩを起点とする産業構造の転換である。生成ＡＩは、半導体・光通信・データセンター・電力・品質保証といった基盤産業に、実際の需要という追い風を与えている。ただし「ＡＩ」と口にする企業がすべて堀を持つわけではない。実際、既製のＡＩ関連スコアは機械的なキーワード照合に頼るため、炊飯器や照明の会社までＡＩ関連として満点に並んでしまうことが分かった。「ＡＩ」と口にすることと、ＡＩの実需で稼ぐことは違う。接点と堀を見分ける仕組み――事業の開示まで遡って実需を確かめる証拠主義――が要る。")

banner2("4．三世代の堀 ― 私たちが求める企業像")
body("以上から、私たちは投資対象を三つの世代で定義する。第一に、すでに完成した堀(バフェットの規律で測れる企業)。第二に、資本効率改革で変わる堀(割安×改善の証拠を持つ企業)。第三に、構造変化から生まれる堀(ＡＩ基盤への実需接続を証拠で示せる企業)。この三世代を一つのポートフォリオに束ね、どれか一つの未来に賭けない。これが「三世代の堀」である。")
fig("fig1_moat_v7.png", "図表Ⅰ-1　三世代の堀 ― 完成・変化・新生を一つに束ねる",
    "注：上段は三世代の名前、下段は中身。両立型・分散役を加えた合計20社で構成する。この三世代(何を買うか)を、第Ⅱ章では守・破・離の順番(どう選ぶか)で選び出す。筆者作成。")
p_h = para(before=5, after=7)
_pr = p_h._p.get_or_add_pPr(); _shd = OxmlElement("w:shd")
_shd.set(qn("w:val"), "clear"); _shd.set(qn("w:color"), "auto"); _shd.set(qn("w:fill"), GOLD_BG)
_pr.append(_shd)
X.setfont(p_h.add_run("　投資仮説　"), GO, 10.5, True, NAVYC)
X.setfont(p_h.add_run("完成した堀・変わる堀・生まれる堀を証拠で選んで一つに束ねたポートフォリオは、堀の“現在”だけを測る新バフェット型の選定を、リスク調整後の効率(シャープ)と下落の小ささで上回る。過度な集中を避けた三世代の分散が、攻め(リターン)と守り(下落の抑制)の両面で効く――これが、三世代を等しく束ねる理由である。"), MIN, 10.5, True, BLACK)

banner2("5．何をもって「バフェットを超える」とするか ― 四つの条件")
body("「超える」を気分で語らないために、確かめられる四つの条件で先に定義しておく。本レポートの残りは、この定義の検証である。")
table(["条件", "内容", "どこで確かめるか"],
      [["① 同じ土俵に立つ", "バフェットの選び方を先行研究の式だけで再現し、その方法が選ぶ5社を実際に持つ。比較相手を自分のポートフォリオの中に持ち歩く", "第Ⅱ章「守」"],
       ["② 死角を補う", "バフェットの式では選ばれない15社を、同じ厳しさの証拠で選ぶ。彼の物差しが見ない未来を、彼と同じ規律で測る", "第Ⅱ章「離」"],
       ["③ 壊れにくい", "選定が一つの仕掛け・一つのテーマに依存しない。判定の目安=条件を一つ変えても20社中11社以上が維持される(候補の広さの変更は別扱い)", "第Ⅱ章末の検査"],
       ["④ 総体で上回る", "主対照『新バフェット型』(品質基準で選んだ大型優良)に対し、リターン・効率ＩＲ・下落の小ささで上回る。参考の純正グレアム型・市場ＴＯＰＩＸには有意に上回る。物差しは比較前に宣言", "第Ⅲ章の検証"]],
      [34, 102, 34], fs=10)
para("図表Ⅰ-2　「超える」の定義 ― 四つの条件と検証の場所。", MIN, 10, after=5)
body("言い換えれば、三世代の堀が目的地の地図(何を買うか)であり、守破離はそこへ歩く順番(どう選ぶか)である。順番には理由がある。①〜③を後から誰でも検証できる形にするため、独自色を最後(離)に置いた。独自の物差しを最後に導入するからこそ、「どこからが私たちの主張か」をいつでも遡って示せる。")

banner2("6．検証の順番 ― 守・破・離")
body("守=条件①、離=条件②、章末の頑健性検査=条件③、第Ⅲ章の対照比較=条件④に対応する(下表)。物差しは比較の前に宣言する。")
table(["段階", "選ぶ堀とやること", "どこまでが借り物か", "検証する条件"],
      [["守", "完成した堀5社を、先行研究の式だけで選ぶ", "式も使い方も、すべて先行研究のまま", "条件①"],
       ["破", "型を破り、割安×変革の会社(変わる堀)5社を選ぶ", "式は借り物。使い方は私たちの判断(過程を開示)", "(条件②の一部)"],
       ["離", "変わる・生まれる堀＋両立・分散の15社を、独自の物差しで選ぶ", "部品は借り物。組み方が私たちの独自性", "条件②"]],
      [16, 72, 56, 26], fs=10)
para("図表Ⅰ-3　守・破・離の分担 ― 独自性を最後に置く。守破離は目的ではなく、四条件を検証可能にするための手段である。", MIN, 10, after=5)
body("なお、本文の数字はすべて選定の過程で保存した記録(通過社数・除外理由・検査結果)から機械的に転記し、レポート自体を記録データから自動生成している。生成・検証コードと正典データ・乱数シード(頑健性の摂動監査に使用)は提出時点で保存しており、どの数字も出どころまで遡って確かめられる。")
# (v6: SDGs記入欄はユーザー指示で削除。公式様式で必須の場合は復活させること)

# ================= Ⅱ スクリーニング =================
banner1("Ⅱ．スクリーニング ― 守・破・離", newpage=True)
body(f"前章では「超える」を四つの条件で定義し、選ぶ順番として守破離を置いた。本章が選定の中核である。出発点は、金融を除く普通株{F9['n_nonfin']:,}社。まず全銘柄共通の関所(スクリーニング0)で{F9['n_base']:,}社に絞り、そこから守(スクリーニング1)・破(同2)・離(同3)の三本の枝がそれぞれ5社を選び、支え役(同4)が両立型3社と分散役2社を加えて、合流した20社が最終ポートフォリオになる。5＋5＋5＋3＋2＝20。本章の書き方は一つだけ徹底する――各スクリーニングの見出しに「何社から何社へ」を必ず書く。まず全体の設計図を一枚で示す。")
fig("fig2_overview_v9.png", f"図表Ⅱ-1　スクリーニングの全体設計図 ― {F9['n_nonfin']:,}社から20社へ",
    "注：スクリーニング0は全枝共通の関所。1〜4は同じ出発点(1,791社)から役割別に選ぶ並列の枝であり、番号は絞り込みの直列順ではなく読み順。全社数は選定スクリプトと同一条件で再計算し、20社の一致を機械検査した。出所：funnel_branches_v9.json。")
body("段の名前は第Ⅰ章の三世代にそのまま対応する――守＝完成した堀、破＝変わる堀、離＝生まれる堀。軸は「時間」の一本だけである。以降の各スクリーニングは、ねらい(何を落とし、何を残すか)→基準(式)→結果(何社→何社)→通し例(代表1社で式の動きを追う)、という同じ型で繰り返す。データの基準は、価格が2023年4月25日〜2026年6月1日の日次、財務・開示がＥＤＩＮＥＴ提出書類(2026年7月8日取得。各社は取得日時点の直近有価証券報告書で、決算期の違いにより最大12カ月の鮮度差がある)である。")

step_chip(0, f"共通の関所 ― すべての枝の出発点({F9['n_nonfin']:,}社→{F9['n_base']:,}社)")
body(f"ねらい: 「選ぶ」前に、「選んではいけない会社」を全枝共通で除く。三つの関所を通す。第一に投資適格(監理・整理銘柄などを除外)、第二に流動性(60日平均売買代金が基準額以上＝買いたい量を、株価を動かさずに買える)――ここまでで{F9['n_nonfin']:,}社が{F9['n_eligible']:,}社になる。第三に価格履歴3年(756営業日)で{F9['n_base']:,}社。三つ目は選ぶための関所ではなく、確かめるための関所である――過去3年へ当てて成績を検証できない会社を残すと、第Ⅲ章の自己検証が成り立たない。")
body(f"正直な限界も先に書く。出発点の{F9['n_nonfin']:,}社は直近時点の上場一覧で、期間中に上場廃止となった会社を含まない(生存者バイアス)。また履歴3年の関所は、上場3年未満の若い会社を機械的に外す――生まれる堀の候補を、検証可能性と引き換えに少し狭めている(技術補遺§A)。", after=6)

step_chip(1, f"守 ― 完成した堀を品質の七関門で選ぶ({F9['n_base']:,}社→5社)")
body("ねらい: バフェットの規律「良い会社を、高すぎない価格で買い、長期保有に耐えない企業を避ける」を、先行研究の式だけで再現する。独自の重み付けは使わず、七つの関門を一つずつ通す。ここで選ぶ5社は「何を超えたのか」を測る基準線であり、以降一度も入れ替えない。")
def verdict(text):
    p = para(after=6)
    X.setfont(p.add_run("→ 判定　"), GO, 10.5, True, TEALC)
    X.setfont(p.add_run(text), MIN, 10.5, True, BLACK)


body("七つの関門を、式そのものから一つずつ見ていく。どの式も先行研究のまま使い、私たちの独自の重みは一切加えない。")
banner3("関門1　資本を高い率で回せるか ― ＲＯＥ")
body("自己資本利益率(ＲＯＥ)は、株主の資本をどれだけ効率よく利益に変えるかを示す。高く持続的なＲＯＥは、まねされにくい強み＝堀の現れである。バフェットの超過成績が「割安×優良×安全」で説明できることを示したFrazzini, Kabiller and Pedersen (2018)、および品質factorを定式化したAsness, Frazzini and Pedersen (2019, QMJ)で、収益性は品質の中核に置かれる。", after=2)
formula("1", "ＲＯＥ ― 資本の収益率(Frazzini, Kabiller and Pedersen, 2018)",
        r"\mathrm{ROE}_i = \frac{\mathrm{NI}_i}{\mathrm{Equity}_i} \ \ge\ 0.15",
        "ＮＩ＝純利益　Equity＝自己資本。15％以上を高収益の基準とする(バフェットが繰り返し目安とする水準)。")
verdict(f"よって、ＲＯＥ15％以上の企業を「資本を高い率で回せる優良企業」と評価する(1-1: 残り{F9['shu']['steps'][0]['n']:,}社)。")
banner3("関門2　価格支配力＝堀があるか ― 営業利益率")
body("高く安定した営業利益率は、値下げ競争に巻き込まれない価格支配力＝堀の現れである。薄利の会社は、原価や競合の変化で利益が消えやすい(QMJ の収益性)。", after=2)
formula("2", "営業利益率 ― 価格支配力(堀)の代理(Asness, Frazzini and Pedersen, 2019)",
        r"\mathrm{OPM}_i = \frac{\mathrm{OperatingIncome}_i}{\mathrm{Revenue}_i} \ \ge\ 0.10",
        "本業の利益÷売上高。10％以上を、堀(価格支配力)のある水準とする。")
verdict(f"よって、営業利益率10％以上の企業を「価格支配力＝堀を持つ」と評価する(1-2: 残り{F9['shu']['steps'][1]['n']:,}社)。")
banner3("関門3　財務は健全か ― 自己資本比率")
body("負債に頼らず高いＲＯＥを出せることが、真の実力である。借入で見かけのＲＯＥを膨らませた会社は、不況で脆い。QMJの「安全」の考え方に沿い、自己資本比率で財務の健全さを測る(Asness et al., 2019)。", after=2)
formula("3", "自己資本比率 ― 財務の安全(Asness, Frazzini and Pedersen, 2019)",
        r"\mathrm{EqR}_i = \frac{\mathrm{Equity}_i}{\mathrm{TotalAssets}_i} \ \ge\ 0.50",
        "自己資本÷総資産。50％以上を、負債に頼らない健全な財務とする。")
verdict(f"よって、自己資本比率50％以上の企業を「財務が健全」と評価する(1-3: 残り{F9['shu']['steps'][2]['n']:,}社)。")
banner3("関門4　利益は予測可能か ― 直近3期無赤字")
body("堀のある企業は、不況でも赤字に沈みにくい。逆に、直近数年に赤字を出す会社は利益が読みにくく、長期保有の前提が崩れる。直近3期に営業・純・営業キャッシュフローのいずれも赤字がないことを、予測可能性の条件とする。", after=2)
formula("4", "無赤字 ― 予測可能性(Asness, Frazzini and Pedersen, 2019)",
        r"\mathrm{LossYears}^{3\mathrm{yr}}_i = 0 \quad (\text{operating, net, operating-CF})",
        "直近3期の営業損失・純損失・営業CF赤字の年数がいずれもゼロの企業だけを通す。")
verdict(f"よって、直近3期に一度も赤字のない企業を「利益が予測可能」と評価する(1-4: 残り{F9['shu']['steps'][3]['n']:,}社)。")
banner3("関門5　利益に現金の裏づけがあるか ― 営業ＣＦ")
body("会計上の利益は見積りや会計方針で膨らませる余地がある。本業で実際に現金を生んでいるか(営業キャッシュフローがプラスか)を、利益の質の条件とする。帳簿利益と現金の乖離が大きい会社は後で失速しやすい(Sloan, 1996)。", after=2)
formula("5", "営業キャッシュフロー ― 利益の質(Sloan, 1996)",
        r"\mathrm{CFO}_i \ >\ 0",
        "ＣＦＯ＝営業キャッシュフロー(本業の現金収入)。プラスの企業だけを通す。")
verdict(f"よって、本業で現金を生んでいる企業を「利益が本物」と評価する(1-5: 残り{F9['shu']['steps'][4]['n']:,}社)。")
banner3("関門6　事業は伸びているか ― 増収かつ増益")
body("堀は成長とともに広がり、縮む事業では狭まる。減収・減益の会社は、堀が痩せている可能性がある。前期比で売上と営業利益がともに減っていないこと(増収かつ増益)を、非縮小の条件とする(QMJ の成長性)。", after=2)
formula("6", "増収かつ増益 ― 非縮小(Asness, Frazzini and Pedersen, 2019)",
        r"g^{\mathrm{Rev}}_i \ \ge\ 0 \quad \land \quad g^{\mathrm{OI}}_i \ \ge\ 0",
        "ｇ＝前期比の伸び率。売上(Rev)と営業利益(OI)がともに前期を下回らない企業を通す。")
verdict(f"よって、増収かつ増益の企業を「事業が縮小していない」と評価する(1-6: 残り{F9['shu']['steps'][5]['n']:,}社)。")
banner3("関門7　実際に買えるか ― 売買のしやすさ")
body("売買が少ない株は、買うだけで値段が動いてしまい、計画どおりの投資ができない。流動性が価格に影響することはAmihud (2002)が実証しており、その非流動性指標も日々の売買代金を基礎とする。私たちは同じ基礎量を、実務で標準的な「一定額以上の売買代金があるか」という関門の形で用いた。", after=2)
formula("7", "60日平均売買代金 ― 売買のしやすさ(Amihud, 2002)",
        r"\mathrm{ADV}^{60}_i = \frac{1}{60}\sum_{t=1}^{60} P_{i,t}\,V_{i,t}",
        "Ｐ＝日々の株価　Ｖ＝日々の出来高(売買された株数)。直近60営業日の売買代金の平均。1日あたり1,000万円を基準とする。")
verdict("よって、1日平均およそ1,000万円以上の売買がある企業を「実際に買える」と評価する。この段階では300万〜1,000万円の会社も「要確認」の印をつけて通過させ、Top5選定では1,000万円以上を明確に要求する(実装どおりに開示)。")
body(f"結果: 七つの関門で、適格{F9['n_base']:,}社は次のように絞られた。落ちた会社は「どの関門を満たさなかったか」で理由が特定できる――落ちた理由をすべて残すことが、選んだ理由の裏づけになる。")
_sh = F9["shu"]["steps"]
_reasons = ["資本を高い率で回せていない", "価格支配力＝堀が薄い", "負債が重く長期保有に不向き",
            "利益が不安定で読みにくい", "利益に現金の裏づけがない", "事業が縮小している"]
table(["守の関門", "通過社数", "落ちる主な理由"],
      [[f"1-{i+1}　{st['label']}", f"{st['n']:,}社", r] for i, (st, r) in enumerate(zip(_sh[:6], _reasons))]
      + [[f"1-7　{_sh[6]['label']}", f"{_sh[6]['n']:,}社", "時価総額データの在庫がなく割安順位を付けられない"]],
      [66, 20, 70], fs=10)
para(f"図表Ⅱ-2　守の結果 ― 適格{F9['n_base']:,}社が{F9['shu']['n_quality']}社、価格ランク可能{F9['shu']['n_priceable']}社まで絞られる。関門7(売買代金)は共通の関所(スクリーニング0)で先に適用済み。出所：funnel_branches_v9.json。", MIN, 10, after=5)
g = DATA["6920"]
body("式が実際にどう働くかを、通過企業の一社・レーザーテック(6920)で追ってみる。半導体マスク検査で世界的地位をもつ、完成した堀の典型例である。")
table(["品質関門", "レーザーテックの値", "判定の目安", "判定"],
      [["高収益 ＲＯＥ(式1)", f"{g['roe']:.0%}", "15％以上", "通過"],
       ["堀・価格支配力 営業利益率(式2)", f"{g['opm']:.0%}", "10％以上", "通過"],
       ["財務健全 自己資本比率(式3)", f"{g['eqr']:.0%}", "50％以上", "通過"],
       ["予測可能性 直近3期無赤字(式4)", "営業・純・営業CFとも赤字なし" if g['loss_free'] else "赤字あり", "3期連続で赤字ゼロ", "通過"],
       ["現金創出 営業ＣＦ(式5)", "本業の現金収入プラス", "営業ＣＦ＞0", "通過"],
       ["非縮小 増収増益(式6)", f"増収＋{g['rev_g']:.0%}・増益＋{g['oi_g']:.0%}", "ともに0％以上", "通過"],
       ["実際に買えるか 売買代金(式7)", f"1日平均 {g['adv_oku']:.1f}億円", "約0.1億円以上", "通過"]],
      [42, 44, 40, 20], fs=10)
para("図表Ⅱ-3　通し数値例：レーザーテック(6920)の品質関門。出所：選定の正典データ(data_real_v7)より筆者作成。", MIN, 10, after=6)
b5 = [DATA[c] for c in ["3092", "4716", "7014", "8136", "6920"]]
body(f"最後の絞り込み: 品質{F9['shu']['n_quality']}社のうち価格ランクが可能な{F9['shu']['n_priceable']}社を、割安×優良の複合順位(収益力ＲＯＥと益回りの順位和＝Greenblattの考え方)で並べ、同一業種は原則2社までとして上位5社を固定した。選ばれたのは、"
     + "、".join(f"{d['code']} {d['name_ja']}({SEC_JP.get(d['sector'], d['sector'])})" for d in b5)
     + "の5社である。なお、この順位の上位12社が第Ⅲ章で比較する主対照『新バフェット型』であり、守の5社はその最上位5社に一致する――比較相手を自分のポートフォリオの中に持ち歩く、という条件①はこの一致で担保される。")
_t5 = table(["コード・社名(完成した堀Top5)", "業種", "ＲＯＥ", "益回り(Ｅ／Ｐ)"],
      [[f"{d['code']} {d['name_ja'].replace('株式会社', '')}", SEC_JP.get(d['sector'], d['sector']),
        f"{d['roe']:.0%}", f"{d['ep']:.1%}"] for d in b5],
      [62, 34, 26, 34], fs=10)
for _r5 in _t5.rows[:-1]:
    for _c5 in _r5.cells:
        for _pp5 in _c5.paragraphs:
            _pp5.paragraph_format.keep_with_next = True
para("図表Ⅱ-4　守Top5の中身 ― 割安×優良の複合順位(ＲＯＥの順位＋益回りの順位)の上位5社。同一業種は2社まで。出所：data_real_v7.json。", MIN, 10, after=5)
fig("fig2_shu_v9.png", f"図表Ⅱ-5　スクリーニング1(守)の絞り込み ― 適格{F9['n_base']:,}社から完成した堀Top5へ",
    "注：棒の長さは残った会社数(対数目盛)。1-1〜1-6は品質の六関門、1-7は価格ランク可能、1-8で5社に固定。出所：funnel_branches_v9.json。")
verdict("受け渡し: 守の5社(基準線)が確定した。しかし守の品質レンズは「いま完成している堀」しか映さない――次の破は、その死角(割安×変革)を探しに行く。")

step_chip(2, f"破 ― 型を破り、変わる堀を選ぶ({F9['n_base']:,}社→5社)")
body(f"ねらい: 守の品質レンズは「いま完成している堀」しか映さない。破では、その型を一度『破り』、いまは割安に放置されているが、東証の資本効率改革や脱炭素の流れでこれから評価が変わる会社(変わる堀)を探す。適格{F9['n_base']:,}社のうち変わる堀に分類されるのは{F9['ha']['steps'][0]['n']}社――ここから黒字と最低限の収益性で絞り、変わる堀の点数の上位5社を選ぶ。まず点数の作り方(手順1〜3と式10)、次に結果(何社→何社)を示す。")
banner3("手順1　単位の違う7つの部品を、どう公平に比べるか ― 順位化")
body("変わる堀の点数は、先行研究由来の7つの部品から組み立てる(下表)。部品は単位がばらばらで、そのままでは足し合わせられない。そこで各指標を「市場の中で下から何番目か」という順位の割合(0〜1)に置き換える。これは資産価格研究がＢ／Ｍなどの指標で銘柄を上位30％・中位40％・下位30％のように順位で区切って比べる慣行(Fama and French, 1993 ほか)を、区切りではなく連続的な順位点として使うものである。", after=2)
table(["部品", "何を測るか", "出典(先行研究)"],
      [["① 純資産比(Ｂ／Ｍ)", "純資産に対して安いか", "Fama and French (1993)"],
       ["② 益回り(Ｅ／Ｐ)", "利益に対して安いか", "Basu (1977; 1983)"],
       ["③ 粗利÷総資産", "収益力は本物か", "Novy-Marx (2013)"],
       ["④ 財務6項目の合格割合", "財務は健全か", "Piotroski (2000) の簡約版"],
       ["⑤ 利益と現金の差", "利益の質(会計の水増しがないか)", "Sloan (1996)"],
       ["⑥ 危険よけ", "債務超過・連続赤字の回避", "Altman (1968)・Ohlson (1980) の考え方"],
       ["⑦ 売買のしやすさ", "60日平均売買代金", "Amihud (2002)"]],
      [42, 62, 52], fs=10)
para("図表Ⅱ-6　変わる堀の点数を組み立てる7つの部品と出典。測る量はすべて先行研究にあるものだけを使う。", MIN, 10, after=5)
formula("8", "順位化 ― 市場の中の相対的な位置に置き換える(Fama and French, 1993)",
        r"r_{i,k} = \frac{\operatorname{rank}_k(x_{i,k})}{N}",
        "ｘ＝指標ｋの値　rank＝市場の中での順位(良い方が大きい)　Ｎ＝算出できた会社数。0〜1の点数になり、単位の違う指標同士を公平に比べられる。")
verdict("よって、7つの部品すべてを0〜1の順位点に変換して比べる。式の定義は一切変えない。")
banner3("手順2　順位点をどう束ねるか ― 重みつき合成点(変革の論理で設計)")
body("各順位点に重みを付けて足し合わせ、危険な兆候(データ欠け・極端に小さい会社など)には減点を与える。複数のシグナルを一つの合成点に束ねる形は、9項目の合計点で勝ち組と負け組を分けたPiotroski (2000)のＦスコアと同型であり、私たちはその「合計」を「重みつき合計−減点」に一般化した(式9)。重みは過去リターンへの当てはめでは決めない――どの重みなら過去の成績が良かったか、という決め方は後出しの罠だからである。『安さ』より『改善の実行』を重くする、という変革の論理から概念的に設計し、決めた後に順位の安定性を監査した(重みの全ベクトルと監査は技術補遺§C)。", after=2)
formula("9", "合成点 ― 候補づくりのための重みつき合計(Piotroski, 2000)",
        r"S_i = \sum_{k=1}^{7} w_k\, r_{i,k} - P_i",
        "ｗ＝指標ｋの重み(変革の論理で設計し、全設定を記録)　ｒ＝順位点(式8)　Ｐ＝減点(異常値・データ欠け・超小型株・一時的利益)。候補を広げるための道具であり、最終20社の選定には使わない。")
verdict("よって、合成点の上位から候補リストを作る。ただし、この点数で最終選定はしない。")
banner3("手順3　順位の付け方を変えても結果は同じか ― 頑健さの監査")
body("順位化のやり方は一つではない。市場全体での順位・業種の中での順位・外れ値に強い標準化など複数のやり方で変わる堀の順位を付け直し、選ばれる顔ぶれが入れ替わらないかを確かめた。付け方に敏感な会社には確認フラグを付け、離の個別確認へ引き継いだ(詳細は技術補遺§C・§E)。重みの大きさを予測力と読んではいけない。", after=2)
banner3("式(10)の導出 ― 変わる堀の定義を、測れる四つの要素に分ける")
body("式(9)の合成点を、変わる堀に合わせて具体化したものが式(10)である(破の選定に使う独自の物差し)。ここから式(10)〜(12)は私たちの設計だが、導出は三つの原則に従う――第一に、部品は先行研究と前掲の順位点の再利用。第二に、重みは成績で決めない(過去リターンへの当てはめは後出しの罠)。第三に、点数(魅力の大きさ)と証拠(裏づけの強さ)を混ぜない。変わる堀の定義『割安に放置されているが、資本効率改革の実行でこれから評価が変わる会社』を、四つの測れる要素に分ける――(a)割安ギャップＶ(いま割安か＝純資産比・益回りの順位点)、(b)資本効率の改善Ｃ(ＲＯＥ・利益率が良くなっているか)、(c)株主還元の強化Ｒ(配当・自社株買いの実行)、(d)改革シグナルＫ(中期計画・ＰＢＲ改善策の開示)。これらを重みつきで束ね、『安いだけの会社』を弾く減点を差し引く。", after=2)
formula("10", "変わる堀の点数(破の選定に使用・私たちの設計)",
        r"S^{\mathrm{Trans}}_i = w_V V_i + w_C C_i + w_R R_i + w_K K_i - P^{\mathrm{Trap}}_i",
        "Ｖ＝割安ギャップ(純資産比・益回りの順位点)　Ｃ＝資本効率の改善　Ｒ＝株主還元の強化　Ｋ＝改革シグナル(開示)　Ｐ＝割安のワナ減点。改善の実行(Ｃ・Ｒ・Ｋ)を割安(Ｖ)より重くし、決めた後に各重みを±20％動かしても顔ぶれが変わらないことを監査した(全重みと監査は技術補遺§C)。")
verdict("よって、割安さに「変革の実行」が重なる会社ほど高得点とし、安いだけの会社は減点で弾く。")
_ha = F9["ha"]["steps"]
_hnames = "、".join(f"{d['code']} {d['name_ja'].replace('株式会社', '')}" for d in sorted([x for x in DATA.values() if x["role"] == "Transformation Core"], key=lambda z: -z["w"]))
body(f"結果: 適格{F9['n_base']:,}社 →(2-1 変わる堀に分類){_ha[0]['n']}社 →(2-2 黒字){_ha[1]['n']}社 →(2-3 ＲＯＥ≧5％){_ha[2]['n']}社 →(2-4 変わる堀の点数の上位・同一業種2社まで・他の役割との重複を除く)5社。選ばれたのは、{_hnames}である。")
fig("fig2_ha_v9.png", f"図表Ⅱ-7　スクリーニング2(破)の絞り込み ― 適格{F9['n_base']:,}社から変わる堀5社へ",
    "注：棒の長さは残った会社数(対数目盛)。変わる堀の点数(式10)は2-4の並べ替えにのみ使い、最終20社の配分には使わない。出所：funnel_branches_v9.json。")
_j = DATA["9022"]
body(f"通し例: 東海旅客鉄道(9022)。変わる堀に分類された{_ha[0]['n']}社の一社で、黒字(2-2)・ＲＯＥ≧5％(2-3)を通過し、変わる堀の点数の順位で上位に入った(2-4)。純資産比{_j['bm']:.2f}倍の割安に、数字で確認できる改善の証拠(証拠の強さ 水準{_j['evid']}・式12で後述)が重なる――「安いだけ」ではない変わる堀の典型である。")
body("破の限界も先に書く。時間を遡った完全な将来予測の検証(学習期間と検証期間の厳密な分割)は、使える年度数の制約で完了していない。したがってこの候補づくりに将来リターンを予測する力があるとは主張しない。割安×変革の候補を広く・壊れにくく作ること――それが破の役割である。", after=2)
verdict("受け渡し: 変わる堀5社が確定した。しかし破の物差しは会計数値に表れる変革しか測れない――次の離は、数値になる前の「生まれる堀」を測りに行く。")

step_chip(3, f"離 ― 生まれる堀を、点数でなく事業の証拠で選ぶ({F9['n_base']:,}社→5社)")
body(f"ねらい: 生まれる堀(ＡＩ・半導体の実需が育てる未来の競争優位)は、会計数値にまだ表れない。study のキーワード型の点数(未来の堀の点数)は、適格{F9['n_base']:,}社のうち{F9['ri']['keyword_path']['fm_category_base']}社を未来分類に挙げる。しかし、この点数をそのまま選定に使うことはできなかった。理由は単純で、決定的である――全上場では{F9['ri']['keyword_path']['tie_n']}社が同じ点数で並び、火災報知機のホーチキ・時計のシチズン・鉄道信号の日本信号が、半導体マスク検査装置で世界的地位を持つレーザーテックと同点になる。社名・業種へのキーワード照合が生む飽和であり、点数では『本当にＡＩ・半導体の実需を持つ会社』を選別できない。だから、この経路は破棄した(3-1)。")
body("そこで離の選定は、点数の並べ替えから証拠の確認へ切り替えた。候補を精査し、事業セグメントの開示まで遡ってＡＩ・半導体・光通信の実需(製品・顧客・投資)を確認できた会社だけを通す(3-2)。そのうえで、全枝共通の適格ガード――黒字・最低限の収益性ＲＯＥ≧5％・流動性・履歴3年――を課す(3-3)。疑わしきは除外する保守設計である。まず点数の設計(式11)と証拠の関所(式12)を示し、次に結果を示す。")
banner3("式(11)の導出 ― 生まれる堀は「堀の源泉」を未来方向へ対応づける")
body("第Ⅰ章でバフェットの堀を、ブランドや特許などの無形の強み・乗り換えの手間・利用者が増えるほど価値が増す仕組み・圧倒的なコスト優位、と分解した。生まれる堀は「これらがこれから形成される力」であり、会計数値に表れにくい。そこで、無形資産が企業価値の主役になったことを示す文献(Lev and Gu, 2016；Peters and Taylor, 2017)と、研究開発投資が株式リターンに結びつくことを示した実証(Chan, Lakonishok and Sougiannis, 2001)に沿って、堀の源泉を六つの測れる部品に対応づけた――無形資産Ｉ・技術力Ｎ・急所を握る度合いＢ(価格決定力)・ＡＩ基盤への実需接続Ａ・データ顧客基盤Ｄ・信頼と安全Ｔである。", after=2)
formula("11", "生まれる堀の点数(離の評価に使用・私たちの設計)",
        r"S^{\mathrm{Emerg}}_i = 100\,(w_I I_i + w_N N_i + w_B B_i + w_A A_i + w_D D_i + w_T T_i) + B^{\mathrm{Evi}}_i - P^{\mathrm{Hype}}_i",
        "Ｉ＝無形資産　Ｎ＝技術力　Ｂ＝急所を握る度合い(価格決定力)　Ａ＝ＡＩ基盤への実需接続　Ｄ＝データ・顧客基盤　Ｔ＝信頼・安全　w＝各重み(ＡＩ接続Ａを最重視)　加点Ｂ^Evi＝証拠水準{1,2,3}を{2,4,8}点に対応づけ　減点Ｐ^Hype＝キーワードの接点のみ。")
body(f"重みは、中心仮説「ＡＩ基盤への実需接続が新しい堀を生む」からＡ(ＡＩ接続)を最重視し、堀の古典的源泉(無形資産・急所を握る度合い)と技術力が続く――過去リターンでは決めない。加点は証拠水準が1段上がるごとに2点→4点→8点と倍にした(水準3=数字まで確認はそれだけ希少)。ただし正直に書く――この点数の土台になる開示データのキーワード照合が前述のとおり飽和する({F9['ri']['keyword_path']['tie_n']}社が同点)ため、式(11)は選定の決定打にはならない。役割は二つに限る: 候補の点検・比較の物差しと、キーワードの接点しかない会社への減点(Ｐ^Hype)である。全重みは技術補遺§Bに開示した。", after=2)
verdict("よって、生まれる堀の点数は評価の物差しにとどめ、選定の決定は次の証拠の関所(式12)と事業検証(3-2)に委ねる。")
banner3("式(12)の導出 ― 点数と証拠を分ける関所")
body("最後の式(12)は、点数が高くても証拠が弱い会社を弾く「関所」である。なぜ点数と証拠を分けるのか。点数は設計の裁量が入る連続値だが、証拠は開示資料まで遡って誰でも確認できる段階値だからである。証拠は「キーワードの接点のみ(水準1)→製品・顧客・投資計画の具体性(水準2)→売上・受注・投資額の数字(水準3)」という検証可能性の階段で順序づけた。役割ごとの判定論理も定義から導いた: 両立型は「変わる」と「生まれる」の二つの物語が両方必要だから弱い方(min)で判定し、変わる堀は改革の経路が複数ありどれか一つ確認できれば足りるから強い方(max)で判定する。", after=2)
formula("12", "証拠の強さ(点数と分けて管理・私たちの設計)",
        r"L^{\mathrm{final}}_i = \begin{cases} \min(L^{TQ}_i, L^{EM}_i) & \text{両立型} \\ L^{EM}_i & \text{生まれる堀} \\ \max(L^{TQ}_i, L^{TS}_i, L^{TR}_i) & \text{変わる堀} \\ \max(L^{TQ}_i, L^{EM}_i) & \text{その他} \end{cases}",
        "ＴＱ＝数字の改善で確認できる変わる証拠　ＴＳ＝株主還元の実行の証拠　ＴＲ＝改革の開示(計画文書)の証拠　ＥＭ＝新テーマ(生まれる堀)の開示証拠。読み方: min＝両方の証拠が必要(弱い方で判定)、max＝どれか一つ強ければ十分。水準1＝キーワードの接点のみ／水準2＝製品・顧客・投資計画の具体性を確認／水準3＝売上・受注・投資額の数字まで確認。なおＴＳ・ＴＲは今回のデータでは全社が同値となり判定に効いておらず、変わる堀の水準は実質ＴＱで決まる(技術補遺§B)。")
_e3 = sum(1 for v in DATA.values() if v.get("evid") == 3); _e2 = sum(1 for v in DATA.values() if v.get("evid") == 2)
verdict(f"よって、点数がどれほど高くても、証拠の水準が足りない会社は役割の候補から外す。最終20社の証拠内訳は、事業で実需を確認できた水準3が{_e3}社・水準2が{_e2}社である(守の完成した堀は品質基準で選び、証拠水準は破・離の判定に用いる)。")
_ri = F9["ri"]["verified_path"]
_rnames = "、".join(f"{d['code']} {d['name_ja'].replace('株式会社', '').replace('　Ｈｏｌｄｉｎｇｓ', 'ＨＤ')}" for d in sorted([x for x in DATA.values() if x["role"] == "Emerging Core"], key=lambda z: -z["w"]))
body(f"結果: 候補を精査し、事業セグメントの開示でＡＩ・半導体・光通信の実需を確認できたのは{_ri[0]['n']}社(うち予備2社)。何をどの開示で確認したかを検証台帳として残す(下表)――誰でも同じ資料で追試できる。適格ガード(3-3)を通過した5社――{_rnames}――を生まれる堀とした。")
table(["コード・社名", "開示で確認した実需(セグメント・製品)", "採否"],
      [["6777 santecホールディングス", "ＡＩデータセンター向け波長可変レーザ・光通信部品", "採用"],
       ["6871 日本マイクロニクス", "半導体プローブカード(先端ウエハテスト)", "採用"],
       ["6590 芝浦メカトロニクス", "半導体・ＦＰＤ製造装置(洗浄・成膜)", "採用"],
       ["6387 サムコ", "化合物半導体(GaN/SiC)製造装置", "採用"],
       ["6627 テラプローブ", "半導体ウエハテスト受託", "採用"],
       ["6951 日本電子", "電子顕微鏡・電子線計測(半導体計測)", "予備"],
       ["6941 山一電機", "半導体テストソケット", "予備"]],
      [52, 86, 18], fs=10)
para(f"図表Ⅱ-8　離の検証台帳 ― 実需を確認できた{_ri[0]['n']}社(予備2社を含む)。確認先は各社の事業セグメント開示(有価証券報告書・決算説明資料)。", MIN, 10, after=5)
fig("fig2_ri_v9.png", f"図表Ⅱ-9　スクリーニング3(離)の分岐 ― キーワード経路を破棄し、事業の証拠で5社へ",
    "注：破棄した経路(3-1)も隠さず描く。事業検証(3-2)は開示資料に基づき、誰でも同じ資料で追試できる。出所：funnel_branches_v9.json。")
_s = DATA["6777"]
body(f"通し例: ｓａｎｔｅｃ ＨＤ(6777)。事業セグメントの開示でＡＩデータセンター向け光通信部品の実需(製品・顧客)を確認できた{_ri[0]['n']}社の一社(3-2)。黒字・ＲＯＥ≧5％・流動性・履歴3年の適格ガードを通過(3-3)。未来の堀の偏差値{_s['fmoat_hensa']}・証拠の強さ 水準{_s['evid']}(式12の階段で判定)。", after=6)

step_chip(4, f"両立型・分散役 ― 支え役を選ぶ(適格プール{F9['dual']['pool_n']:,}社→3社＋2社)")
body(f"ねらい: 三本の枝(守・破・離)は、それぞれの世代に特化するがゆえに隙間を残す。両立型は「変わる」と「生まれる」の橋渡し役、分散役は業種・テーマの偏りの調整役である。主役ではなく支え役なので、枠は3社＋2社と少数にした。")
body(f"両立型(4-1): 黒字×ＲＯＥ≧5％の適格プール{F9['dual']['pool_n']:,}社から、現在の堀と未来の堀の両方の順位がともに上位の3社を選ぶ(同一業種2社まで)。判定は式(12)の弱い方(min)――当てはめ例のキーエンス(6861)は「変わる証拠(高収益の持続)」と「新テーマの証拠(工場自動化・ＡＩの実需)」の両方を事業の開示で確認できて初めて、両立型として通す。分散役(4-2): 選定済みの銘柄が使っていない業種に限り、総合点の上位2社を選ぶ({F9['bridge']['pool_n']:,}社→2社)。")
_dnames = "、".join(f"{d['code']} {d['name_ja'].replace('株式会社', '')}" for d in sorted([x for x in DATA.values() if x["role"] == "Dual Moat"], key=lambda z: -z["w"]))
_bnames = "、".join(f"{d['code']} {d['name_ja'].replace('株式会社', '')}" for d in sorted([x for x in DATA.values() if x["role"] == "Bridge / Diversifier"], key=lambda z: -z["w"]))
body(f"結果: 両立型は{_dnames}の3社、分散役は{_bnames}の2社である。", after=6)

banner2("合流 ― 5＋5＋5＋3＋2＝20社。なぜこの構成か")
body("最終20社は、完成した堀5社(守の5社を固定)・変わる堀5社・生まれる堀5社・両立型3社・分散役2社で構成した。この構成は第Ⅰ章の「超える」の定義から導いている。")
bullet("完成した堀5社を持つ理由", "条件①(同じ土俵)の実行である。この5社を持たなければ、「何を超えたのか」を測る比較相手がいなくなる。しかも三世代のうち、完成した堀だけが仮説ではなく現在の事実である。")
bullet("5・5・5の均等", "三つの世代のどれか一つの未来に賭けない、という設計原理をそのまま数にした。")
bullet("3・2の少数枠", "両立型は変わる堀と生まれる堀の橋渡し、分散役は業種・テーマの偏りの調整役。主役ではなく脇役なので少数にした。")
body(f"「バフェットが選ぶはずの5社は、本当に必要なのか」という問いには、二つで答える。第一に設計上の理由――この5社を持たなければ、「何を超えたのか」を測る比較相手(新バフェット)と同じ土俵に立てない。しかも三世代のうち、完成した堀だけが仮説ではなく現在の事実である。第二に頑健性――守の品質ゲートの判定の目安・順位・業種上限を{ROB['n_variants']}通りに動かしても、完成した堀5社は最小でも{ROB['min_overlap']}社が一致し(ＺＯＺＯ・日本オラクル・名村造船・サンリオは全通りで不変)、恣意的な選定でないことが確かめられる(技術補遺§E)。なお、この構成が唯一の最適だという証明はできない。最適かどうかは未来の成績でしか分からないからだ。私たちが保証するのは、定義との整合と、次に示す壊れにくさである。")
fig("fig2_ri_v7.png", "図表Ⅱ-10　合流 ― 三世代の堀から最終20社へ(矢印は1対1の対応)",
    "注：守の5社は証拠の関所を通らず固定。両立型は二つの点数の両立を、証拠は厳しい方で判定する。筆者作成。")

banner3("選定は壊れにくいか ― 条件③の検証(要約)")
body(f"「複雑なルールを後から調整して、結論ありきで20社を選んだのではないか」という疑いには、三通りの頑健性検査で答えた。第一に、守の品質ゲートの判定の目安・順位・業種上限を{ROB['n_variants']}通り動かしても、完成した堀5社は最小でも{ROB['min_overlap']}社が一致し、核となる4社(ＺＯＺＯ・日本オラクル・名村造船・サンリオ)は全通りで不変だった。第二に、配分の重みを均等・役割予算・最小分散のどれにしても、新バフェット型を上回る結論は変わらない。第三に、利上げ相場からＡＩ相場まで四つの相場局面すべてで超過が一貫した(第Ⅲ章の多期間検証)。事前に宣言した判定の目安(条件を一つ変えても大半が一致)を満たし、条件③は達成である。全表は技術補遺§E。", after=2)
point("本章の要点: 守で完成した堀5社(条件①)、破・離で死角を補う15社(条件②)を選び、判定の目安・重み・相場の三通りの頑健性検査で壊れにくさ(条件③)を確かめた(全表は技術補遺§E)。")

# ================= Ⅲ ポートフォリオ =================
banner1("Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス", newpage=True)
body("前章で最終20社が決まった。本章では500万円の配分を決め、20社の顔ぶれを紹介し、最後に第Ⅰ章の条件④を対照群との比較で検証する。")
banner2("1．500万円の配り方 ― 成績の予想を使わない配分")
body("どの銘柄が上がるかという予想は、一切使わない。予想を使えば、せっかく証拠で選んだ20社に、予想の当て推量が混ざってしまうからだ。配分は次の四段階で決めた。")
bullet("役割予算(第一・第二段階)", "20社を五つの役割に分け(前章)、役割ごとに予算を決める。完成・変化・新生の三世代に各28％、両立型10％、分散役6％――三世代へ同額を置くのは、どの未来にも賭けないという設計の続きである。")
bullet("役割内の配分と執行(第三・第四段階)", "同じ役割の中では、値動きが穏やかで・売買しやすく・証拠が強く・データが信頼できる会社を厚くし(式13)、1銘柄の上限をかけ、1株から買える単元未満株(金額指定)で目標比率どおりに配分する(式14)。")
formula("13", "役割予算つきの配分 ― 選定済み20社への配り方(Markowitz, 1952)",
        r"\omega_i = B_{r(i)} \cdot \frac{\rho_i}{\sum_{j:\,r(j)=r(i)} \rho_j}, \qquad \rho_i = \frac{\ell_i\, e_i\, c_i}{\max(\sigma_i,\,0.10)}",
        "Ｂ＝役割ごとの予算(28/28/28/10/6％)　ℓ＝売買のしやすさ　ｅ＝証拠の強さ　ｃ＝データの信頼度　σ＝1年の値動きの大きさ。上限を超えた分は同じ役割の中で配り直す。")
verdict("よって、同じ役割の中では「値動きが穏やかで・売買しやすく・証拠が強い」会社ほど厚く持つ。")
formula("14", "単元未満株での配分(日本取引所グループ「単元株制度」)",
        r"q_i = \frac{B\,\omega_i}{P_i}",
        f"ｑ＝購入株数(単元未満株=1株未満の端株も可)　Ｂ＝総予算500万円　ω＝目標比率(式13)　Ｐ＝株価。大型優良は株価が高く通常の売買単位(100株)では収まらないため、金額指定の単元未満株を前提とする。")
verdict(f"よって、目標比率どおりに金額を割り当て、総予算500万円のうち{INVEST:,}円(使用率{INVEST/5_000_000:.1%})を投資に充てた。")
wmax = max(DATA.values(), key=lambda z: z["w"])
body(f"最終配分と20社の顔ぶれを、図表Ⅲ-1に一覧で示す（配分・事業概要・堀の偏差値を一表に集約）。役割の予算は守28／破28／離28／両立10／分散6％、最大の保有は{wmax['name_ja']}の{wmax['w']*100:.2f}％。堀の偏差値は全上場企業の中での位置を平均50で表した値で、現在の堀＝いまの競争優位、未来の堀＝これから生まれる競争優位の強さを示す。")
prows = []
for role in order:
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        _nm = d['name_ja'].replace("株式会社", "").replace("ホールディングス", "ＨＤ").replace("　Ｈｏｌｄｉｎｇｓ", "ＨＤ")
        _rj = {"Buffett Core": "守", "Transformation Core": "破", "Emerging Core": "離",
               "Dual Moat": "両立", "Bridge / Diversifier": "分散"}[role]
        prows.append([f"{d['code']} {_nm}", _rj,
                      f"{d['w']*100:.1f}％ {d['amtL1']:,}", d.get("business", ""),
                      f"{d['moat_hensa']}", f"{d['fmoat_hensa']}"])
table(["企業（コード・社名）", "役割", "比率・金額(円)", "事業概要（要ＩＲ確認）", "現在\nの堀", "未来\nの堀"], prows,
      [34, 10, 27, 77, 13, 13], fs=10, fixed=True)
para("図表Ⅲ-1　最終ポートフォリオ20社（総予算500万円・単元未満株の目標配分）。役割: 守＝完成した堀／破＝変わる堀／離＝生まれる堀。堀の偏差値＝全上場の中での位置(平均50)。出所：portfolio_v7.json。", MIN, 10, after=6)

banner2("2．銘柄紹介 ― 三世代の堀、20社の顔ぶれ")
body("各社の「選定データが語る事実」は、第Ⅱ章の点数と証拠から自動的に書ける。一方、事業の中身と堀の背景は、各社のＩＲ資料・有価証券報告書で確かめてから書くべきものなので、確認欄として残した(創作はしない)。")


role_desc = {
    "Buffett Core": "守の式だけで選んだ基準線",
    "Transformation Core": "割安×改善の証拠で選んだ変わる会社",
    "Emerging Core": "ＡＩ基盤への実需接続を事業で確かめた会社",
    "Dual Moat": "変わる×生まれるの両立",
    "Bridge / Diversifier": "業種・テーマの偏りの調整役",
}
for role in order:
    _names = "／".join(f"{d['code']} {d['name_ja'].replace('株式会社', '').replace('ホールディングス', 'ＨＤ').replace('　Ｈｏｌｄｉｎｇｓ', 'ＨＤ')}"
                      for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]))
    bullet(f"{role_jp[role]}（{role_desc[role]}）", _names)
body("配分・事業概要・堀の偏差値は図表Ⅲ-1に集約した。各社の証拠水準と役割別スコアの詳細は技術補遺§B。", after=6)

banner2("3．パフォーマンス検証 ― 新バフェット型・市場と比べる")
body("本節では第Ⅰ章の条件④を検証する。物差しは分析の前に宣言する――①年率リターン ②ＴＯＰＩＸ超過 ③市場との連動度(β) ④最大の下落 ⑤市場超過の安定度(ＩＲ)の5つ、期間は「選定に使った3年」と「直近1年」。直近1年も選定の点数計算には使っていないが、配分の値動き係数(式13のσ)は同じ1年の変動を入力に使うため、完全な時点外検証ではない(技術補遺§G)。")
body("比較相手は二つ置く。主対照は『新バフェット型』――同じ出発点から、バフェットの品質基準(高収益・堀・財務健全)だけで選んだ大型優良の上位12社(時価総額加重)である。参考として『純正グレアム型』――割安一辺倒(簿価割れ)の式だけで選んだ20社(等金額)も併記する。さらに市場平均としてＴＯＰＩＸ・日経平均を並べる。いずれも私たちが同じ規律で、後から入れ替えずに計算した。")
body("位置づけに先に限定をつける。両対照とも『式だけの機械選定』であり、バフェット本人の運用(少数集中・保険フロートによる調達・数十年保有)ではない。この比較は「本人に勝った」ことを意味せず、『同じ土俵で式だけを使った場合に、三世代の設計が何を足したか』を測る基準線である。そしてすべては2026年時点の選定を過去に当てた自己検証であり、将来の勝利の証明ではない。")
fig("cum_v7.png", "図表Ⅲ-2　累積リターンの比較(過去3年・期首=1)",
    "注：本ＰＦ・新バフェット型・純正グレアム型・ＴＯＰＩＸ・日経平均。ＴＯＰＩＸ連動指数1306は未調整の株式分割(2026-03-30)を×10補正、指数系列の欠測を前営業日で補完。価格はyfinance調整後終値(配当込み。日経のみ価格指数)。すべて2026年時点選定の自己検証。出所：control_comparison_v7.json・価格正典データ。", maxw=96)
table(["物差し(3年／直近1年)", "本ＰＦ(離)", "新バフェット型(主対照)", "ＴＯＰＩＸ", "日経平均"],
      [["年率リターン", f"{O3['ann_return']:.1%}／{O1['ann_return']:.1%}", f"{C3['ann_return']:.1%}／{C1['ann_return']:.1%}", f"{O3['topix_ann_return']:.1%}／{O1['topix_ann_return']:.1%}", f"{O3['nikkei_ann_return']:.1%}／{O1['nikkei_ann_return']:.1%}"],
       ["ＴＯＰＩＸ超過", f"＋{O3['excess_vs_topix']*100:.1f}／＋{G1:.1f}pt", f"{C3['excess_vs_topix']*100:.1f}／{C1['excess_vs_topix']*100:.1f}pt", "―", "―"],
       ["市場との連動度β", f"{O3['beta_vs_topix']:.2f}／{O1['beta_vs_topix']:.2f}", f"{C3['beta_vs_topix']:.2f}／{C1['beta_vs_topix']:.2f}", "1.00", "―"],
       ["最大の下落", f"{O3['max_drawdown']:.1%}／{O1['max_drawdown']:.1%}", f"{C3['max_drawdown']:.1%}／{C1['max_drawdown']:.1%}", f"{O3['topix_max_drawdown']:.1%}／{O1['topix_max_drawdown']:.1%}", "―"],
       ["超過の安定度ＩＲ", f"＋{O3['information_ratio']:.2f}／＋{O1['information_ratio']:.2f}", f"{C3['information_ratio']:.2f}／{C1['information_ratio']:.2f}", "―", "―"]],
      [38, 36, 38, 34, 24], fs=10)
para("図表Ⅲ-3　対照群・参照ベンチマークとの比較(5つの物差し、各セル＝3年／直近1年)。日経平均は値がさハイテク株の比重が高い参照値(配当抜き)。出所：control_comparison_v7.json。", MIN, 10, after=6)
body(f"読み取りは三つ。第一に、本ＰＦは新バフェット型を3年で年{D3:.1f}ポイント・直近1年で年{D1:.1f}ポイント上回り、グレアム型には3年で年{DG3:.1f}ポイント上回った。ただし新バフェット型への超過は統計的には確実でない――日次超過のNewey-Westｔ値は3年{SIG['3y']['ours_vs_buffett']['t_newey_west']:.2f}で有意水準に届かず、「互角〜やや上」と評価する。有意に上回るのはグレアム型(ｔ＝{SIG['3y']['ours_vs_graham']['t_newey_west']:.2f})と市場ＴＯＰＩＸ(ｔ＝{SIG['3y']['ours_vs_topix']['t_newey_west']:.2f})である。第二に、守り――最大の下落({O3['max_drawdown']:.1%} 対 {C3['max_drawdown']:.1%})と連動度β({O3['beta_vs_topix']:.2f} 対 {C3['beta_vs_topix']:.2f})――でも新バフェット型を上回った。三世代分散が攻めと守りの両面で効いている。第三に、これはすべて自己検証であり、将来の市場超過の証明ではない――この限界は言い換えずに残す。")
body(f"役割の分担も確かめた。過去3年の値上がり寄与は、生まれる堀(ＡＩ・半導体)が約{ROLE_CONTRIB['by_role_pct'].get('離 生まれる堀',0):.0f}％と牽引し、完成した堀(守)も約{ROLE_CONTRIB['by_role_pct'].get('守 完成した堀',0):.0f}％と土台以上に寄与、変わる堀(公益・資源)は約{ROLE_CONTRIB['by_role_pct'].get('破 変わる堀',0):.0f}％で下げ相場の緩衝材として効いた。旧来の割安一辺倒(グレアム型)に比べ、三世代へ寄与が分散している(役割別の内訳表は技術補遺§E)。なお本検証は取引コスト・税を考慮せず、毎日同じ比率へ戻す固定重み規約で計算した。株数を固定して買い持ちにすると値動きが変わる(参考値は技術補遺§D)。")

banner2("4．結論と限界 ― 結局、バフェットを超えたのか")
body("第Ⅰ章で立てた投資仮説と四つの条件に、正面から答える。")
_evid3 = sum(1 for v in DATA.values() if v.get("evid") == 3)
_evid2 = sum(1 for v in DATA.values() if v.get("evid") == 2)
table(["「超える」の条件", "結果", "判定"],
      [["① 同じ土俵に立つ", "品質基準を式から再現し、その完成した堀5社を固定保有(第Ⅱ章)", "達成"],
       ["② 死角に届く", f"品質レンズが映さない変わる堀・生まれる堀の15社を、事業で検証した証拠つきで選定(証拠水準3=約{_evid3}社・水準2=約{_evid2}社)", "達成"],
       ["③ 壊れにくい", "目安・重み方式・相場局面の三通りで検査。守5は最小4/5一致、『超え』は不変(第Ⅱ章末)", "達成"],
       ["④ 総体で上回る", f"新バフェット型に3年＋{D3:.1f}pt・1年＋{D1:.1f}pt、グレアム型に3年＋{DG3:.1f}pt。有意に勝つのはグレアム型と市場、新バフェットは有意差なし＝互角〜やや上(図表Ⅲ-3)", "達成(有意差なし)"]],
      [30, 118, 22], fs=10)
para("図表Ⅲ-4　判定表 ― 四条件に対する現時点の答え(すべて2026年時点選定の自己検証)。", MIN, 10, after=6)
body(f"結局、超えたのか。答えは明快に分けて書く。割安一辺倒のグレアム型と市場平均(ＴＯＰＩＸ)には、リスクを抑えつつ統計的にも有意に上回った。師である新バフェット型に対しては、リターン・効率(シャープ)・最大の下落のすべてで、しかも利上げ相場からＡＩ相場まで四つの局面すべてで一貫して上回ったが、その差は統計的には確実とは言えない――正直に「互角〜やや上」と評価する。第Ⅰ章の投資仮説はこの範囲で支持された。最も大切な限界――すべて2026年時点の選定を過去に当てた自己検証であり、将来の約束ではない。胸を張れるのは成績の大きさではなく、三世代の堀という枠組みを事前規律(後から入れ替えない・事業で証拠を確かめる)で実行したことにある。")

body(f"最後に、この設計の弱点を私たち自身の手で記録しておく――強みだけを並べたレポートは検証に耐えないからである。①テーマの偏り: 値上がり寄与の約{ROLE_CONTRIB['by_role_pct'].get('離 生まれる堀',0):.0f}%は生まれる堀(ＡＩ・半導体)に集中し、テーマの重みは約40%(業種HHI0.16)。『ＡＩ・半導体バリューチェーンへの意図的な賭け』として隠さず開示する。テーマが崩れる未来では市場に劣りうるため、守・破に予算を置いて緩衝している。②単元未満株の前提: 株価の高い{N_UNBUY}社は通常の売買単位(100株)では¥500万に収まらず、本配分は単元未満株(金額指定)の取扱いを前提とした目標ウェイトである。③証拠は開示ベース: 現場の実態までは測れず、第Ⅳ章の取材はこの限界を埋めるために設計した。④出発点は生存者のみ: 出発点の{F9['n_nonfin']:,}社は期間中の上場廃止企業を含まず(生存者バイアス)、成績はその分有利に出うる(技術補遺§A)。")
point("グレアム型と市場には有意に超え、新バフェットには全指標・全局面で上回るが有意差なし(互角〜やや上)。テーマ集中とデータの限界まで先に書き残した。")

# ================= Ⅳ インタビュー =================
banner1("Ⅳ．インタビュー・アンケート", newpage=True)
body("前章までの選定と検証は、すべて開示データの上に立っている。本章では、その見立てが現場の実態と合っているかを確かめる取材の計画を示す。質問は思いつきではなく、各社を選んだ理由(点数と証拠)から導いた(実施記録は記入欄)。")
table(["対象(役割)", "選定理由から導いた質問草案"],
      [["6777 santec(生まれる堀)", "①ＡＩデータセンター向け光通信部品の受注動向と生産能力投資　②ＡＩ需要が一巡した場合の下支え事業"],
       ["6590 芝浦メカトロニクス(生まれる堀)", "①半導体製造装置の需要見通し　②競合に対する技術優位の源泉"],
       ["6920 レーザーテック(完成した堀)", "①半導体マスク検査での独占的地位の持続性　②研究開発投資の方向"],
       ["3092 ＺＯＺＯ(完成した堀)", "①出店ブランドと購買データの両面ネットワークの堀　②海外・新規事業の成長余地"],
       ["1662 石油資源開発(変わる堀)", "①資源価格の変動への耐性と株主還元の方針　②ＣＣＳ(ＣＯ２貯留)など変革の進捗"],
       ["9503 関西電力(変わる堀)", "①原発再稼働と資本効率改革の進捗　②株主還元と設備投資の優先順位"],
       ["6861 キーエンス(両立型)", "①高い営業利益率(価格支配力)の持続性　②工場自動化・ＡＩ活用による事業機会"],
       ["3449 テクノフレックス(分散役)", "①ニッチ配管部材の参入障壁　②財務の安定性と成長計画"]],
      [52, 118], fs=10)
banner3("実施記録(記入・取材ごとに1列)")
table(["項目", "取材1(記入)", "取材2(記入)"],
      [["企業名・対象者(名字)・部署", "［記入］", "［記入］"],
       ["日時・方法(対面/オンライン/書面)", "［記入］", "［記入］"],
       ["質問と回答の要点", "［記入：質問草案の番号と対応させる］", "［記入］"],
       ["写真(オンライン画面も可)", "［写真貼付欄］\n\n\n", "［写真貼付欄］\n\n\n"],
       ["分析への反映", "［記入：証拠水準・役割の見立てをどう見直したか］", "［記入］"]],
      [38, 66, 66], fs=10)
banner3("取材の還流(総括・記入)")
para("書き方の型：「取材の結果、○○社の△△が確認でき、証拠水準の見立てを□□と見直した／仮説どおりだった。第Ⅱ章の××の判断は現場の実態と整合していた」――分析へ何が返ってきたかを必ず書く。", MIN, 10, after=2)
fill_line("総括(3〜5文)")
para("(未実施の場合は本章を「今後の課題」とし、その旨を明記する。回答の創作はしない。)", MIN, 10, after=6)

# ================= Ⅴ 学んだこと =================
banner1("Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと")
for lab, hint in [("直面した困難と乗り越え方", "例の問い：仮説はどこで崩れたか。候補を広げすぎて確認が回らなくなったとき、どう折り合いをつけたか"),
                  ("インタビューと結びついた学び", "例の問い：取材で確かめて初めて分かったことは何か。開示データだけでは見えなかったものは何か"),
                  ("チームでの役割分担と今後の課題", "例の問い：互いの誤りをどう見つけ合ったか。単元未満株の扱い・証拠確認の自動化など残った宿題"),
                  ("謝辞", "取材先・助言をくれた方・支えてくれた人へ(個人は名字のみ)")]:
    fill_line(lab, hint)
para("(提出者の実体験として記入する。下書きの創作はしない。)", MIN, 10, after=6)

# ================= 参考文献 =================
banner1("参考文献")
body("本文で直接引用した文献に加え、選定・検証・データ取得の設計にあたって参照した文献を、分野別にすべて記載する(正典: docs/references_master.md)。新バフェットの品質基準はBuffett's Alpha(Frazzini, Kabiller and Pedersen 2018)とQuality Minus Junk(Asness, Frazzini and Pedersen 2019)に、割安×優良の複合順位はGreenblatt(2006)に依拠する。", after=5)


def ref_para(r):
    p = doc.add_paragraph(); pp = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360"); pp.append(ind)
    p.paragraph_format.space_after = Pt(1)
    X.setfont(p.add_run(r), MIN, 10, False, BLACK)


REF_GROUPS = [
 ("一次資料 ― バフェット株主への手紙(原文)", [
  "Buffett, W. E. (1988) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/1988.html (2026年7月18日)",
  "Buffett, W. E. (1989) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/1989.html (2026年7月18日)",
  "Buffett, W. E. (2007) Chairman's Letter to the Shareholders of Berkshire Hathaway Inc. Berkshire Hathaway Inc., https://www.berkshirehathaway.com/letters/2007ltr.pdf (2026年7月18日)",
 ]),
 ("英語文献Ⅰ ― 投資・会計・無形資産(第Ⅱ章 守・離の式典拠)", [
  "Altman, E. I. (1968) Financial Ratios, Discriminant Analysis and the Prediction of Corporate Bankruptcy. The Journal of Finance, 23(4), pp.589-609.",
  "Amihud, Y. (2002) Illiquidity and Stock Returns: Cross-Section and Time-Series Effects. Journal of Financial Markets, 5(1), pp.31-56.",
  "Asness, C. S., Frazzini, A., & Pedersen, L. H. (2019) Quality Minus Junk. Review of Accounting Studies, 24, pp.34-112.",
  "Basu, S. (1977) Investment Performance of Common Stocks in Relation to Their Price-Earnings Ratios: A Test of the Efficient Market Hypothesis. The Journal of Finance, 32(3), pp.663-682.",
  "Basu, S. (1983) The Relationship between Earnings' Yield, Market Value and Return for NYSE Common Stocks: Further Evidence. Journal of Financial Economics, 12(1), pp.129-156.",
  "Chan, L. K. C., Lakonishok, J., & Sougiannis, T. (2001) The Stock Market Valuation of Research and Development Expenditures. The Journal of Finance, 56(6), pp.2431-2456.",
  "Fama, E. F., & French, K. R. (1993) Common Risk Factors in the Returns on Stocks and Bonds. Journal of Financial Economics, 33(1), pp.3-56.",
  "Fama, E. F., & French, K. R. (2015) A Five-Factor Asset Pricing Model. Journal of Financial Economics, 116(1), pp.1-22.",
  "Frazzini, A., Kabiller, D., & Pedersen, L. H. (2018) Buffett's Alpha. Financial Analysts Journal, 74(4), pp.35-55.",
  "Jensen, M. C. (1968) The Performance of Mutual Funds in the Period 1945-1964. The Journal of Finance, 23(2), pp.389-416.",
  "Lev, B., & Gu, F. (2016) The End of Accounting and the Path Forward for Investors and Managers, John Wiley & Sons.",
  "Markowitz, H. (1952) Portfolio Selection. The Journal of Finance, 7(1), pp.77-91.",
  "Novy-Marx, R. (2013) The Other Side of Value: The Gross Profitability Premium. Journal of Financial Economics, 108(1), pp.1-28.",
  "Ohlson, J. A. (1980) Financial Ratios and the Probabilistic Prediction of Bankruptcy. Journal of Accounting Research, 18(1), pp.109-131.",
  "Peters, R. H., & Taylor, L. A. (2017) Intangible Capital and the Investment-q Relation. Journal of Financial Economics, 123(2), pp.251-272.",
  "Piotroski, J. D. (2000) Value Investing: The Use of Historical Financial Statement Information to Separate Winners from Losers. Journal of Accounting Research, 38, Supplement, pp.1-41.",
  "Sharpe, W. F. (1966) Mutual Fund Performance. The Journal of Business, 39(1), pp.119-138.",
  "Sharpe, W. F. (1994) The Sharpe Ratio. The Journal of Portfolio Management, 21(1), pp.49-58.",
  "Sloan, R. G. (1996) Do Stock Prices Fully Reflect Information in Accruals and Cash Flows about Future Earnings? The Accounting Review, 71(3), pp.289-315.",
 ]),
 ("英語文献Ⅱ ― 検証・過学習対策・感度分析(壊れにくさ検査の典拠)", [
  "Bailey, D. H., & López de Prado, M. (2012) The Sharpe Ratio Efficient Frontier. Journal of Risk, 15(2), pp.3-44.",
  "Bailey, D. H., Borwein, J., López de Prado, M., & Zhu, Q. J. (2014) Pseudo-Mathematics and Financial Charlatanism: The Effects of Backtest Overfitting on Out-of-Sample Performance. Notices of the American Mathematical Society, 61(5), pp.458-471.",
  "Bailey, D. H., & López de Prado, M. (2014) The Deflated Sharpe Ratio: Correcting for Selection Bias, Backtest Overfitting, and Non-Normality. The Journal of Portfolio Management, 40(5), pp.94-107.",
  "Bailey, D. H., Borwein, J. M., López de Prado, M., & Zhu, Q. J. (2016) The Probability of Backtest Overfitting. Journal of Computational Finance, 20(4), pp.39-69.",
  "Jaccard, P. (1901) Étude comparative de la distribution florale dans une portion des Alpes et des Jura. Bulletin de la Société Vaudoise des Sciences Naturelles, 37, pp.547-579.",
  "López de Prado, M. (2018) Advances in Financial Machine Learning, John Wiley & Sons.",
  "Romano, J. P., & Wolf, M. (2005) Stepwise Multiple Testing as Formalized Data Snooping. Econometrica, 73(4), pp.1237-1282.",
  "Saltelli, A., Ratto, M., Andres, T., Campolongo, F., Cariboni, J., Gatelli, D., Saisana, M., & Tarantola, S. (2008) Global Sensitivity Analysis: The Primer, John Wiley & Sons.",
  "Sobol', I. M. (2001) Global Sensitivity Indices for Nonlinear Mathematical Models and Their Monte Carlo Estimates. Mathematics and Computers in Simulation, 55(1-3), pp.271-280.",
  "White, H. (2000) A Reality Check for Data Snooping. Econometrica, 68(5), pp.1097-1126.",
 ]),
 ("背景調査で参照した文献", [
  "Bergstra, J., & Bengio, Y. (2012) Random Search for Hyper-Parameter Optimization. Journal of Machine Learning Research, 13, pp.281-305.",
  "Deb, K., Pratap, A., Agarwal, S., & Meyarivan, T. (2002) A Fast and Elitist Multiobjective Genetic Algorithm: NSGA-II. IEEE Transactions on Evolutionary Computation, 6(2), pp.182-197.",
  "Beneish, M. D. (1999) The Detection of Earnings Manipulation. Financial Analysts Journal, 55(5), pp.24-36.",
  "French, K. R. \"Data Library\", https://mba.tuck.dartmouth.edu/pages/faculty/ken.french/data_library.html (2026年7月8日)",
 ]),
 ("日本語文献", [
  "経済産業省(2014)『持続的成長への競争力とインセンティブ ―企業と投資家の望ましい関係構築―(伊藤レポート)』経済産業省．",
  "東京証券取引所(2023)「資本コストや株価を意識した経営の実現に向けた対応について」株式会社東京証券取引所．",
  "日本取引所グループ「単元株制度」，https://www.jpx.co.jp/ (2026年7月8日閲覧)",
 ]),
 ("データ・実装資料", [
  "金融庁「EDINET」，https://disclosure2.edinet-fsa.go.jp/ (2026年7月8日閲覧)",
  "金融庁 企画市場局 企業開示課(2026)『EDINET API 仕様書(Version 2)』金融庁．",
  "日本取引所グループ「上場会社情報」，https://www.jpx.co.jp/listing/co-search/ (2026年7月8日閲覧)",
  "日本取引所グループ「東証上場銘柄一覧」(データファイル)．",
  "SciPy Developers, \"scipy.optimize.minimize\", SciPy Documentation, https://docs.scipy.org/ (2026年7月8日)",
 ]),
]
for gtitle, grefs in REF_GROUPS:
    banner3(gtitle)
    for r in grefs:
        ref_para(r)

# (v6: 用語の手引きはユーザー指示で削除。専門用語は本文初出時の言い換えでカバー)

# ================= 保存+自己検査 =================
out = ED / f"beyond_buffett_stockleague_{VER}.docx"
doc.core_properties.comments = VER
doc.save(str(out))
json.dump({h: None for h, l in HEADINGS if l == 1}, open(ED / f"headings_{VER}.json", "w"), ensure_ascii=False)

# 検査1: 禁止語(平易化スイープ)
BANNED = ["候補宇宙", "合格ライン", "分位", "アクルーアル", "ガードレール", "正規化", "ロバスト",
          "レビュー可能性", "ボトルネック", "インサンプル", "in-sample", "母集団", "トラップ",
          "ヘッジ", "ユニバース", "アブレーション", "標本内", "ファクター"]
texts = []
for p in doc.paragraphs:
    texts.append(p.text)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            texts.append(c.text)
full = "\n".join(texts)
bad = [w for w in BANNED if w in full]

# 検査2: 全文字10pt以上(w:szは半ポイント単位)
small = set()
for el in doc.element.iter():
    if el.tag == qn("w:sz") or el.tag == qn("w:szCs"):
        v = el.get(qn("w:val"))
        if v and float(v) < 20:
            small.add(v)

print("saved", out.name, "| ver:", VER, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
print("BANNED-CHECK:", "PASS" if not bad else f"FAIL {bad}")
print("FONT>=10pt-CHECK:", "PASS" if not small else f"FAIL sz(half-pt)={sorted(small)}")
print(f"INVEST={INVEST:,} CASH={CASH:,}")
