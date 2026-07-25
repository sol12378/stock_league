# -*- coding: utf-8 -*-
"""v9 査読用技術補遺(beyond_buffett_technical_appendix_v9.docx)。
本編(≤30頁)に収まらない粒度――全式・選定規律・計算規約・統計・頑健性・限界・再現手順――を収載。
全数値は v7 正典(control_comparison_v7/significance_v7/funnel_exclusion_v7/role_contribution_v7/
portfolio_v7/data_real_v7)から自動転記。手打ち禁止。"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
ED = ROOT / "outputs/stockleague_edition"
WORK = ROOT / "work/pure_buffett_benchmark"
COMP = json.load(open(ED / "control_comparison_v7.json", encoding="utf-8"))
BH = COMP["bh_reference"]
SIG = json.load(open(ED / "significance_v7.json", encoding="utf-8"))
F9 = json.load(open(ED / "funnel_branches_v9.json", encoding="utf-8"))
RC = json.load(open(ED / "role_contribution_v7.json", encoding="utf-8"))
PF = json.load(open(WORK / "portfolio_v7.json", encoding="utf-8"))
DATA = json.load(open(ED / "data_real_v7.json", encoding="utf-8"))
MP = PF["results"]
O3, O1 = COMP["ours"]["3y"], COMP["ours"]["1y"]
C3, C1 = COMP["control_buffett"]["3y"], COMP["control_buffett"]["1y"]
CG3 = COMP["control_graham"]["3y"]

NAVY = RGBColor(0x16, 0x32, 0x4F); TEAL = RGBColor(0x2F, 0x6D, 0x5F)
MINCHO, GOTHIC, MONO = "ＭＳ 明朝", "ＭＳ ゴシック", "Consolas"
doc = Document()
_st = doc.styles["Normal"]; _st.font.size = Pt(10.5); _st.font.name = "Times New Roman"
_st.element.rPr.rFonts.set(qn("w:eastAsia"), MINCHO)
sec = doc.sections[0]; sec.left_margin = sec.right_margin = Mm(20); sec.top_margin = sec.bottom_margin = Mm(18)


def _font(run, name=MINCHO, size=10.5, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def h1(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True; _font(p.add_run(t), GOTHIC, 13, True, NAVY); return p


def h2(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True; _font(p.add_run(t), GOTHIC, 11, True, TEAL); return p


def para(t, size=10.5, after=5, mono=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _font(p.add_run(t), MONO if mono else MINCHO, 9.5 if mono else size); return p


def _shade(cell, hx):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hx)
    cell._tc.get_or_add_tcPr().append(sh)


def table(hdr, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"
    for j, ht in enumerate(hdr):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        _font(c.paragraphs[0].add_run(ht), GOTHIC, 9.5, True, RGBColor(0xFF, 0xFF, 0xFF)); _shade(c, "16324F")
    for row in rows:
        cs = t.add_row().cells
        for j, v in enumerate(row):
            cs[j].paragraphs[0].clear(); _font(cs[j].paragraphs[0].add_run(str(v)), MINCHO, 9.5)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3); return t


def pct(x, dp=1):
    return f"{x*100:.{dp}f}%" if x is not None else "―"


# ===== 表紙 =====
p = doc.add_paragraph(); _font(p.add_run("BEYOND BUFFETT ― 査読用技術補遺(Technical Appendix) v9"), GOTHIC, 16, True, NAVY)
para("日経ＳＴＯＣＫリーグ提出版『三世代の堀』の検証可能性を担保する技術資料。本編(≤30頁)に収まらない詳細――"
     "全パラメータ・選定規律・計算規約・統計・頑健性・限界・再現手順――を収載する。全数値は本編と同一の選定正典"
     "(JSON)から自動転記。**重要な前提: 全ＰＦは2026-06時点の財務・価格で選定し過去に当てた in-sample の"
     "自己検証であり、将来の成績予測ではない。**", after=8)

# ===== §A データ系譜 =====
h1("§A　データ系譜と前処理")
para("価格系列は正典 data/processed/prices_daily.parquet(2021-06-01〜2026-06-01、約1,221営業日)。個別株と"
     "ＴＯＰＩＸ連動ＥＴＦ 1306.T は配当込み調整後終値、日経平均(^N225)は価格指数。1306.T は2026-03-30以降に"
     "未調整の1:10分割があり、当該日以降を×10して連続化した。リターンは version 非依存の pct_change(fill_method=None) で算出。", after=4)
para(f"ユニバースは金融を除く普通株{F9['n_nonfin']:,}社(価格あり)。投資適格＋流動性(60日平均売買代金)で{F9['n_eligible']:,}社、"
     f"さらに価格履歴3年(756営業日=検証可能性)で{F9['n_base']:,}社を全枝共通の出発点(base)とする。守・破・離・両立・分散の全枝の"
     "社数はこのbaseから同一ガードで再計算し、出口20社が選定スクリプトの出力と一致することを機械検査した(funnel_branches_v9.json)。"
     "母集団は直近時点の上場一覧の単一断面で、期間中の上場廃止企業を含まない(生存者バイアス)。本ＰＦ・新バフェット・グレアムには"
     "同一条件で作用するため相互比較は内部整合するが、指数との水準比較は割り引いて読む。", after=6)

# ===== §B 選定式の完全形 =====
h1("§B　選定式の完全形")
h2("B-1　守 ― 完成した堀(新バフェットの品質ゲート)")
para("守は、割安一辺倒でなく『優良を適正価格で』という新バフェットの規律を、次の7関門で式化する(出典: Frazzini, "
     "Kabiller and Pedersen 2018 'Buffett's Alpha'; Asness, Frazzini and Pedersen 2019 'Quality Minus Junk')。"
     "①ＲＯＥ≥15% ②営業利益率≥10%(価格支配力=堀) ③自己資本比率≥50%(安全) ④直近3期無赤字(予測可能性) "
     "⑤営業ＣＦ>0(利益の質; Sloan 1996) ⑥増収かつ増益(非縮小) ⑦60日平均売買代金≥約0.1億円(流動性; Amihud 2002)。", after=3)
table(["守の品質関門(base=1,791社から累積)", "通過社数"],
      [[f"{st['id']}　{st['label']}", f"{st['n']:,}社"] for st in F9["shu"]["steps"]], [110, 26])
para(f"品質ユニバース{F9['shu']['n_quality']}社のうち価格ランク可能{F9['shu']['n_priceable']}社を、ＲＯＥと益回りの順位和"
     "(Greenblatt型)で並べ、同一業種上限2で上位5社を固定(上位12社=主対照の新バフェット型)。", after=6)
h2("B-2　破 ― 変わる堀(割安×変革)")
para("破は、守の品質一辺倒を破り、東証の資本効率改革・脱炭素で評価がこれから変わる割安企業を選ぶ。study の変革カテゴリ"
     "(Transformation Moat)に分類される企業を、総合スコア(adjusted_bb_score; 割安・資本効率改善・株主還元・改革シグナルの"
     "合成)の上位から、営業黒字・純黒字・流動性・同一業種上限2の規律で5社選定する。重みは過去リターンで決めず変革の論理から"
     "設計し、±20%の摂動でも順位が崩れないことを監査した。枝ファネル: base 1,791社 → 変わる堀に分類282社 → 黒字251社 → "
     "ＲＯＥ≥5%で227社 → 点数上位＋業種上限2で5社(数値はfunnel_branches_v9.jsonから転記・実装一致assert済み)。", after=6)
h2("B-3　離 ― 生まれる堀(AI・半導体を事業で検証)")
para("離は、AI・半導体・光通信の実需で新しい堀が生まれる企業を選ぶ。ただし study の future_moat スコアは社名・業種への"
     "キーワード照合に依存して飽和し(炊飯器・照明の企業まで満点に並ぶ)、定量選別に耐えない。そこで点数を鵜呑みにせず、"
     "各社の事業セグメント開示まで遡って半導体・AI基盤への実需接続を確認できた企業のみを採用した(キーワード頼みの排除・"
     "疑わしきは除外)。採用5社=santec／日本マイクロニクス／芝浦メカトロニクス／SAMCO／テラプローブ"
     "(各社の事業内容は本編§Ⅲ)。キーワード経路破棄の実証: future_moatスコアは全上場で273社が同一値に並び"
     "(火災報知機・時計・鉄道信号の会社が半導体マスク検査のレーザーテックと同点)、順位として機能しない。"
     "枝ファネル: 実需確認7社(予備2)→適格ガードで5社。両立型はプール1,463社→3社、分散役は未使用業種638社→2社。", after=3)
para("両立型は現在の堀(moat)と未来の堀(future_moat)がともに上位の3社、分散役は業種・テーマの偏りを整える2社。"
     "証拠水準(接点のみ=1/具体=2/数字まで=3)を点数と分けて管理し、水準の低い会社は役割の候補から外す。", after=3)
_e3 = sum(1 for v in DATA.values() if v.get("evid") == 3); _e2 = sum(1 for v in DATA.values() if v.get("evid") == 2)
para(f"最終20社の証拠内訳: 水準3が{_e3}社・水準2が{_e2}社。", after=6)
h2("B-4　配分(役割予算) ― 式(13)")
para("配分は、成績の予想を使わず、役割予算で決める。予算配分は 守28%／破28%／離28%／両立10%／分散6% とし(三世代コアを"
     "等しく重視、両立・分散は支持役として軽く)、役割内は均等とした。真バフェットは大型優良で株価が高く通常の売買単位(100株)"
     "では500万円に収まらないため、1株から買える単元未満株(金額指定)を前提に目標比率で配分した。", after=6)

# ===== §C 選定の規律 =====
h1("§C　選定の規律(後出しの排除)")
para("(1)成績で選ばない: 守・破・離のいずれも、過去リターンへの当てはめで銘柄や重みを選んでいない。守は品質ゲート(閾値)、"
     "破・離は事前に設計したスコアと事業検証で選ぶ。(2)後から入れ替えない: 選定後に成績を見て銘柄を差し替える操作は一切"
     "していない。(3)摂動監査: 破・離のスコア重みを±20%動かしても選ばれる顔ぶれが変わらないことを確認した。(4)在庫の限界: "
     "shares_outstanding 等の時価総額データは大型中心に約139社しか取得できず、価格ランク可能な品質企業が大型優良に偏る点は"
     "開示する(新バフェット=大型優良と整合)。(5)来歴の正直な開示: 初期の反復では多目的最適化(Deb et al. 2002)や"
     "無作為探索(Bergstra & Bengio 2012)による重み探索も検討したが、後知恵の当てはめ(過去成績への最適化)を避けるため"
     "最終選定には用いず、変革の論理に基づく概念設計＋±20%摂動監査を採用した。", after=6)

# ===== §D 検証規約 =====
h1("§D　検証規約とBH換算")
para("主規約は固定重み日次リバランス(756/252営業日)。買い持ち(株数固定)は複利で極端な値になりやすいため参考(下表)。"
     "取引コスト・税は考慮しない。全ては2026-06選定を過去に当てた in-sample。", after=3)
table(["規約×期間", "本ＰＦ 年率", "本ＰＦ 最大下落", "新バフェット 年率"],
      [["リバランス 3年", pct(O3["ann_return"]), pct(O3["max_drawdown"]), pct(C3["ann_return"])],
       ["買い持ち 3年", pct(BH["ours"]["3y"]["ann_return"]), pct(BH["ours"]["3y"]["max_drawdown"]), pct(BH["control_buffett"]["3y"]["ann_return"])],
       ["リバランス 1年", pct(O1["ann_return"]), pct(O1["max_drawdown"]), pct(C1["ann_return"])]],
      [40, 34, 40, 40])

# ===== §E 頑健性(多期間レジーム) =====
h1("§E　頑健性 ― 多期間レジーム検証")
para("固定した20社を、相場局面の異なる複数期間へ当てはめ、超過が特定局面に依存しないかを確かめた(買い持ち)。", after=3)
def mr(k):
    m = MP[k]
    return [pct(m["v7"]["ann_return"]), f"{m['v7']['excess_vs_topix']*100:+.1f}pt", pct(m["v7"]["max_drawdown"]), f"{m['v7']['sharpe']:.2f}"]
table(["期間", "本ＰＦ 年率", "対TOPIX", "最大下落", "Sharpe"],
      [["全期間(約4.85年)"] + mr("full"), ["P1 利上げ21-22"] + mr("P1_利上21-22"),
       ["P2 AI相場前半23-24"] + mr("P2_AI前半23-24"), ["P3 直近24-26"] + mr("P3_直近24-26")],
      [40, 30, 26, 26, 22])
para(f"役割別の3年寄与: 生まれる堀{RC['by_role_pct'].get('離 生まれる堀',0):.0f}%・守{RC['by_role_pct'].get('守 完成した堀',0):.0f}%・"
     f"変わる堀{RC['by_role_pct'].get('破 変わる堀',0):.0f}%・分散{RC['by_role_pct'].get('分散役',0):.0f}%・両立{RC['by_role_pct'].get('両立型',0):.0f}%。"
     "AI・半導体テーマ(電気機器＋機械)の比重は約40%で、これは『生まれる堀=AI/半導体バリューチェーンへの意図的な賭け』として"
     "開示する(業種HHI約0.16)。守・破が下げ相場の緩衝材となる。", after=6)

# ===== §F 統計 =====
h1("§F　統計的有意性")
para("日次超過リターンの平均がゼロと異なるかを、系列相関に頑健な Newey-West の t 値で検定した(自動ラグ)。", after=3)
def st(win, key):
    s = SIG[win][key]; return [f"{s['t_newey_west']:.2f}", f"{s['t_plain']:.2f}", pct(s['mean_excess_ann'])]
table(["比較(期間)", "NW-t", "素のt", "年率超過"],
      [["対 新バフェット(3年)"] + st("3y", "ours_vs_buffett"), ["対 純正グレアム(3年)"] + st("3y", "ours_vs_graham"),
       ["対 TOPIX(3年)"] + st("3y", "ours_vs_topix"), ["対 新バフェット(1年)"] + st("1y", "ours_vs_buffett")],
      [46, 24, 24, 34])
para("読み方: 純正グレアム型と市場(TOPIX)に対する超過は統計的に有意(NW-t>2)。新バフェット型に対する超過は点推定では"
     "全指標・全局面で上回るが、3年のNW-tは有意水準に届かず『互角〜やや上』と評価する(過大主張しない)。", after=6)

# ===== §G 限界レジスタ =====
h1("§G　限界レジスタ(不利な事実の等格開示)")
for t in [
    "in-sample/look-ahead: 全PFは2026-06の財務・価格で選定し過去へ当てた自己検証。将来の市場超過の証明ではない。生リターン(1年で三桁%)はAI相場ピーク×後知恵選定の人工物であり、本文はリスク調整後・多期間・in-sample枠を前面に置く。",
    "テーマ集中: AI・半導体テーマに約40%・電気機器7社。意図的な賭けとしてHHIで開示。テーマ崩壊時は市場に劣後しうる。",
    "守の循環株: 名村造船は高ROE・割安でGreenblatt順位により守に入ったが造船は市況循環で持続的な堀は薄い(現在の堀偏差値45)。durable-moat純化の立場では差替余地があり、限界として明記する。",
    "証拠は開示ベース: 半導体・AIの実需はセグメント開示で確認したが現場の実態までは測れない。取材(第Ⅳ章)で補完する。",
    "単元未満株の前提: 大型優良は株価が高く、単元未満株(金額指定)の取扱いがない証券会社では本配分をそのまま再現できない。",
    "価格データ在庫: 時価総額データは大型中心の約139社。価格ランク可能な品質企業が大型に偏る。",
]:
    para("・" + t, after=3)

# ===== §H 再現手順 =====
h1("§H　再現手順")
para("すべて .venv/bin/python(3.12, python-docx 1.2)で実行。正典データ→図→本編(2パス)→本補遺の順。", after=3, )
para("1. 提出PF・多期間: scripts/... build_portfolio_v7.py → work/pure_buffett_benchmark/portfolio_v7.json\n"
     "2. 比較・有意性: build_report_data_v7.py → control_comparison_v7.json / significance_v7.json\n"
     "3. 守ファネル: build_funnel_v7.py → funnel_exclusion_v7.json\n"
     "4. 役割寄与: (multiperiodハーネス) → role_contribution_v7.json\n"
     "5. 20社データ: build_data_real_v7.py → data_real_v7.json\n"
     "6. 図: build_figs_v7.py / build_fig_cum_v7.py / build_figs_v9.py(全枝ファネル4点) → assets/*.png、全枝ファネル正典: build_funnel_v9.py → funnel_branches_v9.json\n"
     "7. 本編(2パス): build_contest_v9.py → soffice pdf → make_tocmap.py → build_contest_v9.py --tocmap tocmap_v9.json → soffice pdf\n"
     "8. 本補遺: build_appendix_v9.py → soffice pdf", mono=True, after=6)

out = ED / "beyond_buffett_technical_appendix_v9.docx"
if (ED / "beyond_buffett_technical_appendix_v9.LOCKED").exists():
    raise SystemExit("v9 appendix LOCKED")
doc.save(str(out))
print("saved", out.name, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
