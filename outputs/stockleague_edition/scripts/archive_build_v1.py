# -*- coding: utf-8 -*-
"""STOCKリーグ提出版（contest edition）生成器。
入賞作様式：全面表紙／基礎学習枠／要旨（造語コンセプト）／頁番号つき目次（2パス）／
Ⅰ物語序論→Ⅱスクリーニング（判定基準表＋式6本）→Ⅲ配分・銘柄紹介・パフォーマンス→
Ⅳインタビュー（質問草案）→Ⅴ学んだこと→参考文献。紺×ティール。実データ準拠・捏造なし。
使い方: build_contest.py [--tocmap tocmap.json]
"""
import sys, json, re
from pathlib import Path

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE))
sys.path.insert(0, "/Users/satouryuuichi/Desktop/product/hobby/stock_league/outputs/final_revision/scripts")
import docxlib as X
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
ED = HERE.parent
ASSETS = ED / "assets"
FIG_T = ROOT / "outputs/beyond_buffett_fable_loop_final/phase7_final_report/final_figures"  # teal figs
DATA = json.load(open(ROOT / "outputs/explanatory_revision/data_real.json", encoding="utf-8"))

NAVY = "16324F"; NAVY2 = "1F4568"; TEAL = "2F6D5F"; TEAL_L = "DCE8E4"; GOLD = "C9A227"
NAVYC = RGBColor(0x16, 0x32, 0x4F); TEALC = RGBColor(0x2F, 0x6D, 0x5F)
GO, MIN, WHITE, BLACK = X.GO, X.MIN, X.WHITE, X.BLACK
order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
role_jp = {"Buffett Core": "完成した堀", "Transformation Core": "変わる堀",
           "Emerging Core": "生まれる堀", "Dual Moat": "両立型", "Bridge / Diversifier": "分散役"}

TOCMAP = {}
if len(sys.argv) > 2 and sys.argv[1] == "--tocmap":
    TOCMAP = json.load(open(sys.argv[2], encoding="utf-8"))

doc = Document()
X.setup_styles(doc)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)

HEADINGS = []  # (text, level) for TOC


def para(text="", name=MIN, size=10, bold=False, align=None, before=0, after=4,
         color=None, indent=False, line=1.28, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if keep:
        pf.keep_with_next = True
    if not indent:
        X._clear_indent(p)
    if text:
        X.setfont(p.add_run(text), name, size, bold, color or BLACK)
    return p


def banner1(text):
    HEADINGS.append((text, 1))
    p = para(before=12, after=5, keep=True); X.shade_cell.__wrapped__ if False else None
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    pr.append(shd)
    X.setfont(p.add_run("　" + text), GO, 14, True, WHITE)
    return p


def banner2(text):
    HEADINGS.append((text, 2))
    p = para(before=9, after=4, keep=True)
    pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), TEAL_L)
    pr.append(shd)
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:left")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "28"); b.set(qn("w:space"), "3"); b.set(qn("w:color"), NAVY)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run("　" + text), GO, 11.5, True, NAVYC)
    return p


def banner3(text):
    p = para(before=6, after=3, keep=True)
    pr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "dashed"); b.set(qn("w:sz"), "10"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), TEAL)
    pbdr.append(b); pr.append(pbdr)
    X.setfont(p.add_run(text), GO, 10.5, True, TEALC)
    return p


def body(text, size=10):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_after = Pt(4); pf.line_spacing = 1.3
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), "100"); ind.set(qn("w:firstLine"), "200"); pp.append(ind)
    X.setfont(p.add_run(text), MIN, size, False, BLACK)
    return p


def quote(text, who):
    p = para(before=6, after=2, line=1.5)
    pp = p._p.get_or_add_pPr(); ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480"); ind.set(qn("w:right"), "480"); pp.append(ind)
    r = p.add_run("「" + text + "」"); X.setfont(r, MIN, 11, True, NAVYC); r.font.italic = True
    p2 = para(after=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    X.setfont(p2.add_run("――　" + who), MIN, 9, False, BLACK)


def fig(fname, title, note, maxw=150, src=None):
    fp = (src or FIG_T) / fname
    if not fp.exists():
        fp = ASSETS / fname
    pt = para(before=6, after=2, keep=True)
    X.setfont(pt.add_run(title), GO, 9.5, True, NAVYC)
    if fp.exists():
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        X._clear_indent(pi); pi.paragraph_format.keep_with_next = True
        w, h = Image.open(fp).size; wmm = min(maxw, w / 220 * 25.4)
        if wmm * h / w > 165:
            pi.add_run().add_picture(str(fp), height=Mm(165))
        else:
            pi.add_run().add_picture(str(fp), width=Mm(wmm))
    para(note, MIN, 8, after=8)


def table(hdr, rows, widths, hdr_fill=NAVY, fs=8.5, zebra=True):
    t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(hdr):
        c = t.rows[0].cells[j]; X.shade_cell(c, hdr_fill)
        pp = c.paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
        X.setfont(pp.add_run(h), GO, fs, True, WHITE)
    for i, r in enumerate(rows):
        cells = t.add_row().cells
        for j in range(len(hdr)):
            if zebra and i % 2 == 1:
                X.shade_cell(cells[j], "F2F6F4")
            pp = cells[j].paragraphs[0]; X._clear_indent(pp); pp.paragraph_format.space_after = Pt(0)
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            X.setfont(pp.add_run(str(r[j]) if j < len(r) else ""), MIN, fs, False, BLACK)
    for r in t.rows:
        X.cant_split(r)
        for j, w in enumerate(widths):
            r.cells[j].width = Mm(w)
    para("", after=4)
    return t


def formula(num, title, latex, defn):
    t = doc.add_table(rows=3, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in t.rows:
        X.cant_split(r); r.cells[0].width = Mm(166)
    c0 = t.rows[0].cells[0]; X.shade_cell(c0, NAVY); X.cell_borders(c0, top=True, bottom=True)
    p0 = c0.paragraphs[0]; X._clear_indent(p0); p0.paragraph_format.space_after = Pt(0); p0.paragraph_format.keep_with_next = True
    X.setfont(p0.add_run(f"　式（{num}）{title}"), GO, 9.5, True, WHITE)
    c1 = t.rows[1].cells[0]; X.cell_borders(c1, bottom=True)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; X._clear_indent(p1)
    p1.paragraph_format.space_before = Pt(5); p1.paragraph_format.space_after = Pt(5); p1.paragraph_format.keep_with_next = True
    om = X.latex_to_omml(latex)
    from lxml import etree
    M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    pr = om.find(f"{{{M}}}oMathParaPr")
    if pr is None:
        pr = etree.Element(f"{{{M}}}oMathParaPr"); om.insert(0, pr)
    jc = pr.find(f"{{{M}}}jc")
    if jc is None:
        jc = etree.SubElement(pr, f"{{{M}}}jc")
    jc.set(f"{{{M}}}val", "center")
    p1._p.append(om)
    c2 = t.rows[2].cells[0]; X.cell_borders(c2, bottom=True)
    p2 = c2.paragraphs[0]; X._clear_indent(p2)
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    X.setfont(p2.add_run(defn), MIN, 8, False, BLACK)
    para("", after=4)


def fill_line(label):
    p = para(after=2)
    X.setfont(p.add_run("・" + label + "："), MIN, 10, True, BLACK)
    X.setfont(p.add_run("［記入］"), MIN, 10, False, RGBColor(0x88, 0x88, 0x88))


def step_chip(no, name):
    p = para(before=8, after=2, keep=True)
    r = p.add_run(f"  STEP {no}  ")
    X.setfont(r, GO, 10, True, WHITE)
    rpr = r._element.get_or_add_rPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), NAVY)
    rpr.append(shd)
    X.setfont(p.add_run("　" + name), GO, 11, True, NAVYC)


# ================= 表紙（全面） =================
sec.top_margin = Mm(0); sec.bottom_margin = Mm(0)
sec.left_margin = sec.right_margin = Mm(0)
pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc.paragraph_format.space_before = Pt(0); pc.paragraph_format.space_after = Pt(0)
pc.add_run().add_picture(str(ASSETS / "cover.png"), width=Mm(210))
# new section with normal margins
ns = doc.add_section(WD_SECTION.NEW_PAGE)
ns.left_margin = ns.right_margin = Mm(20); ns.top_margin = Mm(20); ns.bottom_margin = Mm(20)

# ================= 基礎学習（記入枠） =================
banner1("基礎学習")
body("日経ＳＴＯＣＫリーグ所定の基礎学習ワークシートに基づき、当年度の設問への解答を本節に記入する（配布年度により設問が異なるため、本テンプレートでは枠のみを用意した）。")
fill_line("基礎学習の設問と解答（当年度ワークシートを転記）")
banner3("投資テーマと関連が深い SDGs の目標（3つ以内・記入）")
table(["関連の深い SDGs の目標", "その主な理由（記入）"],
      [["9．産業と技術革新の基盤をつくろう", "［記入例の観点：AI・半導体・光通信など「生まれる堀」への投資が技術革新の基盤形成を支えるため］"],
       ["8．働きがいも経済成長も", "［記入例の観点：資本効率改革で「変わる堀」を持つ企業の再評価が日本経済の成長につながるため］"],
       ["12．つくる責任・つかう責任", "［記入］"]],
      [58, 108], fs=9)
doc.add_page_break()

# ================= 要旨・目次 =================
banner1("要旨")
body("本レポートでは、これからの日本株投資の姿として「三世代の堀（Three-Generation Moat）」を提示する。堀（Moat）とは、他社がまねしにくい持続的な競争優位のことである。私たちは、ウォーレン・バフェットが築いた「完成した堀を割安に買う」という規律を出発点に、堀を時間軸でとらえ直した。すなわち、いま完成している堀だけでなく、資本効率改革によってこれから「変わる堀」、AI・半導体・光通信などの構造変化からこれから「生まれる堀」までを、一つのポートフォリオに束ねる。")
body("選定は、剣道の修行段階になぞらえた守・破・離の三段階スクリーニングで行った。守では先行研究の式だけで完成した堀を5社に絞り、破では式を変えずに候補を1,200社へ広げ、離では変わる堀・生まれる堀を証拠水準つきで測って最終20社を組んだ。銘柄は一度も成績を見て入れ替えていない。私たちは、成績を誇る代わりに、この20社が「壊れにくい構造」を持つことを16通りの検査で確かめた。三世代の堀こそ、バフェットを超える——正しくは、バフェットの規律を未来へ延長する——私たちの答えである。")

banner1("目次")
toc_items = ["基礎学習", "要旨", "Ⅰ．背景・投資テーマ決定", "Ⅱ．スクリーニング ― 守・破・離",
             "Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス", "Ⅳ．インタビュー・アンケート",
             "Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと", "参考文献"]
for it in toc_items:
    p = para(after=2)
    X.setfont(p.add_run(it), MIN, 10, False, BLACK)
    pg = TOCMAP.get(it, "")
    if pg:
        X.setfont(p.add_run("　" + "･" * max(4, 56 - len(it) * 2) + "　" + str(pg)), MIN, 10, False, BLACK)
doc.add_page_break()

# ================= Ⅰ 背景・投資テーマ決定 =================
banner1("Ⅰ．背景・投資テーマ決定")
banner2("0．イントロダクション ― 百年企業の中に、未来の堀が眠っていた")
quote("素晴らしい会社をまずまずの価格で買うほうが、まずまずの会社を素晴らしい価格で買うよりずっと良い。",
      "ウォーレン・バフェット（1989年 株主への手紙）")
body("バフェットは、企業を城に、競争優位を城を守る堀（Moat）にたとえた。深い堀を持つ会社を、高すぎない価格で買い、長く持つ。この単純な規律が、半世紀を超える成果を生んだ。私たちの研究は、この規律への挑戦ではなく、その延長である。")
body("きっかけは、一つの会社だった。フジクラ（5803）——1885年創業の電線メーカーである。長らく「成熟産業の老舗」と見られてきたこの会社は、いま、生成AIのデータセンターを支える光配線の担い手として再評価されている。百年かけて磨いた細径高密度の配線技術が、AIという新しい産業の「ボトルネック」を握ったのだ。完成したはずの堀の内側で、新しい堀が生まれていた。")
body("同じ頃、東京証券取引所は上場企業に「資本コストや株価を意識した経営」を求め、長く割安に放置された企業が、資本効率の改善と株主還元の強化によって変わり始めた。堀は静的ではない。変わる堀があり、生まれる堀がある。ならば、いまの堀だけを測るバフェットの物差しを、時間軸の方向へ延長できないか。これが私たちの問いである。")
banner2("1．背景 ― 二つの地殻変動")
body("第一の地殻変動は資本効率改革である。東京証券取引所（2023）の要請は、伊藤レポート（経済産業省、2014）以来の資本収益性重視の流れを決定づけ、PBR（株価純資産倍率）1倍割れ企業の経営改革を促した。割安に放置された企業の中から、本当に変われる企業を見分けられれば、それは「変わる堀」への投資になる。")
body("第二の地殻変動はAIを起点とする産業構造の転換である。生成AIは、半導体・光通信・データセンター・電力・品質保証といった基盤産業に、実需としての追い風を与えている。ただし「AI」と口にする企業がすべて堀を持つわけではない。私たちの調査では、開示資料でAIに言及しながら具体的な製品・顧客・数量の裏づけを欠く企業が577社に上った。キーワードと堀を見分ける仕組みが要る。")
banner2("2．三世代の堀 ― 私たちが求める企業像")
body("以上から、私たちは投資対象を三つの世代で定義する。第一に、すでに完成した堀（バフェットの規律で測れる企業）。第二に、資本効率改革で変わる堀（割安×改善の証拠を持つ企業）。第三に、構造変化から生まれる堀（AI基盤への実需接続を証拠で示せる企業）。この三世代を一つのポートフォリオに束ね、どれか一つの未来に賭けない。これが「三世代の堀」である。")
fig("moat_timeaxis.png", "図表 Ⅰ-1　三世代の堀（完成・変化・新生）",
    "注：完成した堀（守 Top5）・変わる堀（Transformation）・生まれる堀（Emerging）。筆者作成。")
banner2("3．投資テーマの決定 ― 守・破・離")
body("三世代の堀を、思いつきの銘柄選びにしないため、私たちは選定手順そのものを守・破・離の三段階に設計した。守では先行研究の式を一切変えずに使い、破では式を変えずに使い方だけを最適化し、離で初めて独自の物差しを導入する。独自性を最後に置くことで、独自の物差しがどこで何を変えたかを、いつでも遡って検証できる。")
doc.add_page_break()

# ================= Ⅱ スクリーニング =================
banner1("Ⅱ．スクリーニング ― 守・破・離")
body("本章が選定の中核である。3,099社の非金融普通株から、三段階のスクリーニングで最終20社を選ぶ。各段階では、何を見たか（観点）、何で測ったか（データと式）、どこで区切ったか（基準）、何社が残ったか（結果）を必ず示す。")

step_chip(1, "守 ― 完成した堀を抽出する（→ Buffett Core 5社）")
body("守では、バフェットの規律「良い会社を、高すぎない価格で買い、長期保有に耐えない企業を避ける」を、先行研究の式だけで再現する。独自の重み付きスコアは使わず、関門を一つずつ通す。判定基準は次の通りである。")
table(["観点", "使うデータ・式", "通過基準", "通過社数"],
      [["割安か（純資産）", "B/M＝純資産÷時価総額（式1）", "市場上位30%", "3,099→2,740社（算出可能）"],
       ["割安か（利益）", "E/P＝純利益÷時価総額", "黒字かつ上位50%", "→583社（Value通過）"],
       ["良い会社か", "Gross Profitability（式2）", "市場中央値以上", "（同上583社に含む）"],
       ["財務は健全か", "Piotroski 合格割合", "0.65以上", "→146社"],
       ["利益は本物か", "Sloan アクルーアル", "悪い側上位を除外", "→112社"],
       ["壊れにくいか", "危機ガードレール（債務超過・3期連続赤字）", "該当ゼロ", "→90社"],
       ["実際に買えるか", "60日平均売買代金", "約1,000万円/日以上", "→77社"]],
      [34, 52, 40, 40], fs=8.5)
formula("1", "Ｂ／Ｍ ― 純資産に対して安いか（Fama and French, 1993）",
        r"\mathrm{BM}_i = \frac{\mathrm{BE}_i}{\mathrm{ME}_i}",
        "BE＝自己資本（純資産）　ME＝時価総額（株価×株式数）。大きいほど純資産のわりに株価が安い。例：9470 学研ＨＤは1.526で市場上位30%。")
formula("2", "Gross Profitability ― 資産を粗利に変える力（Novy-Marx, 2013）",
        r"\mathrm{GP}_i = \frac{\mathrm{Revenue}_i - \mathrm{COGS}_i}{\mathrm{TotalAssets}_i}",
        "Revenue＝売上高　COGS＝売上原価　TotalAssets＝総資産。大きいほど資産効率の高い優良企業。例：9470は0.395で中央値以上。")
body("通過77社から、GP→E/P→B/M→Piotroski→Sloan→流動性→時価総額の固定順で上位5社を選び、Buffett Core として固定した（同一業種は原則2社まで）。選ばれたのは、3539 JMホールディングス、4350 メディカルシステムネットワーク、6430 大黒電機、7803 ブシロード、9470 学研ホールディングスである。以降、この5社は一度も入れ替えない。")
fig("phase1_flow.png", "図表 Ⅱ-1　守の絞り込み（3,099社→Top5）",
    "注：各段階の通過社数。出所：Phase1正典データより筆者作成。")

step_chip(2, "破 ― 式を変えずに候補を1,200社へ広げる")
body("守は規律が強い分、通過がわずか77社になる。これでは「変わる堀」「生まれる堀」の候補を評価する余地がない。破では、守の式の定義を一切変えず、閾値・分位・候補数という「使い方」だけを最適化して、次の段階で精査する候補宇宙を1,200社へ広げた。")
table(["観点", "内容", "結果"],
      [["正規化", "各指標を市場内の順位（0〜1）へ変換。式の定義は不変", "指標同士を公平に比較可能に"],
       ["頑健性の監査", "4通りの正規化方式で共通して上位に残るかを確認", "core 970社・robust 1,135社"],
       ["候補数の決定", "広さ・品質・安全性・流動性・レビュー可能性を比較", "Top1200 を正式採用（Top2000は参照）"],
       ["守との整合", "守のTop5が候補宇宙に含まれるか", "5社すべて Top1200 に含まれる"]],
      [30, 84, 52], fs=8.5)
body("なお1,200という数は、数式が出した唯一の正解ではない。候補を広げるほど取りこぼしは減るが、一社ずつ確認できる規模には限りがある。品質・広さ・レビュー可能性を両立させる運営上の判断として1,200社を選び、その判断過程を開示した。破のスコアは候補づくりの道具であり、最終選定には使っていない。")
fig("phase2_flow.png", "図表 Ⅱ-2　破：候補宇宙の形成（Top1200）",
    "注：式は不変のまま使い方を最適化し、入れ子の母集団から Top1200 を正式採用。筆者作成。")

step_chip(3, "離 ― 変わる堀・生まれる堀を測り、20社を組む")
body("離で初めて、私たち独自の物差しを導入する。ただし新しい式を発明したのではない。守で使った先行研究の部品（割安・改善・利益の質）を、時間軸の方向へ組み替えた。変わる堀は式（3）、生まれる堀は式（4）で測り、スコアの根拠の強さは式（5）の証拠水準としてスコアと分けて管理する。")
formula("3", "Transformation Score ― 変わる堀（実装形・選定に使用）",
        r"S^{\mathrm{Trans}}_i = 0.22\,V_i + 0.24\,C_i + 0.10\,F_i + 0.18\,X_i + 0.16\,Q_i + 0.10\,\Phi_i - P^{\mathrm{Trap}}_i",
        "V＝割安度　C＝資本効率の改善　F＝還元余力（営業CF−設備投資）　X＝実行信頼性　Q＝利益の質　Φ＝データ信頼度　P^Trap＝割安のワナ減点。低PBR株選びではない：割安でも改善証拠を欠く15社を除外した。")
formula("4", "Emerging Score ― 生まれる堀（選定に使用）",
        r"S^{\mathrm{Emerg}}_i = 100\,(w_I I_i + w_N N_i + w_B B_i + w_A A_i + w_D D_i + w_T T_i) + B^{\mathrm{Evi}}_i - P^{\mathrm{Hype}}_i",
        "I＝無形資産　N＝技術力　B＝ボトルネック性　A＝AI基盤接続　D＝データ顧客基盤　T＝信頼安全基盤　B^Evi＝証拠加点(≤8)　P^Hype＝キーワードのみ18点減点。AIテーマ株選びではない：言及のみの577社を除外した。")
formula("5", "Evidence Level ― 証拠の強さ（スコアと分離して管理）",
        r"L^{\mathrm{final}}_i = \begin{cases} \min(L^{TQ}_i, L^{EM}_i) & \text{両立型} \\ L^{EM}_i & \text{生まれる堀} \\ \max(L^{TQ}_i, L^{TS}_i, L^{TR}_i) & \text{変わる堀} \\ \max(L^{TQ}_i, L^{EM}_i) & \text{その他} \end{cases}",
        "Level 1＝言及のみ／Level 2＝製品・顧客・投資計画の具体性／Level 3＝売上・受注・投資額の数量根拠。最終20社の分布は L3：15社、L2：4社、L1：1社。")
body("最終20社は、完成した堀5社（守Top5固定）・変わる堀5社・生まれる堀5社・両立型3社・分散役2社の五役割で構成した。同一業種は3社まで、同一テーマは4社までとし、除外した862社は理由別（AIキーワードのみ577・上位互換あり221・財務危険32・割安のワナ17・低PBRのみ15）にすべて記録した。選定後に成績を見た入れ替えは一度も行っていない。")
fig("phase3_flow.png", "図表 Ⅱ-3　離：三世代の堀から最終20社へ",
    "注：二つのスコアと証拠水準から五役割へ。筆者作成。")
banner2("選定は壊れにくいか ― 16通りの分解検査（アブレーション）")
body("最後に、この選定が特定の一要素や一テーマ、後付けの調整に依存していないかを、条件を一つずつ外した16通りの再選定で確かめた。最も影響が大きかったのは「候補をTop100に狭める」場合で、20社中7社しか一致しない。つまり、破で候補を1,200社へ広げた判断こそが選定の背骨である。逆に、どのペナルティを一つ外しても選定は崩壊しない。特定の仕掛け一つに頼った選定ではないことを、数で確かめた。")
fig("ablation_overlap.png", "図表 Ⅱ-4　条件を一つ外したときの最終20社との一致数",
    "注：A8（Top100限定）が最小の7。過去データ上の構造検査であり、成績の主張ではない。出所：ablation_results.csv。")
doc.add_page_break()

# ================= Ⅲ ポートフォリオ =================
banner1("Ⅲ．ポートフォリオの決定・銘柄紹介・パフォーマンス")
banner2("1．500万円の配り方 ― 成績予想を使わない配分")
body("配分は四段階で決めた。①20社を役割ごとに分け、②役割ごとの予算を決め（完成・変化・新生に各25%、両立型15%、分散役10%）、③同じ役割の中では「値動きが穏やかで・売買しやすく・証拠が強く・データが信頼できる」会社を厚くし（式6）、④1銘柄8%の上限と実際に買える株数（式7）に丸める。リターンの予想は一切使っていない。")
formula("6", "役割予算つきリスク調整配分（選定済み20社への配り方）",
        r"\omega_i = B_{r(i)} \cdot \frac{\rho_i}{\sum_{j:\,r(j)=r(i)} \rho_j}, \qquad \rho_i = \frac{\ell_i\, e_i\, c_i}{\max(\sigma_i,\,0.10)}",
        "B＝役割予算（25/25/25/15/10%）　ℓ＝流動性係数　e＝証拠係数　c＝信頼度係数　σ＝1年の値動きの大きさ。8%超過分は同じ役割内で配り直す。")
formula("7", "単元株調整 ― 実際に買える株数へ丸める",
        r"q_i = L_i \cdot \operatorname{floor}\!\left( \frac{B\,\omega_i}{P_i L_i} \right)",
        "q＝購入株数　L＝売買単位（基準1株）　B＝総予算500万円　ω＝目標比率　P＝株価　floor＝切り捨て。投資額4,949,198円・残現金50,801円（消化率99.0%）。")
body("最終配分を図表Ⅲ-1に示す。最大保有はゼンリンの7.46%、役割合計は25/25/25/15/10で、すべて制約内である。実単元（100株）では株価の高い9社が買えないため、単元未満株の利用を前提とする。")
prows = []
for role in order:
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        prows.append([d["code"], d["name"][:16], role_jp[role], d["theme"],
                      f"{d['w']*100:.2f}%", f"{d['qtyL1']:,}", f"{d['amtL1']:,}円"])
table(["コード", "企業名", "役割", "テーマ", "比率", "株数", "金額"], prows,
      [13, 36, 20, 24, 14, 14, 21], fs=8)
para("図表 Ⅲ-1　最終ポートフォリオ（総予算500万円・単元未満株基準）。出所：allocation_final.csv。", MIN, 8, after=8)

banner2("2．銘柄紹介 ― 三世代の堀、20社の顔ぶれ")
body("各社の「定量の骨子」（なぜその役割か）は選定データから記載済みである。事業概要と定性的な堀の根拠は、各社のＩＲ資料・有価証券報告書で確認のうえ記入する（［要ＩＲ確認］欄）。")
mrows = []
draft = {
 "Buffett Core": "守の全関門を通過。割安×優良×健全の完成した堀",
 "Transformation Core": "割安かつ資本効率が改善中。還元余力あり＝変わる堀",
 "Emerging Core": "AI基盤への実需接続を証拠水準L2以上で確認＝生まれる堀",
 "Dual Moat": "変わる堀と生まれる堀の両立を min 式で確認",
 "Bridge / Diversifier": "業種・テーマの分散を担う堅実枠",
}
for role in order:
    for d in sorted([x for x in DATA.values() if x["role"] == role], key=lambda z: -z["w"]):
        quant = f"{role_jp[role]}／T={d['tsc']:.0f}・E={d['esc']:.0f}／証拠L{d['evid']}／比率{d['w']*100:.1f}%"
        mrows.append([f"{d['code']} {d['name'][:13]}\n（{d['sector'][:14]}）", quant + "。" + draft[role],
                      "［要ＩＲ確認：事業概要・堀の根拠・リスク］"])
table(["企業（業種）", "定量の骨子（選定データより）", "事業概要・堀の根拠（記入）"], mrows,
      [38, 66, 62], fs=7.5)
para("図表 Ⅲ-2　銘柄紹介（左2列は選定データから自動記載、右列は提出者がＩＲ確認のうえ記入）", MIN, 8, after=8)

banner2("3．パフォーマンスの分析 ― 成績を誇る代わりに、壊れにくさを確かめる")
body("私たちはこのポートフォリオの成績を予測しない。構築に使ったのと同じ期間で成績を測っても（標本内・in-sample）、それは未来の証明にならないからだ。その代わり、過去3年のデータで「リスクの姿」を確かめた。年率の値動きは21.9%、最大の下落は−24.9%で市場（TOPIX −23.3%）並み、市場連動度ベータは0.925。リターンの牽引役は生まれる堀（フジクラ・オルガノなど）で、完成した堀は安定の土台に徹していた。割安株を外すと値動きが35.5%へ悪化することも分かった——変わる堀は、攻めではなく緩衝材として効いている。")
body("誠実のために記す。直近1年では市場に負けている（超過リターンの安定度ＩＲ＝−0.405）。それでも私たちがこの20社を持てるのは、選定が成績ではなく、三世代の堀という構造と証拠に基づくからである。")
fig("drawdown_chart.png", "図表 Ⅲ-3　過去3年の累積と下落（標本内・リスク特性の確認）",
    "注：成績の主張ではない。ベンチマークの1306.Tは未調整の株式分割（2026-03-30）を補正済み。")
fig("role_contribution.png", "図表 Ⅲ-4　役割別の寄与 ― 生まれる堀が牽引、完成した堀は土台",
    "注：過去3年・標本内。出所：phase5_validation_summary.json。")
doc.add_page_break()

# ================= Ⅳ インタビュー =================
banner1("Ⅳ．インタビュー・アンケート")
body("スクリーニングは開示データに基づく。その仮説が実態と合っているかを、投資先企業への取材で確かめる。以下に、選定理由から導いた質問草案を用意した（実施記録は記入欄）。")
table(["対象（役割）", "分析から導いた質問草案"],
      [["5803 フジクラ（生まれる堀）", "①データセンター向け光配線の受注動向と生産能力投資の計画　②AI需要が一巡した場合の下支え事業"],
       ["6368 オルガノ（生まれる堀）", "①半導体向け超純水装置の需要見通し　②価格決定力（ボトルネック性）の源泉"],
       ["9470 学研ＨＤ（完成した堀）", "①教育×医療福祉の資本効率改善の進捗　②株主還元方針の考え方"],
       ["9474 ゼンリン（両立型）", "①地図データベースの参入障壁　②AI時代におけるデータ資産の収益化"],
       ["3697 ＳＨＩＦＴ（両立型）", "①ソフトウェア品質保証の需要構造　②AI活用による事業機会と脅威"]],
      [52, 114], fs=8.5)
banner3("実施記録（記入）")
for lab in ["対象者（企業名・部署・役職）", "実施日・実施方法", "質問と回答", "分析への反映（仮説はどう変わったか）"]:
    fill_line(lab)
para("（未実施の場合は本章を「今後の課題」とし、その旨を明記する。回答の創作はしない。）", MIN, 8, after=6)

# ================= Ⅴ 学んだこと =================
banner1("Ⅴ．日経ＳＴＯＣＫリーグを通じて学んだこと")
for lab in ["仮説が変わった瞬間（例：成績ではなく構造で持つ、という発想への転換）",
            "うまくいかなかったこと・工夫したこと",
            "チームでの役割分担と学び",
            "今後の課題（単元未満株の扱い、証拠水準の自動化など）"]:
    fill_line(lab)
para("（提出者の実体験として記入する。下書きの創作はしない。）", MIN, 8, after=6)

# ================= 参考文献 =================
banner1("参考文献")
refs_en = [
 "Basu, S. (1977) Investment Performance of Common Stocks in Relation to Their Price-Earnings Ratios. The Journal of Finance, 32(3), pp.663-682.",
 "Fama, E. F., & French, K. R. (1993) Common Risk Factors in the Returns on Stocks and Bonds. Journal of Financial Economics, 33(1), pp.3-56.",
 "Frazzini, A., Kabiller, D., & Pedersen, L. H. (2018) Buffett's Alpha. Financial Analysts Journal, 74(4), pp.35-55.",
 "Jensen, M. C. (1968) The Performance of Mutual Funds in the Period 1945-1964. The Journal of Finance, 23(2), pp.389-416.",
 "Novy-Marx, R. (2013) The Other Side of Value: The Gross Profitability Premium. Journal of Financial Economics, 108(1), pp.1-28.",
 "Piotroski, J. D. (2000) Value Investing. Journal of Accounting Research, 38, Supplement, pp.1-41.",
 "Sharpe, W. F. (1966) Mutual Fund Performance. The Journal of Business, 39(1), pp.119-138.",
 "Sloan, R. G. (1996) Do Stock Prices Fully Reflect Information in Accruals and Cash Flows about Future Earnings? The Accounting Review, 71(3), pp.289-315.",
]
refs_ja = [
 "経済産業省（2014）『持続的成長への競争力とインセンティブ（伊藤レポート）』経済産業省．",
 "東京証券取引所（2023）「資本コストや株価を意識した経営の実現に向けた対応について」株式会社東京証券取引所．",
 "バフェット, W.（1989）「株主への手紙」Berkshire Hathaway Inc.（筆者訳）．",
 "金融庁「EDINET」，https://disclosure2.edinet-fsa.go.jp/（2026年7月8日）",
 "日本取引所グループ「上場会社情報」，https://www.jpx.co.jp/listing/co-search/（2026年7月8日）",
]
banner3("英語文献")
for r in refs_en:
    p = doc.add_paragraph(); pp = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360"); pp.append(ind)
    p.paragraph_format.space_after = Pt(1)
    X.setfont(p.add_run(r), MIN, 8.5, False, BLACK)
banner3("日本語文献")
for r in refs_ja:
    p = doc.add_paragraph(); pp = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360"); pp.append(ind)
    p.paragraph_format.space_after = Pt(1)
    X.setfont(p.add_run(r), MIN, 8.5, False, BLACK)

out = ED / "beyond_buffett_stockleague_v1.docx"
doc.save(str(out))
json.dump([h for h, l in HEADINGS if l == 1], open(ED / "headings_v1.json", "w"), ensure_ascii=False)
print("saved", out.name, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
