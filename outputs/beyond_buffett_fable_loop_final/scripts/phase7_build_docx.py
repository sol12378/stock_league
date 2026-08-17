#!/usr/bin/env python
"""Build final_report.docx in NIKKEI STOCK LEAGUE prize-report style from
final_report.md. Renders: teal banner headings (L1 filled / L2 light box /
L3 underlined), formula boxes ("model name (citation)" banner + centered eq +
variable defs + 図表番号 caption below), data tables (teal header + caption
below), and fill-in TEMPLATE tables (portfolio / 銘柄紹介 / interview / learned)
for the submitter to complete. Equations are rendered as readable text
(no LaTeX engine); finalise with Word's equation editor using
final_report_latex_equations.md.
"""
import re
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

ROOT = Path("/Users/satouryuuichi/Desktop/product/hobby/stock_league")
OUT = ROOT / "outputs/beyond_buffett_fable_loop_final"
P7 = OUT / "phase7_final_report"
FIG = P7 / "final_figures"
MIN = "ＭＳ 明朝"
GO = "ＭＳ ゴシック"
TEAL = "2F6D5F"
TEAL_LT = "DCE8E4"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEALC = RGBColor(0x2F, 0x6D, 0x5F)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = sec.bottom_margin = Mm(22)


# ---- low-level helpers ----------------------------------------------------
def setfont(run, name=MIN, size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    latin = "Times New Roman" if name == MIN else name
    rfonts.set(qn("w:ascii"), latin); rfonts.set(qn("w:hAnsi"), latin)
    if color is not None:
        run.font.color.rgb = color


def shade(el, fill):
    """add w:shd to a paragraph's pPr or a cell's tcPr element."""
    pr = el.get_or_add_pPr() if el.tag.endswith("}p") else el.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pr.append(shd)


def border(p, edges=("bottom",), color=TEAL, sz="18", val="single"):
    pr = p._element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for e in edges:
        b = OxmlElement(f"w:{e}")
        b.set(qn("w:val"), val); b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "3"); b.set(qn("w:color"), color)
        pbdr.append(b)
    pr.append(pbdr)


def para(text="", name=MIN, size=9, bold=False, align=None, before=0, after=2,
         color=None, indent=None, hang=None, line=None, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if keep:
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    if line is not None:
        p.paragraph_format.line_spacing = line
    if indent is not None:
        p.paragraph_format.left_indent = Pt(indent)
    if hang is not None:
        p.paragraph_format.left_indent = Pt(hang)
        p.paragraph_format.first_line_indent = Pt(-hang)
    if text:
        setfont(p.add_run(text), name, size, bold, color)
    return p


def banner(text, level):
    if level == 1:
        p = para(before=10, after=4, keep=True); shade(p._element, TEAL)
        setfont(p.add_run("　" + text), GO, 13, True, WHITE)
    elif level == 2:
        p = para(before=8, after=3, keep=True); shade(p._element, TEAL_LT)
        border(p, edges=("left",), color=TEAL, sz="30")
        setfont(p.add_run("　" + text), GO, 11, True, TEALC)
    else:
        p = para(before=6, after=2, keep=True); border(p, edges=("bottom",), color=TEAL, sz="12", val="dashed")
        setfont(p.add_run(text), GO, 10.5, True, TEALC)


# ---- LaTeX -> readable text ----------------------------------------------
def tex(s):
    s = s.strip()
    s = re.sub(r"\\dfrac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", s)
    s = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", s)
    s = s.replace(r"\times", "×").replace(r"\cdot", "·")
    s = s.replace(r"\le", "≤").replace(r"\ge", "≥").replace(r"\sum", "Σ")
    s = s.replace(r"\Omega", "Ω").replace(r"\rho", "ρ").replace(r"\sigma", "σ")
    s = s.replace(r"\Phi", "Φ").replace(r"\omega", "ω").replace(r"\alpha", "α")
    s = s.replace(r"\mathbf", "").replace(r"\mathrm", "").replace(r"\text", "")
    s = s.replace(r"\left", "").replace(r"\right", "").replace(r"\quad", "    ")
    s = re.sub(r"\\bigg\|", "|", s)
    s = s.replace("{", "").replace("}", "")
    s = s.replace(r"\\", "  /  ")  # cases separator
    s = re.sub(r"\\[a-zA-Z]+", "", s)  # drop leftover commands
    s = s.replace("_", "_").replace("^", "^")
    return re.sub(r"\s+", " ", s).strip()


EQDIR = FIG / "_eq"
EQDIR.mkdir(exist_ok=True)
_eqn = [0]
plt.rcParams["mathtext.fontset"] = "cm"
plt.rcParams["font.family"] = "Hiragino Sans"


def render_eq(latex):
    """render a LaTeX equation to a tight PNG via mathtext; handle cases{}."""
    _eqn[0] += 1
    path = EQDIR / f"eq_{_eqn[0]}.png"
    if r"\begin{cases}" in latex:
        lhs = latex.split("=", 1)[0].strip()
        body = re.search(r"\\begin\{cases\}(.+?)\\end\{cases\}", latex, re.S).group(1)
        branches = [b.strip() for b in body.split(r"\\") if b.strip()]
        fig = plt.figure(figsize=(7.2, 0.55 * len(branches) + 0.4))
        fig.text(0.02, 0.5, f"${lhs} =$", ha="left", va="center", fontsize=17)
        y = 0.86
        for br in branches:
            expr, cond = (br.split("&") + [""])[:2]
            expr = expr.replace(r"\dfrac", r"\frac").strip()
            cond = re.sub(r"\\text\{(.+?)\}", r"\1", cond).strip()
            fig.text(0.30, y, f"${expr}$", ha="left", va="center", fontsize=15)
            fig.text(0.66, y, f"（{cond}）", ha="left", va="center", fontsize=11)
            y -= 1.0 / len(branches) * 0.92
    else:
        lx = (latex.replace(r"\dfrac", r"\frac").replace(r"\,", " ").replace(r"\!", "")
              .replace(r"\bigg", "").replace(r"\big", "")
              .replace(r"\le ", r"\leq ").replace(r"\ge ", r"\geq ")
              .replace(r"\text{-}", "-"))
        fig = plt.figure(figsize=(9, 0.9))
        try:
            fig.text(0.5, 0.5, f"${lx}$", ha="center", va="center", fontsize=17)
            fig.canvas.draw()   # force parse now so we can catch errors
        except Exception:
            plt.close(fig)
            fig = plt.figure(figsize=(9, 0.7))
            fig.text(0.5, 0.5, deftext(lx), ha="center", va="center", fontsize=13)
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.08, facecolor="white")
    plt.close(fig)
    return path


def deftext(s):
    """clean LaTeX tokens out of a variable-definition line for prose display."""
    s = re.sub(r"\\(mathcal|mathrm|mathbf|text|operatorname)\{([^{}]*)\}", r"\2", s)
    s = (s.replace(r"\Omega", "Ω").replace(r"\Phi", "Φ").replace(r"\rho", "ρ")
         .replace(r"\sigma", "σ").replace(r"\alpha", "α").replace(r"\omega", "ω")
         .replace(r"\le", "≤").replace(r"\ge", "≥").replace(r"\times", "×").replace(r"\cdot", "·"))
    s = re.sub(r"_\{([^{}]*)\}", r"_\1", s)
    s = re.sub(r"\^\{([^{}]*)\}", r"^\1", s)
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    return s.replace("{", "").replace("}", "")


def formula_box(caption, eq, defn):
    defn = deftext(defn)
    num, title = (caption.split("｜", 1) + [""])[:2] if "｜" in caption else (caption, "")
    # model-name banner
    p = para(before=6, after=0, keep=True); shade(p._element, TEAL)
    setfont(p.add_run("　" + (title or num)), GO, 9.5, True, WHITE)
    # equation image (centered), width capped to usable page width
    try:
        eqp = render_eq(eq)
        w_px, h_px = Image.open(eqp).size
        w_mm = min(160.0, w_px / 220 * 25.4)
        pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pe.paragraph_format.space_before = Pt(3); pe.paragraph_format.space_after = Pt(3)
        pe.paragraph_format.keep_with_next = True
        pe.add_run().add_picture(str(eqp), width=Mm(w_mm))
    except Exception as e:
        pe = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3, keep=True)
        setfont(pe.add_run(tex(eq)), "Cambria Math", 11)
    # variable definitions
    para(defn, MIN, 8, after=1, keep=True)
    # 図表番号 caption below
    pc = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=6)
    setfont(pc.add_run(f"{num.strip()}　{title.strip()}"), GO, 9, True)


def fit_picture(run, path, max_w=165.0, max_h=200.0):
    """add a picture scaled to fit within max_w x max_h (mm), preserving aspect."""
    w_px, h_px = Image.open(path).size
    if (max_w * h_px / w_px) <= max_h:
        run.add_picture(str(path), width=Mm(max_w))
    else:
        run.add_picture(str(path), height=Mm(max_h))


def figure(fname, caption, note):
    fp = FIG / fname
    if not fp.exists():
        para(f"［図表未生成: {fname}］", MIN, 8); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    fit_picture(p.add_run(), fp, max_w=160.0, max_h=190.0)
    pc = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1, keep=True)
    setfont(pc.add_run(caption), GO, 9, True)
    para(note, MIN, 8, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)


def set_cell(cell, text, name=MIN, size=8.5, bold=False, color=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    if text:
        setfont(p.add_run(text), name, size, bold, color)
    if fill:
        shade(cell._tc, fill)


def data_table(headers, rows, caption, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, GO, 9, True, WHITE, TEAL)
    for r in rows:
        cells = t.add_row().cells
        for j, v in enumerate(r):
            set_cell(cells[j], str(v))
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Mm(w)
    pc = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=6)
    setfont(pc.add_run(caption), GO, 9, True)


# ---- INSERT builders ------------------------------------------------------
def insert_portfolio_table():
    f = pd.read_csv(OUT / "phase4_portfolio_allocation/allocation_final.csv")
    order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
    f["__o"] = f.final_role.map({r: i for i, r in enumerate(order)})
    f = f.sort_values(["__o", "target_weight_final"], ascending=[True, False])
    rows = []
    for _, r in f.iterrows():
        rows.append([int(r.code_n), r.company_name[:18], r.final_role, r.theme,
                     f"{r.target_weight_final*100:.2f}%", f"{int(r.qty_L1)}",
                     f"¥{int(r.amount_L1):,}"])
    data_table(["コード", "企業名", "役割", "テーマ", "目標比率", "株数(L=1)", "購入金額(L=1)"],
               rows, "図表 III-1　投資企業ポートフォリオ（配分案C・単元未満株 L=1 基準）",
               widths=[14, 40, 30, 26, 16, 16, 26])
    para("注：総予算 500 万円、投資額 ¥4,949,198、残現金 ¥50,801（消化率 99.0%）。"
         "L=100（実単元）では株価の高い 9 社が購入不可となり消化率 46.7%（本文注2）。", MIN, 8, after=6)


def insert_meigara_template():
    f = pd.read_csv(OUT / "phase4_portfolio_allocation/allocation_final.csv")
    order = ["Buffett Core", "Transformation Core", "Emerging Core", "Dual Moat", "Bridge / Diversifier"]
    f["__o"] = f.final_role.map({r: i for i, r in enumerate(order)})
    f = f.sort_values(["__o", "target_weight_final"], ascending=[True, False])
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["企業名（コード）／業種／役割・テーマ", "企業概要（記入）", "Moat の根拠：変わる／生まれる（記入）"]):
        set_cell(t.rows[0].cells[j], h, GO, 9, True, WHITE, TEAL)
    for _, r in f.iterrows():
        cells = t.add_row().cells
        left = f"{r.company_name[:20]}（{int(r.code_n)}）\n{r.sector}\n{r.final_role}／{r.theme}"
        set_cell(cells[0], "", fill=TEAL_LT)
        pp = cells[0].paragraphs[0]
        setfont(pp.add_run(f"{r.company_name[:20]}（{int(r.code_n)}）"), GO, 8.5, True, TEALC)
        for line in [r.sector, f"{r.final_role}／{r.theme}"]:
            pl = cells[0].add_paragraph(); pl.paragraph_format.space_after = Pt(0)
            setfont(pl.add_run(line), MIN, 8)
        set_cell(cells[1], "")   # blank for submitter
        set_cell(cells[2], "")
    for r in t.rows:
        r.cells[0].width = Mm(42); r.cells[1].width = Mm(60); r.cells[2].width = Mm(66)
    pc = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=6)
    setfont(pc.add_run("図表 III-2　銘柄紹介（記入テンプレート：左列は自動記入、右2列を提出者が記入）"), GO, 9, True)


def insert_interview_template():
    p = para(before=4, after=2); shade(p._element, TEAL_LT); border(p, edges=("left",), color=TEAL, sz="30")
    setfont(p.add_run("　ご質問内容（記入）"), GO, 10, True, TEALC)
    for i in range(1, 4):
        para(f"{i}．［質問{i}を記入］", MIN, 9, after=1, hang=14)
    para("（下記は投資対象企業ごとの記入枠。インタビューは 2〜4 社を目安に記入。ZOOM 等の実施日時・担当者・参加者を記す。）",
         MIN, 8, before=2, after=4)
    for k in range(1, 4):
        t = doc.add_table(rows=4, cols=2); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell(t.rows[0].cells[0], f"企業名（記入）", GO, 9.5, True, WHITE, TEAL)
        set_cell(t.rows[0].cells[1], "", fill=TEAL)
        # merge header row
        t.rows[0].cells[0].merge(t.rows[0].cells[1])
        for ri, lab in enumerate(["日時", "ご担当者様", "参加者"], start=1):
            set_cell(t.rows[ri].cells[0], lab, GO, 9, True, TEALC, TEAL_LT)
            set_cell(t.rows[ri].cells[1], "")
        for r in t.rows:
            r.cells[0].width = Mm(28)
        for i in range(1, 4):
            para(f"{i}．［回答を記入］", MIN, 9, after=1, hang=14)
        para("", after=4)


def insert_learned_template():
    prompts = [
        "チームでの取り組みの経緯（結成・役割分担・打ち合わせの工夫など）を記入。",
        "スクリーニング（守・破・離）を進めるうえで難しかったこと・工夫したことを記入。",
        "「専門家同士の提携」ならぬ「三世代の Moat」というテーマに辿り着くまでの学びを記入。",
        "投資・企業分析・チームワークについて得た教訓を記入。",
        "（STOCK リーグは氏名・謝辞の記載が可能。必要に応じて指導教員・協力企業への謝辞を記入。）",
    ]
    for pr in prompts:
        para("・" + pr, MIN, 9, after=2, hang=10)
    para("［本文を記入］", MIN, 9, before=4, after=6, color=RGBColor(0x88, 0x88, 0x88))


def insert_references():
    refs = (P7 / "final_references.md").read_text(encoding="utf-8").splitlines()
    emit = False
    for rl in refs:
        r = rl.rstrip()
        if r.startswith("## 英語文献"):
            banner("英語文献", 3); emit = True; continue
        if r.startswith("## 日本語文献"):
            banner("日本語文献", 3); continue
        if r.startswith("## 本文引用"):
            break
        if emit and r.strip() and not r.startswith("#") and r.strip() != "---":
            para(re.sub(r"\*(.+?)\*", r"\1", r.strip()), MIN, 9, after=2, hang=18)


INSERTS = {
    "portfolio_table": insert_portfolio_table,
    "meigara_template": insert_meigara_template,
    "interview_template": insert_interview_template,
    "learned_template": insert_learned_template,
    "references": insert_references,
}

# ---- cover ---------------------------------------------------------------
md = (P7 / "final_report.md").read_text(encoding="utf-8").splitlines()
cover_meta = []
mtxt = "\n".join(md)
cm = re.search(r"<!-- COVERMETA(.+?)-->", mtxt, re.S)
if cm:
    cover_meta = [l.strip() for l in cm.group(1).strip().splitlines() if l.strip()]

for _ in range(4):
    doc.add_paragraph()
para("BEYOND BUFFETT", GO, 26, True, WD_ALIGN_PARAGRAPH.CENTER, after=6, color=TEALC)
para("― Moat の時間軸拡張による日本株ポートフォリオ ―", MIN, 13, True, WD_ALIGN_PARAGRAPH.CENTER, after=3)
para('"INVESTING IN COMPLETED, CHANGING, AND EMERGING MOATS"', MIN, 11, False,
     WD_ALIGN_PARAGRAPH.CENTER, after=48)
for line in cover_meta:
    para(line, MIN, 11, False, WD_ALIGN_PARAGRAPH.RIGHT, after=3)
doc.add_page_break()

# ---- body parse ----------------------------------------------------------
i = 0
in_ref = False
while i < len(md):
    s = md[i].rstrip()
    # skip comments except FIG/INSERT/COVERMETA already handled
    fig = re.match(r"<!-- FIG:\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*-->", s)
    ins = re.match(r"<!-- INSERT:\s*(\w+)\s*-->", s)
    if s.startswith("<!--") and not (fig or ins):
        # skip whole comment block if multi-line
        while i < len(md) and "-->" not in md[i]:
            i += 1
        i += 1; continue
    if fig:
        figure(fig.group(1), fig.group(2), fig.group(3)); i += 1; continue
    if ins:
        fn = INSERTS.get(ins.group(1))
        if fn:
            fn()
        i += 1; continue
    if s.startswith("```formula"):
        cap = eq = defn = ""
        i += 1
        while i < len(md) and not md[i].startswith("```"):
            ln = md[i]
            if ln.startswith("CAPTION:"):
                cap = ln[8:].strip()
            elif ln.startswith("EQ:"):
                eq = ln[3:].strip()
            elif ln.startswith("DEF:"):
                defn = ln[4:].strip()
            i += 1
        i += 1
        formula_box(cap, eq, defn); continue
    if s.strip() == "---":
        i += 1; continue
    if s.startswith("# "):
        i += 1; continue  # title on cover
    if s.startswith("## ") and s[3:].strip() in ("― Moat の時間軸拡張による日本株ポートフォリオ ―",):
        i += 1; continue
    if s.startswith("### ") and "INVESTING" in s:
        i += 1; continue
    m1 = re.match(r"^## (.+)", s)
    if m1:
        title = m1.group(1).strip()
        in_ref = title.startswith("参考文献")
        if title[:1] in "ⅠⅡⅢⅣⅤ" or title.startswith("参考文献"):
            doc.add_page_break()
        banner(title, 1); i += 1; continue
    m2 = re.match(r"^### (.+)", s)
    if m2:
        banner(m2.group(1).strip(), 2); i += 1; continue
    m3 = re.match(r"^#### (.+)", s)
    if m3:
        banner(m3.group(1).strip(), 3); i += 1; continue
    if not s.strip():
        i += 1; continue
    txt = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    txt = re.sub(r"`(.+?)`", r"\1", txt)
    txt = re.sub(r"^-\s+", "・", txt)
    para(txt, MIN, 9.5, after=5, line=1.35)
    i += 1

doc.save(str(P7 / "final_report.docx"))
print("saved", P7 / "final_report.docx", "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
